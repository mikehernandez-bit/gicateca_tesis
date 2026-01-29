
# ==========================================
# 4. UTILIDADES RESTAURADAS (Para compatibilidad con Maestría/Posgrado)
# ==========================================
def encontrar_recurso(nombre_archivo):
    """Busca un recurso en BASE_DIR o STATIC_ASSETS_DIR."""
    candidatos = [
        os.path.join(BASE_DIR, nombre_archivo),
        os.path.join(STATIC_ASSETS_DIR, nombre_archivo),
        os.path.join(BASE_DIR, "assets", nombre_archivo)
    ]
    for ruta in candidatos:
        if os.path.exists(ruta):
            return ruta
    return None

def agregar_nota_guia(doc, texto):
    """Agrega una nota tipo guía (estilo azul/gris)."""
    if not texto: return
    try:
        # Crea una tabla de 1 celda para el borde/fondo
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Cm(15.0)  # Ancho fijo aprox
        
        cell = table.cell(0, 0)
        # Fondo gris claro F2F8FD
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F2F8FD')
        tc_pr.append(shd)
        
        # Párrafo dentro
        cell.text = texto
        for p in cell.paragraphs:
            p.style = None
            if p.runs:
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    except Exception as e:
        print(f"Error agregando nota guia: {e}")

def crear_tabla_estilizada(doc, data):
    """Crea una tabla bonita desde JSON {headers: [], rows: []}."""
    if not data: return
    headers = data.get('headers', [])
    rows = data.get('rows', [])
    if not headers: return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)

    # Filas
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = str(val)
                for p in row_cells[i].paragraphs:
                   if p.runs:
                       p.runs[0].font.size = Pt(10)
