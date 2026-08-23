from docx import Document

import app.engine.renderers  # noqa: F401
from app.engine.normalizer import normalize
from app.engine.registry import render_block, render_blocks


def test_canonical_placeholder_asset_renders_caption_and_image() -> None:
    doc = Document()
    render_block(
        doc,
        {
            "type": "image",
            "titulo": "Arquitectura del sistema",
            "ruta": "assets/placeholder_figura.png",
            "fuente": "Placeholder tecnico controlado",
        },
    )

    texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    assert any("Arquitectura del sistema" in text for text in texts)
    assert len(doc.inline_shapes) == 1


def test_document_derived_abbreviations_drop_unreferenced_template_examples() -> None:
    data = {
        "_meta": {"id": "test-format", "university": "unac"},
        "configuracion": {"ruta_logo": "app/static/assets/LogoUNAC.png"},
        "caratula": {
            "universidad": "UNIVERSIDAD TEST",
            "facultad": "FACULTAD TEST",
            "titulo_placeholder": "TITULO TEST",
        },
        "preliminares": {
            "indices": [
                {
                    "titulo": "INDICE DE ABREVIATURAS",
                    "items": [{"texto": "OMS: Organizacion Mundial de la Salud"}],
                }
            ]
        },
        "cuerpo": [
            {
                "titulo": "IV. METODOLOGIA",
                "contenido": [
                    {
                        "texto": "4.1 Diseno metodologico",
                        "_ai_content": (
                            "La IA aplicada al IoT mejora los KPI operativos y el analisis se proceso con SPSS."
                        ),
                    }
                ],
            }
        ],
    }

    blocks = normalize(data)
    abbr_tables = [block for block in blocks if block["type"] == "abbreviations_table"]
    assert len(abbr_tables) == 1

    rows = abbr_tables[0]["rows"]
    siglas = {row["sigla"] for row in rows}
    assert {"IA", "IoT", "SPSS", "KPI"}.issubset(siglas)
    assert "OMS" not in siglas


def test_abbreviation_extraction_rejects_authors_years_and_common_words() -> None:
    data = {
        "_meta": {"id": "test-format", "university": "unac"},
        "configuracion": {"ruta_logo": "app/static/assets/LogoUNAC.png"},
        "caratula": {
            "universidad": "UNIVERSIDAD TEST",
            "facultad": "FACULTAD TEST",
            "titulo_placeholder": "TITULO TEST",
        },
        "preliminares": {
            "indices": {"abreviaturas": "INDICE DE ABREVIATURAS"},
        },
        "cuerpo": [
            {
                "titulo": "II. MARCO TEORICO",
                "contenido": [
                    {
                        "texto": "2.1 Bases teoricas",
                        "_ai_content": (
                            "Internet de las Cosas (IoT), Objetivos de Desarrollo Sostenible (ODS), "
                            "Society of Automotive Engineers (SAE), Universidad Nacional del Callao (UNAC), "
                            "Failure Mode and Effects Analysis (FMEA), Mean Time Between Failures (MTBF) "
                            "y Statistical Package for the Social Sciences (SPSS) se usan en el estudio. "
                            "Mobley (2002), Mendoza (2021), Chen (2019) y palabras como fallos, datos o mantenimiento "
                            "no deben entrar como siglas."
                        ),
                    }
                ],
            }
        ],
    }

    blocks = normalize(data)
    abbr_tables = [block for block in blocks if block["type"] == "abbreviations_table"]
    assert len(abbr_tables) == 1

    rows = abbr_tables[0]["rows"]
    mapping = {row["sigla"]: row["meaning"] for row in rows}
    assert mapping["IoT"] == "Internet de las Cosas"
    assert mapping["ODS"] == "Objetivos de Desarrollo Sostenible"
    assert mapping["SAE"] == "Society of Automotive Engineers"
    assert mapping["UNAC"] == "Universidad Nacional del Callao"
    assert mapping["FMEA"] == "Failure Mode and Effects Analysis"
    assert mapping["MTBF"] == "Tiempo Medio Entre Fallas"
    assert mapping["SPSS"] == "Statistical Package for the Social Sciences"

    forbidden = {"COSAS", "MOBLEY", "MENDOZA", "CHEN", "FALLOS", "DATOS", "2021", "2002", "2019"}
    assert forbidden.isdisjoint({sigla.upper() for sigla in mapping})


def test_maintenance_abbreviations_scan_tables_captions_and_are_sorted_in_spanish() -> None:
    data = {
        "_meta": {"id": "unac-proyecto-cuant", "university": "unac"},
        "caratula": {
            "universidad": "UNIVERSIDAD NACIONAL DEL CALLAO",
            "titulo": "PLAN RCM PARA MOTONIVELADORAS CAT 24M",
        },
        "preliminares": {
            "indices": {"abreviaturas": "INDICE DE ABREVIATURAS"},
        },
        "cuerpo": [
            {
                "titulo": "II. MARCO TEORICO",
                "contenido": [
                    {
                        "texto": "2.2 Bases teoricas",
                        "_ai_content": "El RCM se evalua mediante MTBF y MTTR.",
                    },
                    {
                        "tipo": "tabla",
                        "titulo": "Tabla 3.1 Indicadores AMEF de la flota CAT 24M",
                        "encabezados": ["AMEF", "MTBF", "MTTR"],
                        "filas": [["RCM", "12", "3"]],
                    },
                ],
            }
        ],
    }

    rows = next(block["rows"] for block in normalize(data) if block["type"] == "abbreviations_table")
    mapping = {row["sigla"]: row["meaning"] for row in rows}

    assert list(mapping) == sorted(mapping)
    assert mapping["AMEF"] == "Análisis de Modo y Efecto de Falla"
    assert mapping["CAT"] == "Caterpillar"
    assert mapping["MTBF"] == "Tiempo Medio Entre Fallas"
    assert mapping["MTTR"] == "Tiempo Medio Para Reparar"
    assert mapping["RCM"] == "Mantenimiento Centrado en Confiabilidad"


def test_cited_corporate_author_is_added_to_abbreviation_index() -> None:
    data = {
        "_meta": {"id": "unac-proyecto-cuant", "university": "unac"},
        "preliminares": {
            "indices": {"abreviaturas": "INDICE DE ABREVIATURAS"},
        },
        "cuerpo": [
            {
                "titulo": "I. PLANTEAMIENTO DEL PROBLEMA",
                "contenido": [
                    {
                        "texto": "1.1 Realidad problemática",
                        "_ai_content": "La guía sustenta el análisis [[CITE:SRC_GMG_2020]].",
                    }
                ],
            }
        ],
        "finales": {
            "referencias": {
                "_ai_content": (
                    "[[SOURCE:SRC_GMG_2020]] GMG. (2020). Guía técnica de "
                    "mantenimiento de equipos. Technical Maintenance Review, 8(2), 1-20."
                )
            }
        },
    }

    rows = next(block["rows"] for block in normalize(data) if block["type"] == "abbreviations_table")

    assert {row["sigla"] for row in rows} == {"GMG"}
    assert rows[0]["meaning"] == "Global Mining Guidelines Group"


def test_annex_image_uses_annex_heading_without_figure_caption() -> None:
    data = {
        "_meta": {"id": "test-format", "university": "unac"},
        "configuracion": {"ruta_logo": "app/static/assets/LogoUNAC.png"},
        "caratula": {
            "universidad": "UNIVERSIDAD TEST",
            "facultad": "FACULTAD TEST",
            "titulo_placeholder": "TITULO TEST",
        },
        "cuerpo": [{"titulo": "CAP I"}],
        "finales": {
            "anexos": {
                "titulo_seccion": "ANEXOS",
                "lista": [
                    {
                        "titulo": "Anexo 2: Diagrama de flujo del proceso de recoleccion de datos",
                        "_ai_content": [
                            {
                                "tipo": "figura",
                                "titulo": "Diagrama de flujo del proceso de recoleccion de datos",
                                "caption": "Figura 15. Diagrama de flujo del proceso de recoleccion de datos",
                                "ruta_placeholder": "assets/placeholder_figura.png",
                                "fuente": "Placeholder tecnico controlado",
                            }
                        ],
                    }
                ],
            }
        },
    }

    blocks = normalize(data)
    annex_heading = [
        block
        for block in blocks
        if block["type"] == "paragraph_bold"
        and "Anexo 2: Diagrama de flujo del proceso de recoleccion de datos" in block["text"]
    ]
    assert annex_heading

    image_blocks = [block for block in blocks if block["type"] == "image"]
    assert len(image_blocks) == 1
    assert image_blocks[0]["omit_caption"] is True

    doc = Document()
    render_blocks(doc, blocks)
    texts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    assert "Anexo 2: Diagrama de flujo del proceso de recoleccion de datos" in texts
    assert not any(text.startswith("Figura ") for text in texts)
    assert len(doc.inline_shapes) >= 2
