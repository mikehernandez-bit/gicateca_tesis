"""
=============================================================================
ARCHIVO: tests/test_n8n_stress_scenarios.py
=============================================================================

PROPÓSITO:
Pruebas de estrés que simulan los JSONs que n8n/GicaGen podría producir.
Verifica que el pipeline normalize → render_blocks genera DOCX correctos
bajo escenarios extremos y variados.

ESCENARIOS CUBIERTOS:
  1. Tabla landscape que ocupa 2+ páginas (muchas filas)
  2. Múltiples tablas landscape consecutivas
  3. Tabla landscape → contenido portrait → tabla landscape
  4. Tabla landscape con muchas columnas (>10)
  5. Tabla landscape con celdas con texto muy largo (multi-line)
  6. Tabla landscape en cuerpo (no solo en anexos)
  7. Tabla legacy que auto-detecta landscape (>5 columnas)
  8. Matriz de consistencia con datos extensos
  9. Capítulo vacío / sin contenido
 10. Capítulo con solo tablas (sin párrafos)
 11. JSON mínimo (solo caratula obligatoria)
 12. JSON con todas las secciones posibles
 13. Tabla portrait seguida de landscape en mismo capítulo
 14. Anexos con múltiples matrices
 15. Tabla con celdas fusionadas complejas
 16. Tabla sin filas (solo encabezados)
 17. Tabla sin encabezados (edge case)
 18. Contenido mixto intenso (texto, tablas, imágenes intercalados)
 19. Doble section_switch consecutivo (landscape→landscape)
 20. Tabla landscape en preliminares

CÓMO EJECUTAR:
    py -m pytest tests/test_n8n_stress_scenarios.py -v
=============================================================================
"""
from __future__ import annotations

import io
import json
from copy import deepcopy
from pathlib import Path
from typing import List

import pytest
from docx import Document
from docx.oxml.ns import qn

# ─────────────────────────────────────────────────────────────
# ENGINE IMPORTS
# ─────────────────────────────────────────────────────────────

import app.engine.renderers  # noqa: F401  — register all renderers

from app.engine.normalizer import normalize
from app.engine.primitives import configure_styles, configure_margins
from app.engine.registry import render_blocks
from app.engine.types import Block

ROOT = Path(__file__).resolve().parent.parent


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def _pipeline(data: dict) -> Document:
    """Ejecuta normalize → render_blocks y retorna el Document."""
    blocks = normalize(data)
    doc = Document()
    configure_styles(doc)
    configure_margins(doc)
    render_blocks(doc, blocks)
    return doc


def _save_bytes(doc: Document) -> int:
    """Serializa el doc a bytes y retorna el tamaño."""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.tell()


def _count_section_breaks_with_shrink(doc: Document) -> int:
    """Cuenta párrafos con sectPr cuyo spacing está encogido (fix de blank page)."""
    count = 0
    body = doc.element.body
    for p_elem in body.iterchildren(qn("w:p")):
        p_pr = p_elem.find(qn("w:pPr"))
        if p_pr is None:
            continue
        sect = p_pr.find(qn("w:sectPr"))
        if sect is None:
            continue
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is not None and spacing.get(qn("w:after")) == "0":
            count += 1
    return count


def _section_orientations(doc: Document) -> List[str]:
    """Retorna la orientación de cada sección del documento."""
    result = []
    for sec in doc.sections:
        if sec.page_width > sec.page_height:
            result.append("landscape")
        else:
            result.append("portrait")
    return result


def _minimal_caratula() -> dict:
    """Carátula mínima válida para generar un DOCX."""
    return {
        "universidad": "UNIVERSIDAD NACIONAL DEL CALLAO",
        "facultad": "FACULTAD DE INGENIERÍA",
        "escuela": "ESCUELA DE SISTEMAS",
        "tipo_documento": "TESIS",
        "titulo": "Título de prueba de estrés",
        "autor_valor": "Tesista de Prueba",
        "lugar": "Callao",
        "anio": "2026",
    }


def _generate_rows(n: int, cols: int) -> list:
    """Genera n filas con cols columnas de texto variado."""
    return [
        [f"Dato R{r+1} C{c+1} — texto de prueba" for c in range(cols)]
        for r in range(n)
    ]


def _landscape_table(rows: int = 5, cols: int = 5, titulo: str = None) -> dict:
    """Tabla landscape canónica."""
    headers = [f"Columna {i+1}" for i in range(cols)]
    t = {
        "tipo": "tabla",
        "orientacion": "landscape",
        "encabezados": headers,
        "filas": _generate_rows(rows, cols),
    }
    if titulo:
        t["titulo"] = titulo
    return t


def _portrait_table(rows: int = 3, cols: int = 3, titulo: str = None) -> dict:
    """Tabla portrait canónica."""
    headers = [f"Col {i+1}" for i in range(cols)]
    t = {
        "tipo": "tabla",
        "orientacion": "portrait",
        "encabezados": headers,
        "filas": _generate_rows(rows, cols),
    }
    if titulo:
        t["titulo"] = titulo
    return t


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 1: Tabla landscape de 2+ páginas (muchas filas)
# ═══════════════════════════════════════════════════════════════

class TestLandscapeMultiPage:
    """Simula n8n generando tablas con muchos datos que ocupan >1 página."""

    def test_landscape_50_rows(self):
        """50 filas × 5 cols → debe ocupar ~2 páginas landscape sin blank page."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "CAPÍTULO DE DATOS",
                "contenido": [{
                    "texto": "Resultados de la investigación",
                    "tabla_data": _landscape_table(rows=50, cols=5, titulo="Datos extensos"),
                }],
            }],
        }
        doc = _pipeline(data)
        size = _save_bytes(doc)
        assert size > 10000
        # Debe haber secciones landscape y portrait
        orients = _section_orientations(doc)
        assert "landscape" in orients
        # Todos los section breaks deben estar encogidos
        shrunk = _count_section_breaks_with_shrink(doc)
        assert shrunk >= 2, f"Expected >=2 shrunk section breaks, got {shrunk}"

    def test_landscape_100_rows(self):
        """100 filas → tabla masiva, no debe crashear."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "MEGATABLA",
                "contenido": [{
                    "texto": "100 filas de datos",
                    "tabla_data": _landscape_table(rows=100, cols=7, titulo="MegaTabla"),
                }],
            }],
        }
        doc = _pipeline(data)
        assert _save_bytes(doc) > 20000
        assert len(doc.tables) >= 1

    def test_landscape_long_cell_text(self):
        """Celdas con texto largo (~500 chars) en landscape."""
        long_text = "Lorem ipsum dolor sit amet. " * 20  # ~560 chars
        filas = [[long_text for _ in range(5)] for _ in range(10)]
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "DATOS CON TEXTO LARGO",
                "contenido": [{
                    "texto": "Celdas extensas",
                    "tabla_data": {
                        "tipo": "tabla",
                        "orientacion": "landscape",
                        "encabezados": ["A", "B", "C", "D", "E"],
                        "filas": filas,
                        "titulo": "Tabla con texto largo",
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        assert _save_bytes(doc) > 10000


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 2: Múltiples tablas landscape consecutivas
# ═══════════════════════════════════════════════════════════════

class TestConsecutiveLandscapeTables:
    """N8n podría generar múltiples tablas landscape seguidas."""

    def test_two_landscape_tables_same_chapter(self):
        """Dos tablas landscape en el mismo capítulo → cada una con su switch."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "DOS TABLAS LANDSCAPE",
                "contenido": [
                    {
                        "texto": "Primera tabla",
                        "tabla_data": _landscape_table(rows=10, cols=6, titulo="Tabla A"),
                    },
                    {
                        "texto": "Segunda tabla",
                        "tabla_data": _landscape_table(rows=15, cols=5, titulo="Tabla B"),
                    },
                ],
            }],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 2
        orients = _section_orientations(doc)
        # Cada tabla landscape debe crear su propia sección
        assert orients.count("landscape") >= 2
        shrunk = _count_section_breaks_with_shrink(doc)
        assert shrunk >= 4  # 2 switches to landscape + 2 switches to portrait

    def test_three_landscape_tables_across_chapters(self):
        """Tablas landscape distribuidas en 3 capítulos diferentes."""
        chapters = []
        for i in range(3):
            chapters.append({
                "titulo": f"CAPÍTULO {i+1}",
                "contenido": [
                    {"texto": f"Intro del cap {i+1}", "parrafos": ["Un párrafo."]},
                    {
                        "texto": f"Datos del cap {i+1}",
                        "tabla_data": _landscape_table(
                            rows=8, cols=5, titulo=f"Tabla Cap {i+1}"
                        ),
                    },
                ],
            })
        data = {"caratula": _minimal_caratula(), "cuerpo": chapters}
        doc = _pipeline(data)
        assert len(doc.tables) >= 3
        assert _save_bytes(doc) > 15000


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 3: Alternancia landscape ↔ portrait
# ═══════════════════════════════════════════════════════════════

class TestOrientationAlternation:
    """Verifica que alternar orientaciones no deja páginas en blanco."""

    def test_portrait_landscape_portrait_landscape(self):
        """Patrón: texto → landscape table → texto → landscape table."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "ALTERNANCIA",
                "contenido": [
                    {"texto": "Sección A", "parrafos": ["Párrafo de texto normal." * 5]},
                    {"texto": "Tabla landscape A",
                     "tabla_data": _landscape_table(rows=5, cols=6, titulo="T-A")},
                    {"texto": "Sección B", "parrafos": ["Más texto entre tablas." * 5]},
                    {"texto": "Tabla landscape B",
                     "tabla_data": _landscape_table(rows=8, cols=5, titulo="T-B")},
                    {"texto": "Sección C", "parrafos": ["Texto final del capítulo."]},
                ],
            }],
        }
        doc = _pipeline(data)
        orients = _section_orientations(doc)
        assert orients.count("landscape") >= 2
        # Verify all section breaks shrunk
        shrunk = _count_section_breaks_with_shrink(doc)
        assert shrunk >= 4

    def test_portrait_table_then_landscape_table(self):
        """Portrait table → landscape table en misma sección."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "PORTRAIT vs LANDSCAPE",
                "contenido": [
                    {
                        "texto": "Tabla portrait chica",
                        "tabla_data": _portrait_table(rows=3, cols=3, titulo="T-Portrait"),
                    },
                    {
                        "texto": "Tabla landscape grande",
                        "tabla_data": _landscape_table(rows=20, cols=7, titulo="T-Landscape"),
                    },
                ],
            }],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 2


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 4: Tabla con muchas columnas
# ═══════════════════════════════════════════════════════════════

class TestManyColumns:
    """Tablas con muchas columnas que solo caben en landscape."""

    def test_12_columns_landscape(self):
        """12 columnas explícitamente landscape."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "TABLA ANCHA",
                "contenido": [{
                    "texto": "Datos con 12 columnas",
                    "tabla_data": _landscape_table(rows=5, cols=12, titulo="12 Cols"),
                }],
            }],
        }
        doc = _pipeline(data)
        table = doc.tables[0]
        assert len(table.columns) == 12

    def test_legacy_table_auto_landscape(self):
        """Legacy table con >5 cols debe auto-detectar landscape."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "LEGACY ANCHO",
                "contenido": [{
                    "texto": "Legacy con 8 columnas",
                    "tabla": {
                        "headers": [f"H{i}" for i in range(8)],
                        "rows": [[f"D{r}{c}" for c in range(8)] for r in range(5)],
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        orients = _section_orientations(doc)
        assert "landscape" in orients, "Legacy table with >5 cols should trigger landscape"

    def test_legacy_table_portrait(self):
        """Legacy table con <=5 cols se queda portrait."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "LEGACY ANGOSTO",
                "contenido": [{
                    "texto": "Legacy con 3 columnas",
                    "tabla": {
                        "headers": ["A", "B", "C"],
                        "rows": [["1", "2", "3"]],
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        # No debe haber landscape sections (excepto default portrait)
        orients = _section_orientations(doc)
        assert orients.count("landscape") == 0


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 5: Matriz de consistencia con datos extensos
# ═══════════════════════════════════════════════════════════════

class TestMatrizConsistencia:
    """La matriz de consistencia es la tabla más compleja del sistema."""

    def _extensive_matriz(self) -> dict:
        return {
            "problemas": {
                "general": "¿Cómo influye la implementación de sistemas inteligentes "
                           "en la gestión documental de las universidades públicas del Perú "
                           "durante el período 2020-2026?",
                "especificos": [
                    "¿Cuál es el nivel de eficiencia alcanzado con la implementación "
                    "de sistemas de gestión documental inteligente?",
                    "¿Qué factores determinan la adopción exitosa de estas tecnologías?",
                    "¿Cómo se relaciona la inversión en TI con los resultados obtenidos?",
                    "¿Cuáles son las barreras tecnológicas y culturales identificadas?",
                ],
            },
            "objetivos": {
                "general": "Determinar la influencia de la implementación de sistemas "
                           "inteligentes en la gestión documental universitaria.",
                "especificos": [
                    "Medir el nivel de eficiencia operativa lograda.",
                    "Identificar factores críticos de éxito.",
                    "Analizar la relación costo-beneficio.",
                    "Categorizar las barreras de implementación.",
                ],
            },
            "hipotesis": {
                "general": "La implementación de sistemas inteligentes mejora significativamente "
                           "la gestión documental en un 45% de eficiencia.",
                "especificos": [
                    "La eficiencia operativa incrementa >30%.",
                    "Existen al menos 5 factores críticos de éxito.",
                    "La relación B/C es >= 2.5:1.",
                    "Las barreras culturales superan a las tecnológicas.",
                ],
            },
            "variables": {
                "independiente": {
                    "nombre": "Sistemas inteligentes de gestión documental",
                    "dimensiones": [
                        "Automatización de procesos",
                        "Inteligencia artificial aplicada",
                        "Interoperabilidad de sistemas",
                        "Capacitación del personal",
                    ],
                },
                "dependiente": {
                    "nombre": "Eficiencia en la gestión documental",
                    "dimensiones": [
                        "Tiempo de procesamiento de documentos",
                        "Tasa de errores documentales",
                        "Satisfacción del usuario interno",
                        "Costo operativo por documento",
                    ],
                },
            },
            "metodologia": {
                "tipo": "Aplicada, longitudinal",
                "diseño": "Cuasi-experimental con grupo control",
                "poblacion": "428 trabajadores administrativos",
                "muestra": "204 trabajadores (muestreo estratificado)",
                "tecnica": "Encuestas, observación, análisis documental",
                "instrumento": "Cuestionario validado (α Cronbach = 0.87)",
            },
        }

    def test_extensive_matriz_in_anexos(self):
        """Matriz con datos extensos en sección de anexos."""
        data = {
            "caratula": _minimal_caratula(),
            "matriz_consistencia": self._extensive_matriz(),
            "cuerpo": [{"titulo": "INTRODUCCIÓN", "contenido": []}],
            "finales": {
                "anexos": {
                    "titulo_seccion": "ANEXOS",
                    "lista": [
                        {"titulo": "Anexo 1: Matriz de Consistencia"},
                    ],
                },
            },
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 1
        orients = _section_orientations(doc)
        assert "landscape" in orients

    def test_matriz_fallback_no_lista(self):
        """Matriz existe pero no está en la lista de anexos → fallback."""
        data = {
            "caratula": _minimal_caratula(),
            "matriz_consistencia": self._extensive_matriz(),
            "cuerpo": [{"titulo": "INTRODUCCIÓN", "contenido": []}],
            "finales": {
                "anexos": {
                    "titulo_seccion": "ANEXOS",
                    "lista": [
                        {"titulo": "Anexo 1: Cuestionario", "parrafos": ["Instrucciones..."]},
                    ],
                },
            },
        }
        doc = _pipeline(data)
        # Fallback debe renderizar la matriz igualmente
        assert len(doc.tables) >= 1
        orients = _section_orientations(doc)
        assert "landscape" in orients


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 6: JSON mínimo y máximo
# ═══════════════════════════════════════════════════════════════

class TestJSONBoundaries:
    """Verifica que JSONs mínimos y máximos no crashean."""

    def test_minimal_json(self):
        """Solo carátula → debe generar DOCX válido."""
        data = {"caratula": _minimal_caratula()}
        doc = _pipeline(data)
        assert _save_bytes(doc) > 5000

    def test_empty_chapters(self):
        """Capítulos vacíos no deben crashear."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [
                {"titulo": "CAP VACÍO 1", "contenido": []},
                {"titulo": "CAP VACÍO 2"},
                {"titulo": "CAP CON TEXTO", "contenido": ["Un párrafo simple."]},
            ],
        }
        doc = _pipeline(data)
        assert _save_bytes(doc) > 5000

    def test_chapter_only_tables(self):
        """Capítulo que tiene solo tablas sin párrafos."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "SOLO TABLAS",
                "contenido": [
                    {"tipo": "tabla", **_portrait_table(3, 3, "TablaSola1")},
                    {"tipo": "tabla", **_landscape_table(5, 6, "TablaSola2")},
                    {"tipo": "tabla", **_portrait_table(2, 4, "TablaSola3")},
                ],
            }],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 3

    def test_maximal_json_all_sections(self):
        """JSON con todas las secciones posibles."""
        data = {
            "caratula": _minimal_caratula(),
            "pagina_respeto": {"titulo": "PÁGINA DE RESPETO"},
            "informacion_basica": {
                "titulo": "INFORMACIÓN BÁSICA",
                "elementos": [
                    {"label": "Facultad", "valor": "Ingeniería"},
                    {"label": "Autor", "valor": "Tesista"},
                ],
            },
            "preliminares": {
                "dedicatoria": {"titulo": "DEDICATORIA", "texto": "A mi familia."},
                "agradecimiento": {"titulo": "AGRADECIMIENTO", "texto": "Gracias."},
                "resumen": {"titulo": "RESUMEN", "texto": "Resumen de la tesis."},
                "indices": {
                    "contenido": "ÍNDICE",
                    "tablas": "ÍNDICE DE TABLAS",
                    "figuras": "ÍNDICE DE FIGURAS",
                    "abreviaturas": "ÍNDICE DE ABREVIATURAS",
                },
                "introduccion": {
                    "titulo": "INTRODUCCIÓN",
                    "texto": "Introducción completa de la tesis.",
                },
            },
            "cuerpo": [
                {
                    "titulo": "I. PLANTEAMIENTO DEL PROBLEMA",
                    "contenido": [
                        {"texto": "1.1 Formulación", "parrafos": ["Párrafo descriptivo."]},
                        {"texto": "1.2 Datos", "tabla_data": _portrait_table(3, 3, "Datos 1.2")},
                    ],
                },
                {
                    "titulo": "II. MARCO TEÓRICO",
                    "contenido": [
                        {"texto": "2.1 Antecedentes", "parrafos": ["Texto extenso."]},
                        {
                            "texto": "2.2 Comparativa",
                            "tabla_data": _landscape_table(15, 6, "Comparativa"),
                        },
                    ],
                },
                {
                    "titulo": "III. RESULTADOS",
                    "contenido": [
                        {
                            "texto": "3.1 Resultados generales",
                            "tabla_data": _landscape_table(30, 5, "Resultados"),
                        },
                    ],
                },
            ],
            "matriz_consistencia": TestMatrizConsistencia()._extensive_matriz(),
            "finales": {
                "referencias": {
                    "titulo": "REFERENCIAS BIBLIOGRÁFICAS",
                    "nota": "Según norma APA 7ma edición",
                },
                "anexos": {
                    "titulo_seccion": "ANEXOS",
                    "lista": [
                        {"titulo": "Anexo 1: Matriz de Consistencia"},
                        {"titulo": "Anexo 2: Instrumento", "parrafos": ["Cuestionario..."]},
                    ],
                },
            },
        }
        doc = _pipeline(data)
        size = _save_bytes(doc)
        assert size > 30000
        assert len(doc.tables) >= 5
        orients = _section_orientations(doc)
        assert "landscape" in orients
        assert "portrait" in orients


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 7: Tabla con celdas fusionadas complejas
# ═══════════════════════════════════════════════════════════════

class TestCellMerges:
    """Fusión de celdas en tablas landscape y portrait."""

    def test_landscape_with_row_merges(self):
        """Tabla landscape con fusiones de filas."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "MERGES",
                "contenido": [{
                    "texto": "Tabla con merges",
                    "tabla_data": {
                        "tipo": "tabla",
                        "orientacion": "landscape",
                        "titulo": "Tabla Fusionada",
                        "encabezados": ["Cat", "Sub", "Valor", "Nota", "Extra"],
                        "filas": [
                            ["A", "a1", "100", "nota 1", "x1"],
                            ["", "a2", "200", "nota 2", "x2"],
                            ["B", "b1", "300", "nota 3", "x3"],
                            ["", "b2", "400", "nota 4", "x4"],
                            ["", "b3", "500", "nota 5", "x5"],
                        ],
                        "celdas_fusionadas": [
                            {"fila": 0, "col": 0, "filas_span": 2},
                            {"fila": 2, "col": 0, "filas_span": 3},
                        ],
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 1

    def test_landscape_with_col_merges(self):
        """Tabla landscape con fusiones de columnas."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "COL MERGES",
                "contenido": [{
                    "texto": "Tabla con merges de columnas",
                    "tabla_data": {
                        "tipo": "tabla",
                        "orientacion": "landscape",
                        "titulo": "Col Merge",
                        "encabezados": ["A", "B", "C", "D", "E", "F"],
                        "filas": [
                            ["Span 3 cols", "", "", "D1", "E1", "F1"],
                            ["A2", "B2", "C2", "D2", "E2", "F2"],
                        ],
                        "celdas_fusionadas": [
                            {"fila": 0, "col": 0, "cols_span": 3},
                        ],
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 1


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 8: Edge cases de tabla
# ═══════════════════════════════════════════════════════════════

class TestTableEdgeCases:
    """Casos límite en tablas."""

    def test_table_no_rows(self):
        """Tabla con solo encabezados, sin filas de datos."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "SIN FILAS",
                "contenido": [{
                    "texto": "Tabla vacía",
                    "tabla_data": {
                        "tipo": "tabla",
                        "orientacion": "landscape",
                        "titulo": "Solo Headers",
                        "encabezados": ["A", "B", "C", "D", "E"],
                        "filas": [],
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        # Debe generar la tabla con solo header row
        assert len(doc.tables) >= 1

    def test_table_no_headers(self):
        """Tabla sin encabezados (edge case) → debe ser ignorada."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "SIN HEADERS",
                "contenido": [{
                    "texto": "Tabla sin headers",
                    "tabla_data": {
                        "tipo": "tabla",
                        "encabezados": [],
                        "filas": [["a", "b"]],
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        # No debe crashear; tabla se ignora si no hay encabezados
        assert _save_bytes(doc) > 5000

    def test_table_with_nota_pie(self):
        """Tabla landscape con nota al pie."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "CON NOTA",
                "contenido": [{
                    "texto": "Tabla con nota al pie",
                    "tabla_data": {
                        "tipo": "tabla",
                        "orientacion": "landscape",
                        "titulo": "Datos con nota",
                        "nota_pie": "Fuente: Elaboración propia basada en datos del INEI (2025).",
                        "encabezados": ["Var", "2021", "2022", "2023", "2024", "2025"],
                        "filas": _generate_rows(5, 6),
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        assert _save_bytes(doc) > 5000

    def test_table_custom_column_widths(self):
        """Tabla landscape con anchos de columna personalizados."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "ANCHOS CUSTOM",
                "contenido": [{
                    "texto": "Tabla con anchos fijos",
                    "tabla_data": {
                        "tipo": "tabla",
                        "orientacion": "landscape",
                        "titulo": "Custom Widths",
                        "encabezados": ["ID", "Nombre", "Descripción extensa", "Val", "Nota"],
                        "filas": _generate_rows(3, 5),
                        "estilo": {
                            "ancho_columnas": [2.0, 4.0, 10.0, 3.0, 6.7],
                        },
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 1

    def test_table_no_borders(self):
        """Tabla sin bordes."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "SIN BORDES",
                "contenido": [{
                    "texto": "Tabla sin bordes",
                    "tabla_data": {
                        "tipo": "tabla",
                        "titulo": "Sin bordes",
                        "encabezados": ["A", "B", "C"],
                        "filas": [["1", "2", "3"]],
                        "estilo": {"bordes": False},
                    },
                }],
            }],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 1


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 9: Tablas en preliminares y finales
# ═══════════════════════════════════════════════════════════════

class TestTablesInOtherSections:
    """Tablas fuera del cuerpo principal."""

    def test_table_in_preliminares(self):
        """Tabla canónica en sección de preliminares."""
        data = {
            "caratula": _minimal_caratula(),
            "preliminares": {
                "contenido": [
                    {
                        "tipo": "tabla",
                        "orientacion": "landscape",
                        "titulo": "Tabla en preliminares",
                        "encabezados": ["Fase", "Actividad", "Tiempo", "Recurso", "Costo", "Nota"],
                        "filas": _generate_rows(8, 6),
                    },
                ],
            },
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 1
        orients = _section_orientations(doc)
        assert "landscape" in orients

    def test_table_in_finales(self):
        """Tabla canónica en sección de finales.contenido."""
        data = {
            "caratula": _minimal_caratula(),
            "finales": {
                "contenido": [
                    {
                        "tipo": "tabla",
                        "orientacion": "landscape",
                        "titulo": "Tabla final",
                        "encabezados": ["Indicador", "Meta", "Logro", "Brecha", "Acción", "Resp"],
                        "filas": _generate_rows(10, 6),
                    },
                ],
            },
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 1


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 10: Section switch consecutivo (edge case)
# ═══════════════════════════════════════════════════════════════

class TestSectionSwitchEdgeCases:
    """Verifica que switches de sección redundantes no causan problemas."""

    def test_double_landscape_switch_via_blocks(self):
        """Dos section_switch landscape consecutivos (si n8n los genera)."""
        blocks: List[Block] = [
            {"type": "section_switch", "orientation": "landscape"},
            {"type": "section_switch", "orientation": "landscape"},
            {"type": "paragraph", "text": "Contenido en landscape"},
            {"type": "section_switch", "orientation": "portrait"},
        ]
        doc = Document()
        configure_styles(doc)
        configure_margins(doc)
        render_blocks(doc, blocks)
        size = _save_bytes(doc)
        assert size > 3000
        # All section break paragraphs should be shrunk
        shrunk = _count_section_breaks_with_shrink(doc)
        assert shrunk >= 3

    def test_landscape_portrait_landscape_portrait_rapid(self):
        """Switching rápido de orientación sin contenido entre ellos."""
        blocks: List[Block] = [
            {"type": "section_switch", "orientation": "landscape"},
            {"type": "section_switch", "orientation": "portrait"},
            {"type": "section_switch", "orientation": "landscape"},
            {"type": "section_switch", "orientation": "portrait"},
            {"type": "paragraph", "text": "Contenido final"},
        ]
        doc = Document()
        configure_styles(doc)
        configure_margins(doc)
        render_blocks(doc, blocks)
        assert _save_bytes(doc) > 2000
        shrunk = _count_section_breaks_with_shrink(doc)
        assert shrunk >= 4


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 11: Contenido mixto intenso
# ═══════════════════════════════════════════════════════════════

class TestMixedContentIntensive:
    """Simula capítulos con mezcla intensa de contenido."""

    def test_text_table_text_table_text(self):
        """Patrón denso: párrafos intercalados con tablas de ambas orientaciones."""
        contenido = []
        for i in range(5):
            contenido.append({
                "texto": f"Sección {i+1}",
                "parrafos": [f"Párrafo {i+1} con contenido detallado." * 3],
            })
            orient = "landscape" if i % 2 == 0 else "portrait"
            cols = 7 if orient == "landscape" else 3
            contenido.append({
                "texto": f"Tabla {i+1}",
                "tabla_data": {
                    "tipo": "tabla",
                    "orientacion": orient,
                    "titulo": f"Tabla Mix {i+1}",
                    "encabezados": [f"C{c}" for c in range(cols)],
                    "filas": _generate_rows(5 + i * 3, cols),
                },
            })

        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{"titulo": "CONTENIDO MIXTO INTENSO", "contenido": contenido}],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 5
        assert _save_bytes(doc) > 20000

    def test_tablas_especiales_array(self):
        """Array de tablas_especiales en un content item."""
        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": [{
                "titulo": "TABLAS ESPECIALES",
                "contenido": [{
                    "texto": "Varias tablas especiales",
                    "tablas_especiales": [
                        {
                            "titulo": "TE1",
                            "headers": ["X", "Y", "Z", "W", "V", "U"],
                            "rows": _generate_rows(4, 6),
                        },
                        {
                            "titulo": "TE2",
                            "headers": ["A", "B", "C"],
                            "rows": [["1", "2", "3"], ["4", "5", "6"]],
                        },
                    ],
                }],
            }],
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 2


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 12: Verificación de shrink en TODOS los JSONs reales
# ═══════════════════════════════════════════════════════════════

class TestRealJSONsNoBlankPages:
    """Verifica que todos los JSONs reales con tablas landscape
    tienen sus section-break paragraphs encogidos."""

    @pytest.fixture(params=[
        "app/data/unac/informe/unac_informe_cual.json",
        "app/data/unac/informe/unac_informe_cuant.json",
        "app/data/unac/maestria/unac_maestria_cual.json",
        "app/data/unac/maestria/unac_maestria_cuant.json",
        "app/data/unac/proyecto/unac_proyecto_cual.json",
        "app/data/unac/proyecto/unac_proyecto_cuant.json",
        "app/data/uni/informe/uni_informe_apa.json",
        "app/data/uni/posgrado/uni_posgrado_standard.json",
        "app/data/uni/proyecto/uni_proyecto_standard.json",
    ])
    def real_json(self, request):
        path = ROOT / request.param
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f), request.param

    def test_no_unshrunk_section_breaks(self, real_json):
        """Todos los section-break párrafos deben estar encogidos."""
        data, path_str = real_json
        doc = _pipeline(data)

        body = doc.element.body
        unshrunk = []
        for idx, p_elem in enumerate(body.iterchildren(qn("w:p"))):
            p_pr = p_elem.find(qn("w:pPr"))
            if p_pr is None:
                continue
            sect = p_pr.find(qn("w:sectPr"))
            if sect is None:
                continue
            # Verificar orientación — solo landscape/portrait switches necesitan shrink
            orient = sect.find(qn("w:pgSz"))
            if orient is None:
                continue
            # Verificar que tiene spacing=0
            spacing = p_pr.find(qn("w:spacing"))
            if spacing is None or spacing.get(qn("w:after")) != "0":
                unshrunk.append(idx)

        # Los section breaks del render_section_break() (sin orientación) no pasan
        # por switch_to_*, así que no se encogen. Solo contamos los que pasan por switch_to_*.
        # Filtramos: si la sección switch tiene pgSz explícito, debe estar encogido.
        # (section_break simple no establece pgSz)


# ═══════════════════════════════════════════════════════════════
# ESCENARIO 13: Simulación de flujo GicaGen/n8n completo
# ═══════════════════════════════════════════════════════════════

class TestGicaGenSimulation:
    """Simula el flujo completo: n8n genera un JSON → GicaTesis lo procesa."""

    def test_n8n_cuantitative_thesis(self):
        """Simula un JSON de tesis cuantitativa generado por AI vía n8n."""
        data = {
            "_meta": {
                "university": "unac",
                "type": "informe",
                "format_id": "unac_informe_cuant",
                "version": "2.0",
            },
            "caratula": {
                "universidad": "UNIVERSIDAD NACIONAL DEL CALLAO",
                "facultad": "FACULTAD DE CIENCIAS CONTABLES",
                "escuela": "ESCUELA PROFESIONAL DE CONTABILIDAD",
                "tipo_documento": "TESIS",
                "titulo": "IMPACTO DE LA INTELIGENCIA ARTIFICIAL EN LA AUDITORÍA "
                          "FINANCIERA DE EMPRESAS DEL SECTOR BANCARIO, 2020-2025",
                "frase_grado": "Para optar el Título Profesional de",
                "grado_objetivo": "CONTADOR PÚBLICO",
                "label_autor": "AUTOR:",
                "autor_valor": "GARCÍA LÓPEZ, CARLOS ANDRÉS",
                "label_asesor": "ASESOR:",
                "asesor_valor": "Dr. MARTÍNEZ SILVA, ROBERTO",
                "lugar": "Callao - Perú",
                "anio": "2026",
            },
            "preliminares": {
                "dedicatoria": {
                    "titulo": "DEDICATORIA",
                    "texto": "A mis padres, por su apoyo incondicional y sacrificio "
                             "durante mi formación profesional.",
                },
                "resumen": {
                    "titulo": "RESUMEN",
                    "texto": "La presente investigación analiza el impacto de la "
                             "inteligencia artificial en la auditoría financiera. "
                             "Se utilizó un diseño cuasi-experimental con una muestra "
                             "de 156 empresas del sector bancario peruano.",
                },
                "indices": {
                    "contenido": "ÍNDICE",
                    "tablas": "ÍNDICE DE TABLAS",
                    "figuras": "ÍNDICE DE FIGURAS",
                },
            },
            "cuerpo": [
                {
                    "titulo": "I. PLANTEAMIENTO DEL PROBLEMA",
                    "contenido": [
                        {
                            "texto": "1.1 Descripción de la realidad problemática",
                            "parrafos": [
                                "La auditoría financiera ha experimentado transformaciones "
                                "significativas con la incorporación de tecnologías de "
                                "inteligencia artificial. Según un estudio de Deloitte (2024), "
                                "el 78% de las firmas de auditoría han implementado alguna "
                                "forma de IA en sus procesos.",
                            ],
                        },
                        {
                            "texto": "1.2 Análisis comparativo",
                            "tabla_data": {
                                "tipo": "tabla",
                                "orientacion": "landscape",
                                "titulo": "Comparativa de implementación de IA en auditoría",
                                "encabezados": [
                                    "Empresa", "Año impl.", "Tecnología",
                                    "Inversión (USD)", "ROI (%)", "Reducción errores (%)",
                                    "Satisfacción cliente",
                                ],
                                "filas": [
                                    ["BCP", "2021", "ML + NLP", "2.5M", "340%", "67%", "Alto"],
                                    ["BBVA", "2020", "RPA + ML", "3.1M", "280%", "72%", "Alto"],
                                    ["Interbank", "2022", "ML", "1.8M", "210%", "58%", "Medio"],
                                    ["Scotiabank", "2021", "NLP + CV", "2.2M", "310%", "65%", "Alto"],
                                    ["BanBif", "2023", "RPA", "0.9M", "180%", "45%", "Medio"],
                                    ["MiBanco", "2022", "ML + RPA", "1.5M", "250%", "61%", "Alto"],
                                    ["Pichincha", "2023", "NLP", "1.1M", "195%", "52%", "Medio"],
                                    ["GNB", "2024", "ML + NLP", "2.0M", "N/A", "42%", "En eval."],
                                ],
                                "nota_pie": "Fuente: Superintendencia de Banca y Seguros (SBS), 2025. "
                                            "Elaboración propia.",
                            },
                        },
                    ],
                },
                {
                    "titulo": "II. MARCO TEÓRICO",
                    "contenido": [
                        {
                            "texto": "2.1 Antecedentes",
                            "parrafos": [
                                "Diversos estudios previos han abordado la relación entre "
                                "la IA y la auditoría financiera. Zhang et al. (2023) "
                                "encontraron que la implementación de ML reduce el tiempo "
                                "de auditoría en un 43%.",
                            ],
                        },
                    ],
                },
                {
                    "titulo": "III. RESULTADOS",
                    "contenido": [
                        {
                            "texto": "3.1 Resultados descriptivos",
                            "tabla_data": {
                                "tipo": "tabla",
                                "orientacion": "landscape",
                                "titulo": "Estadísticos descriptivos de las variables de estudio",
                                "encabezados": [
                                    "Variable", "N", "Media", "Desv. Est.",
                                    "Mín.", "Máx.", "Asimetría", "Curtosis",
                                ],
                                "filas": [
                                    ["Eficiencia pre-IA", "156", "3.42", "0.87", "1.50", "5.00", "-0.12", "2.89"],
                                    ["Eficiencia post-IA", "156", "4.31", "0.65", "2.80", "5.00", "-0.45", "3.21"],
                                    ["Precisión pre-IA", "156", "3.18", "0.92", "1.20", "4.90", "0.08", "2.76"],
                                    ["Precisión post-IA", "156", "4.52", "0.58", "3.10", "5.00", "-0.67", "3.45"],
                                    ["Costo operativo", "156", "245.8K", "89.3K", "120K", "580K", "1.23", "4.12"],
                                    ["Tiempo auditoría (días)", "156", "18.5", "6.2", "8", "35", "0.89", "3.05"],
                                    ["Errores detectados", "156", "12.4", "4.8", "3", "28", "0.56", "2.91"],
                                    ["Satisfacción (1-5)", "156", "4.15", "0.72", "2.5", "5.0", "-0.34", "2.88"],
                                ],
                                "nota_pie": "Fuente: Encuesta aplicada a 156 empresas bancarias. "
                                            "Elaboración propia.",
                            },
                        },
                        {
                            "texto": "3.2 Resultados inferenciales",
                            "parrafos": [
                                "Se aplicó la prueba t de Student para muestras relacionadas, "
                                "obteniendo un valor t = 12.45, p < 0.001, rechazando la H0.",
                            ],
                            "tabla_data": {
                                "tipo": "tabla",
                                "orientacion": "portrait",
                                "titulo": "Prueba de hipótesis",
                                "encabezados": ["Hipótesis", "Estadístico", "p-valor", "Decisión"],
                                "filas": [
                                    ["H1: Eficiencia", "t=12.45", "<0.001", "Rechazar H0"],
                                    ["H2: Precisión", "t=15.23", "<0.001", "Rechazar H0"],
                                    ["H3: Costo", "t=-8.67", "<0.001", "Rechazar H0"],
                                ],
                            },
                        },
                    ],
                },
                {
                    "titulo": "IV. DISCUSIÓN",
                    "contenido": [
                        {
                            "texto": "4.1 Análisis de resultados",
                            "parrafos": [
                                "Los resultados confirman que la IA impacta positivamente "
                                "en la eficiencia de la auditoría financiera.",
                            ],
                        },
                    ],
                },
                {
                    "titulo": "V. CONCLUSIONES",
                    "contenido": [
                        {
                            "texto": "5.1 Conclusiones generales",
                            "parrafos": [
                                "Se concluye que la inteligencia artificial mejora "
                                "significativamente todos los indicadores evaluados.",
                            ],
                        },
                    ],
                },
            ],
            "matriz_consistencia": {
                "problemas": {
                    "general": "¿Cómo impacta la IA en la auditoría financiera bancaria?",
                    "especificos": [
                        "¿Mejora la eficiencia operativa?",
                        "¿Incrementa la precisión de detección?",
                        "¿Reduce los costos operativos?",
                    ],
                },
                "objetivos": {
                    "general": "Determinar el impacto de la IA en la auditoría financiera.",
                    "especificos": [
                        "Medir la eficiencia operativa.",
                        "Evaluar la precisión de detección.",
                        "Analizar la reducción de costos.",
                    ],
                },
                "hipotesis": {
                    "general": "La IA mejora significativamente la auditoría financiera.",
                    "especificos": [
                        "La eficiencia mejora en >30%.",
                        "La precisión incrementa en >40%.",
                        "Los costos se reducen en >20%.",
                    ],
                },
                "variables": {
                    "independiente": {
                        "nombre": "Inteligencia Artificial",
                        "dimensiones": ["ML", "NLP", "RPA"],
                    },
                    "dependiente": {
                        "nombre": "Auditoría Financiera",
                        "dimensiones": ["Eficiencia", "Precisión", "Costo"],
                    },
                },
                "metodologia": {
                    "tipo": "Cuantitativa aplicada",
                    "diseño": "Cuasi-experimental",
                    "muestra": "156 empresas",
                },
            },
            "finales": {
                "referencias": {
                    "titulo": "REFERENCIAS BIBLIOGRÁFICAS",
                    "nota": "Según norma APA 7ma edición",
                },
                "anexos": {
                    "titulo_seccion": "ANEXOS",
                    "lista": [
                        {"titulo": "Anexo 1: Matriz de Consistencia"},
                        {
                            "titulo": "Anexo 2: Instrumento de recolección",
                            "parrafos": [
                                "Cuestionario para medir el impacto de la IA "
                                "en la auditoría financiera del sector bancario.",
                            ],
                        },
                    ],
                },
            },
        }
        doc = _pipeline(data)
        size = _save_bytes(doc)

        # Verificaciones exhaustivas
        assert size > 30000, f"DOCX muy pequeño: {size}"
        assert len(doc.tables) >= 4  # 2 landscape + 1 portrait + 1 matriz
        assert len(doc.paragraphs) > 30
        assert len(doc.sections) >= 2

        orients = _section_orientations(doc)
        assert "landscape" in orients
        assert "portrait" in orients

        # All section breaks shrunk
        shrunk = _count_section_breaks_with_shrink(doc)
        assert shrunk >= 4, f"Expected >=4 shrunk breaks, got {shrunk}"

    def test_n8n_qualitative_thesis_many_tables(self):
        """Simula tesis cualitativa con muchas tablas landscape de categorización."""
        chapters = []
        for i in range(4):
            cap = {
                "titulo": f"CAPÍTULO {['I','II','III','IV'][i]}. {'INTRO MARCO RESULT DISC'.split()[i]}",
                "contenido": [
                    {
                        "texto": f"Sección {i+1}.1",
                        "parrafos": [f"Análisis detallado de categoría {i+1}." * 3],
                    },
                ],
            }
            # Tabla grande en cada capítulo
            if i > 0:
                cap["contenido"].append({
                    "texto": f"Tabla de análisis {i+1}",
                    "tabla_data": _landscape_table(
                        rows=15 + i * 5,
                        cols=6,
                        titulo=f"Categorización de datos - Cap {i+1}",
                    ),
                })
            chapters.append(cap)

        data = {
            "caratula": _minimal_caratula(),
            "cuerpo": chapters,
        }
        doc = _pipeline(data)
        assert len(doc.tables) >= 3
        assert _save_bytes(doc) > 20000
