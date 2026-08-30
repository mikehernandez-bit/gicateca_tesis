from __future__ import annotations

import pytest
from docx import Document

import app.engine.renderers  # noqa: F401
from app.engine.document_validation import (
    DocumentIndexValidationError,
    validate_docx_field_safety,
    validate_unac_project_document,
)
from app.engine.primitives import configure_styles, disable_update_fields, enable_update_fields
from app.engine.registry import render_blocks
from app.engine.normalizer import _attach_cached_index_entries


def _five_table_blocks() -> list[dict]:
    blocks: list[dict] = [
        {
            "type": "toc_field",
            "field_code": ' TOC \\c "Tabla" \\h \\z ',
            "heading_text": "ÍNDICE DE TABLAS",
            "page_label": "Pág.",
        },
        {
            "type": "toc_field",
            "field_code": ' TOC \\c "Figura" \\h \\z ',
            "heading_text": "ÍNDICE DE FIGURAS",
            "page_label": "Pág.",
        },
        {
            "type": "abbreviations_table",
            "rows": [{"sigla": "RCM", "meaning": "Mantenimiento Centrado en Confiabilidad"}],
        },
    ]
    for caption in (
        "Tabla 1.1 Diagnóstico de fallas",
        "Tabla 3.1 Operacionalización de variable independiente",
        "Tabla 3.2 Operacionalización de variable dependiente",
        "Tabla 5.1 Cronograma de actividades",
        "Tabla 6.1 Presupuesto de investigación",
    ):
        blocks.extend(
            [
                {"type": "heading", "text": f"{caption.split()[1].split('.')[0]}. CAPÍTULO", "level": 1},
                {
                    "type": "table",
                    "titulo": caption,
                    "encabezados": ["RCM", "Valor"],
                    "filas": [["RCM", "1"]],
                    "estilo": {"titulo_exacto": True},
                },
            ]
        )
    _attach_cached_index_entries(blocks)
    return blocks


def test_five_tables_produce_five_seq_captions_and_pass_quality_gate() -> None:
    doc = Document()
    configure_styles(doc)
    blocks = _five_table_blocks()
    render_blocks(doc, blocks)
    enable_update_fields(doc)

    report = validate_unac_project_document(doc, blocks)

    assert report["table_captions"] == 5
    assert report["page_headers"] == 2


def test_quality_gate_rejects_exact_table_without_seq_caption() -> None:
    doc = Document()
    configure_styles(doc)
    blocks = _five_table_blocks()
    render_blocks(doc, blocks)

    caption = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("Tabla 6.1"))
    for instr in list(caption._p.xpath(".//w:instrText")):
        instr.getparent().remove(instr)

    with pytest.raises(DocumentIndexValidationError, match="SEQ Tabla"):
        validate_unac_project_document(doc, blocks)


def test_generated_diagrams_count_as_expected_figure_captions() -> None:
    doc = Document()
    configure_styles(doc)
    blocks: list[dict] = [
        {
            "type": "toc_field",
            "field_code": ' TOC \\c "Figura" \\h \\z ',
            "heading_text": "ÍNDICE DE FIGURAS",
            "page_label": "Pág.",
        }
    ]
    blocks.extend(
        {
            "type": "image",
            "titulo": title,
            "ruta": "",
            "diagram_type": "flow",
            "diagram_data": {"labels": ["Entrada", "Análisis", "Resultado"]},
            "fuente": "Elaboración propia",
            "omit_caption": False,
        }
        for title in (
            "Proceso del RCM",
            "Taxonomía ISO 14224",
            "Flujo del AMEF",
            "Sistemas del equipo",
        )
    )

    _attach_cached_index_entries(blocks)

    render_blocks(doc, blocks)
    report = validate_unac_project_document(doc, blocks)

    assert report["figure_captions"] == 4
    assert report["visible_figure_index_entries"] == 4
    assert report["page_headers"] == 1


def test_quality_gate_rejects_empty_visible_figure_index() -> None:
    doc = Document()
    configure_styles(doc)
    blocks: list[dict] = [
        {
            "type": "toc_field",
            "field_code": ' TOC \\c "Figura" \\h \\z ',
            "heading_text": "ÍNDICE DE FIGURAS",
            "page_label": "Pág.",
        },
        {"type": "heading", "text": "II. MARCO TEÓRICO", "level": 1},
        {
            "type": "image",
            "titulo": "Proceso del RCM",
            "ruta": "",
            "diagram_type": "flow",
            "diagram_data": {"labels": ["Entrada", "Proceso", "Salida"]},
        },
    ]

    render_blocks(doc, blocks)

    with pytest.raises(DocumentIndexValidationError, match="índice visible de figuras incompleto"):
        validate_unac_project_document(doc, blocks)


def test_final_docx_preserves_fields_without_update_on_open(tmp_path) -> None:
    doc = Document()
    configure_styles(doc)
    blocks = _five_table_blocks()
    render_blocks(doc, blocks)
    enable_update_fields(doc)
    disable_update_fields(doc)
    output = tmp_path / "safe-fields.docx"
    doc.save(output)

    report = validate_docx_field_safety(output)

    assert report["native_fields"] >= 7
    assert report["update_fields_enabled"] == 0
    assert report["external_relationships"] == 0
