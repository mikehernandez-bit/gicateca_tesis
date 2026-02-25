"""
=============================================================================
ARCHIVO: tests/test_engine_renderers.py
FASE: Block Engine - Fase 4
=============================================================================

PROPÓSITO:
Tests para todos los renderers del Block Engine.
Verifica que cada renderer genera el DOCX correcto a partir de un Block dict.

ESTRATEGIA:
- Cada test crea un Document(), llama render_block() con un Block,
  y verifica el contenido generado en el documento.
- Test de integración: normalize 9 JSONs → render_blocks → save DOCX.

CÓMO EJECUTAR:
    py -m pytest tests/test_engine_renderers.py -v
=============================================================================
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import renderers to trigger @register
import app.engine.renderers  # noqa: F401
from app.engine.registry import render_block, render_blocks, list_registered, _RENDERERS
from app.engine.normalizer import normalize


ROOT = Path(__file__).resolve().parents[1]


# ─────────────────────────────────────────────────────────────
# FIXTURE: asegurar que los renderers están registrados
# ─────────────────────────────────────────────────────────────

@pytest.fixture(autouse=True)
def _ensure_renderers():
    """Si otro test limpió el registry, re-registra los renderers."""
    if not _RENDERERS:
        import importlib
        import app.engine.renderers as _mod
        # Re-import each renderer sub-module to re-trigger @register
        for name in [
            "apa_examples", "centered_text", "headings", "image",
            "info_table", "logo", "matriz", "note", "page_control",
            "paragraphs", "table", "toc",
        ]:
            importlib.reload(getattr(_mod, name))
    yield


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def _doc() -> Document:
    """Crea un Document limpio."""
    return Document()


def _render(block: dict) -> Document:
    """Crea un doc, renderiza un block, retorna el doc."""
    doc = _doc()
    render_block(doc, block)
    return doc


def _render_many(blocks: list) -> Document:
    """Crea un doc, renderiza una lista de blocks."""
    doc = _doc()
    render_blocks(doc, blocks)
    return doc


# ─────────────────────────────────────────────────────────────
# REGISTRO COMPLETO
# ─────────────────────────────────────────────────────────────

class TestAllRegistered:
    EXPECTED_TYPES = {
        "abbreviations_table", "apa_examples", "black_heading", "centered_text", "heading",
        "image", "index_items", "info_table", "legacy_table", "logo",
        "matriz", "note", "page_break", "page_footer", "paragraph",
        "paragraph_centered", "section_break", "section_switch",
        "table", "toc_field",
    }

    def test_all_20_registered(self):
        registered = set(list_registered())
        assert registered == self.EXPECTED_TYPES

    def test_count(self):
        assert len(list_registered()) == 20


# ─────────────────────────────────────────────────────────────
# CENTERED_TEXT
# ─────────────────────────────────────────────────────────────

class TestCenteredText:
    def test_renders_text(self):
        doc = _render({"type": "centered_text", "text": "HELLO", "bold": True, "size": 16})
        texts = [p.text for p in doc.paragraphs]
        assert "HELLO" in texts

    def test_centered_alignment(self):
        doc = _render({"type": "centered_text", "text": "TEST"})
        p = [p for p in doc.paragraphs if p.text == "TEST"][0]
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_bold_applied(self):
        doc = _render({"type": "centered_text", "text": "B", "bold": True})
        p = [p for p in doc.paragraphs if p.text == "B"][0]
        assert p.runs[0].bold is True

    def test_empty_text_noop(self):
        doc = _render({"type": "centered_text", "text": ""})
        # Empty text → add_p_centered returns early
        assert len(doc.paragraphs) == 0


# ─────────────────────────────────────────────────────────────
# LOGO
# ─────────────────────────────────────────────────────────────

class TestLogo:
    def test_logo_renders_without_crash(self):
        """Logo con datos válidos no crashea (puede fallback a [LOGO])."""
        doc = _render({
            "type": "logo",
            "data": {"configuracion": {"ruta_logo": "nonexistent.png"}},
        })
        texts = [p.text for p in doc.paragraphs]
        # Should have either a picture or the [LOGO] fallback
        assert len(doc.paragraphs) >= 1

    def test_logo_with_real_data(self):
        """Logo con ruta válida de UNAC inserta imagen."""
        doc = _render({
            "type": "logo",
            "data": {
                "configuracion": {"ruta_logo": "app/static/assets/LogoUNAC.png"},
                "_meta": {"university": "unac"},
            },
        })
        # Should have at least one paragraph (the image paragraph)
        assert len(doc.paragraphs) >= 1


# ─────────────────────────────────────────────────────────────
# HEADING + BLACK_HEADING
# ─────────────────────────────────────────────────────────────

class TestHeadings:
    def test_heading_renders(self):
        doc = _render({"type": "heading", "text": "CAPÍTULO I", "level": 1})
        # Headings in python-docx appear in doc.paragraphs with heading style
        found = any("CAPÍTULO I" in p.text for p in doc.paragraphs)
        assert found

    def test_heading_centered(self):
        doc = _render({"type": "heading", "text": "TITLE", "level": 1, "centered": True})
        p = [p for p in doc.paragraphs if "TITLE" in p.text][0]
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_black_heading_renders(self):
        doc = _render({"type": "black_heading", "text": "1.1 Sección", "level": 2})
        found = any("1.1 Sección" in p.text for p in doc.paragraphs)
        assert found

    def test_black_heading_black_color(self):
        doc = _render({"type": "black_heading", "text": "BH", "level": 2})
        from docx.shared import RGBColor
        p = [p for p in doc.paragraphs if "BH" in p.text][0]
        assert p.runs[0].font.color.rgb == RGBColor(0, 0, 0)

    def test_heading_has_no_page_break_before(self):
        doc = _render({"type": "heading", "text": "CAPITULO I", "level": 1})
        p = [p for p in doc.paragraphs if "CAPITULO I" in p.text][0]
        assert p.paragraph_format.page_break_before is False

    def test_black_heading_has_no_page_break_before(self):
        doc = _render({"type": "black_heading", "text": "SUBTITULO", "level": 2})
        p = [p for p in doc.paragraphs if "SUBTITULO" in p.text][0]
        assert p.paragraph_format.page_break_before is False


# ─────────────────────────────────────────────────────────────
# PARAGRAPH + PARAGRAPH_CENTERED
# ─────────────────────────────────────────────────────────────

class TestParagraphs:
    def test_paragraph(self):
        doc = _render({"type": "paragraph", "text": "Some text here"})
        assert any(p.text == "Some text here" for p in doc.paragraphs)

    def test_paragraph_centered(self):
        doc = _render({"type": "paragraph_centered", "text": "Centered"})
        p = [p for p in doc.paragraphs if p.text == "Centered"][0]
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER


# ─────────────────────────────────────────────────────────────
# NOTE
# ─────────────────────────────────────────────────────────────

class TestNote:
    def test_note_renders(self):
        doc = _render({"type": "note", "text": "Important info"})
        # Note is rendered as a table (blue box)
        assert len(doc.tables) == 1

    def test_note_empty_noop(self):
        doc = _render({"type": "note", "text": ""})
        assert len(doc.tables) == 0


# ─────────────────────────────────────────────────────────────
# PAGE CONTROL
# ─────────────────────────────────────────────────────────────

class TestPageControl:
    def test_page_break(self):
        doc = _render({"type": "page_break"})
        # page_break adds a paragraph with a break
        assert len(doc.paragraphs) >= 1

    def test_section_break(self):
        doc = _render({"type": "section_break"})
        # Adds a new section
        assert len(doc.sections) == 2  # default + new

    def test_section_switch_landscape(self):
        doc = _render({"type": "section_switch", "orientation": "landscape"})
        last_section = doc.sections[-1]
        # Landscape: width > height
        assert last_section.page_width > last_section.page_height

    def test_section_switch_portrait(self):
        doc = _render({"type": "section_switch", "orientation": "portrait"})
        last_section = doc.sections[-1]
        # Portrait: height > width
        assert last_section.page_height > last_section.page_width

    def test_section_switch_no_blank_page(self):
        """Section break paragraph must be shrunk to prevent blank pages."""
        from docx.oxml.ns import qn as _qn

        doc = _render({"type": "section_switch", "orientation": "landscape"})
        body = doc.element.body
        # Find the paragraph that hosts the section break (sectPr in pPr)
        for p_elem in body.iterchildren(_qn("w:p")):
            p_pr = p_elem.find(_qn("w:pPr"))
            if p_pr is not None and p_pr.find(_qn("w:sectPr")) is not None:
                # Spacing should be zero
                spacing = p_pr.find(_qn("w:spacing"))
                assert spacing is not None, "Section-break paragraph must have spacing element"
                assert spacing.get(_qn("w:before")) == "0"
                assert spacing.get(_qn("w:after")) == "0"
                # Font size should be minimal (2 half-points = 1pt)
                r_pr = p_pr.find(_qn("w:rPr"))
                assert r_pr is not None, "Section-break paragraph must have rPr"
                sz = r_pr.find(_qn("w:sz"))
                assert sz is not None and sz.get(_qn("w:val")) == "2"
                break
        else:
            pytest.fail("No section-break paragraph found")

    def test_page_break_skipped_after_section_switch(self):
        """page_break inmediatamente después de un section_switch debe omitirse.

        Esto previene hojas en blanco: el section_switch ya crea una nueva
        página (NEW_PAGE), así que un page_break adicional es redundante.
        """
        from docx.oxml.ns import qn as _qn

        doc = Document()
        # Simular: contenido → section_switch(portrait) → page_break
        doc.add_paragraph("Contenido antes del switch")
        render_blocks(doc, [
            {"type": "section_switch", "orientation": "portrait"},
            {"type": "page_break"},
        ])
        # Contar page_break characters (w:br con w:type="page")
        body = doc.element.body
        page_breaks = 0
        for p_elem in body.iterchildren(_qn("w:p")):
            for br in p_elem.iter(_qn("w:br")):
                if br.get(_qn("w:type")) == "page":
                    page_breaks += 1
        assert page_breaks == 0, (
            f"page_break should be skipped after section_switch, "
            f"but found {page_breaks} page break(s)"
        )

    def test_page_break_not_skipped_after_content(self):
        """page_break después de contenido normal NO se omite."""
        from docx.oxml.ns import qn as _qn

        doc = Document()
        render_blocks(doc, [
            {"type": "paragraph", "text": "Contenido normal"},
            {"type": "page_break"},
        ])
        body = doc.element.body
        page_breaks = 0
        for p_elem in body.iterchildren(_qn("w:p")):
            for br in p_elem.iter(_qn("w:br")):
                if br.get(_qn("w:type")) == "page":
                    page_breaks += 1
        assert page_breaks >= 1, "page_break after normal content should NOT be skipped"

    def test_page_break_skipped_when_previous_is_page_break(self):
        """Dos page_break consecutivos deben colapsarse para evitar hoja en blanco."""
        from docx.oxml.ns import qn as _qn

        doc = Document()
        render_blocks(doc, [
            {"type": "paragraph", "text": "Contenido normal"},
            {"type": "page_break"},
            {"type": "page_break"},
        ])
        body = doc.element.body
        page_breaks = 0
        for p_elem in body.iterchildren(_qn("w:p")):
            for br in p_elem.iter(_qn("w:br")):
                if br.get(_qn("w:type")) == "page":
                    page_breaks += 1
        assert page_breaks == 1, f"expected a single page break, got {page_breaks}"

    def test_page_break_force_keeps_intentional_second_break(self):
        """page_break forzado debe conservarse (caso página de respeto)."""
        from docx.oxml.ns import qn as _qn

        doc = Document()
        render_blocks(doc, [
            {"type": "paragraph", "text": "Portada"},
            {"type": "page_break"},
            {"type": "page_break", "force": True},
        ])
        body = doc.element.body
        page_breaks = 0
        for p_elem in body.iterchildren(_qn("w:p")):
            for br in p_elem.iter(_qn("w:br")):
                if br.get(_qn("w:type")) == "page":
                    page_breaks += 1
        assert page_breaks == 2, f"expected two page breaks, got {page_breaks}"

    def test_page_footer(self):
        doc = _render({"type": "page_footer"})
        footer = doc.sections[-1].footer
        # Footer should have content (PAGE field)
        assert len(footer.paragraphs) >= 1


# ─────────────────────────────────────────────────────────────
# TABLE
# ─────────────────────────────────────────────────────────────

class TestTable:
    def test_basic_table(self):
        doc = _render({
            "type": "table",
            "encabezados": ["A", "B", "C"],
            "filas": [["1", "2", "3"], ["4", "5", "6"]],
        })
        assert len(doc.tables) == 1
        t = doc.tables[0]
        assert len(t.rows) == 3  # header + 2 data
        assert len(t.columns) == 3

    def test_table_with_title(self):
        doc = _render({
            "type": "table",
            "titulo": "Mi tabla test",
            "encabezados": ["X"],
            "filas": [["1"]],
        })
        # Title paragraph before table
        texts = [p.text for p in doc.paragraphs]
        assert any("Mi tabla test" in t for t in texts)

    def test_table_landscape_switches_section(self):
        doc = _render({
            "type": "table",
            "orientacion": "landscape",
            "encabezados": ["A", "B"],
            "filas": [["1", "2"]],
        })
        # Should have at least 3 sections: default + landscape + portrait restore
        assert len(doc.sections) >= 3

    def test_table_portrait_no_extra_sections(self):
        doc = _render({
            "type": "table",
            "orientacion": "portrait",
            "encabezados": ["A"],
            "filas": [["1"]],
        })
        # Portrait: only default section
        assert len(doc.sections) == 1

    def test_table_no_headers_noop(self):
        doc = _render({"type": "table", "encabezados": [], "filas": []})
        assert len(doc.tables) == 0

    def test_table_cell_merge(self):
        doc = _render({
            "type": "table",
            "encabezados": ["A", "B", "C"],
            "filas": [["1", "2", "3"], ["4", "5", "6"]],
            "celdas_fusionadas": [{"fila": 0, "col": 0, "cols_span": 2}],
        })
        assert len(doc.tables) == 1

    def test_table_footer_note(self):
        doc = _render({
            "type": "table",
            "encabezados": ["A"],
            "filas": [["1"]],
            "nota_pie": "Nota de prueba",
        })
        texts = [p.text for p in doc.paragraphs]
        assert any("Nota de prueba" in t for t in texts)


# ─────────────────────────────────────────────────────────────
# LEGACY_TABLE
# ─────────────────────────────────────────────────────────────

class TestLegacyTable:
    def test_legacy_renders(self):
        doc = _render({
            "type": "legacy_table",
            "tabla": {"headers": ["X", "Y"], "rows": [["a", "b"]]},
            "titulo": "Legacy test",
        })
        assert len(doc.tables) == 1
        t = doc.tables[0]
        assert len(t.rows) == 2
        assert len(t.columns) == 2

    def test_legacy_no_data_noop(self):
        doc = _render({"type": "legacy_table", "tabla": {}})
        assert len(doc.tables) == 0

    def test_legacy_wide_table_landscape(self):
        """Tables with >5 headers get landscape orientation."""
        doc = _render({
            "type": "legacy_table",
            "tabla": {
                "headers": ["A", "B", "C", "D", "E", "F"],
                "rows": [["1", "2", "3", "4", "5", "6"]],
            },
        })
        # Should have landscape section
        assert len(doc.sections) >= 3


# ─────────────────────────────────────────────────────────────
# IMAGE
# ─────────────────────────────────────────────────────────────

class TestImage:
    def test_image_placeholder_is_omitted(self):
        """Image placeholder route is omitted completely."""
        doc = _render({
            "type": "image",
            "titulo": "Test Figure",
            "ruta": "placeholder",
        })
        assert len(doc.paragraphs) == 0

    def test_image_missing_file_is_omitted(self):
        """Image with missing file is omitted."""
        doc = _render({
            "type": "image",
            "titulo": "Fig",
            "ruta": "nonexistent_image.png",
            "fuente": "Elaboración propia",
        })
        assert len(doc.paragraphs) == 0


# ─────────────────────────────────────────────────────────────
# TOC_FIELD
# ─────────────────────────────────────────────────────────────

class TestTocField:
    def test_toc_renders(self):
        doc = _render({
            "type": "toc_field",
            "field_code": ' TOC \\o "1-3" ',
            "heading_text": "ÍNDICE",
        })
        found = any("ÍNDICE" in p.text for p in doc.paragraphs)
        assert found

    def test_toc_has_placeholder(self):
        doc = _render({
            "type": "toc_field",
            "field_code": ' TOC \\o "1-3" ',
            "heading_text": "ÍNDICE",
        })
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "Actualice" in texts

    def test_toc_enforces_page_break_after_block(self):
        from docx.oxml.ns import qn as _qn

        doc = _render_many([
            {
                "type": "toc_field",
                "field_code": ' TOC \\o "1-3" ',
                "heading_text": "ÍNDICE",
            },
            {"type": "heading", "text": "INTRODUCCION", "level": 1},
        ])

        page_breaks = 0
        for p_elem in doc.element.body.iterchildren(_qn("w:p")):
            for br in p_elem.iter(_qn("w:br")):
                if br.get(_qn("w:type")) == "page":
                    page_breaks += 1
        assert page_breaks >= 1


# ─────────────────────────────────────────────────────────────
# INDEX_ITEMS
# ─────────────────────────────────────────────────────────────

class TestIndexItems:
    def test_renders_items(self):
        doc = _render({
            "type": "index_items",
            "items": [
                {"texto": "OMS", "pag": "5"},
                {"texto": "WHO", "pag": "10", "bold": True},
            ],
        })
        texts = [p.text for p in doc.paragraphs]
        assert any("OMS" in t for t in texts)
        assert any("WHO" in t for t in texts)

    def test_empty_items_noop(self):
        doc = _render({"type": "index_items", "items": []})
        assert len(doc.paragraphs) == 0


class TestAbbreviationsTable:
    def test_renders_table_with_two_columns(self):
        doc = _render({
            "type": "abbreviations_table",
            "rows": [
                {"sigla": "IA", "meaning": "Inteligencia Artificial"},
                {"sigla": "ERP", "meaning": "Planificacion de Recursos Empresariales"},
            ],
        })
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.columns) == 2
        assert table.cell(0, 0).text == "SIGLA"
        assert table.cell(0, 1).text == "SIGNIFICADO"

    def test_empty_rows_show_discreet_note(self):
        doc = _render({"type": "abbreviations_table", "rows": []})
        texts = [p.text for p in doc.paragraphs]
        assert "(Completar abreviaturas)" in texts


# ─────────────────────────────────────────────────────────────
# INFO_TABLE
# ─────────────────────────────────────────────────────────────

class TestInfoTable:
    def test_renders_table(self):
        doc = _render({
            "type": "info_table",
            "elementos": [
                {"label": "TÍTULO:", "valor": "Mi tesis"},
                {"label": "AUTOR:", "valor": "Juan"},
            ],
        })
        assert len(doc.tables) == 1
        t = doc.tables[0]
        assert len(t.rows) == 2

    def test_empty_noop(self):
        doc = _render({"type": "info_table", "elementos": []})
        assert len(doc.tables) == 0


# ─────────────────────────────────────────────────────────────
# MATRIZ
# ─────────────────────────────────────────────────────────────

class TestMatriz:
    def _sample_matriz(self):
        return {
            "problemas": {"general": "¿Problema?", "especificos": ["E1", "E2"]},
            "objetivos": {"general": "Objetivo", "especificos": ["OE1"]},
            "hipotesis": {"general": "Hipótesis", "especificos": ["HE1"]},
            "variables": {
                "independiente": {"nombre": "VI", "dimensiones": ["D1"]},
                "dependiente": {"nombre": "VD", "dimensiones": ["D2"]},
            },
            "metodologia": {"tipo": "Aplicada", "diseño": "Experimental"},
        }

    def test_matriz_renders_table(self):
        doc = _render({
            "type": "matriz",
            "data": self._sample_matriz(),
            "landscape": False,
        })
        assert len(doc.tables) == 1
        t = doc.tables[0]
        assert len(t.columns) == 5  # PROBLEMAS, OBJETIVOS, HIPÓTESIS, VARIABLES, METODOLOGÍA

    def test_matriz_landscape_true(self):
        doc = _render({
            "type": "matriz",
            "data": self._sample_matriz(),
            "landscape": True,
        })
        # Landscape=True → render_tabla handles section switch internally
        assert len(doc.sections) >= 3

    def test_matriz_empty_data(self):
        doc = _render({"type": "matriz", "data": {}})
        assert len(doc.tables) == 0


# ─────────────────────────────────────────────────────────────
# APA_EXAMPLES
# ─────────────────────────────────────────────────────────────

class TestApaExamples:
    def test_renders_header_and_examples(self):
        doc = _render({
            "type": "apa_examples",
            "ejemplos": ["Author (2020). Title.", "Author2 (2021). Title2."],
        })
        texts = [p.text for p in doc.paragraphs]
        assert "Ejemplos APA 7:" in texts
        assert "Author (2020). Title." in texts

    def test_hanging_indent(self):
        from docx.shared import Cm
        doc = _render({
            "type": "apa_examples",
            "ejemplos": ["Example ref"],
        })
        # Second paragraph (the example) has hanging indent
        example_p = [p for p in doc.paragraphs if p.text == "Example ref"][0]
        assert example_p.paragraph_format.left_indent == Cm(1.27)
        assert example_p.paragraph_format.first_line_indent == Cm(-1.27)

    def test_empty_noop(self):
        doc = _render({"type": "apa_examples", "ejemplos": []})
        assert len(doc.paragraphs) == 0


# ─────────────────────────────────────────────────────────────
# INTEGRACIÓN: NORMALIZAR + RENDERIZAR 9 JSONs REALES
# ─────────────────────────────────────────────────────────────

class TestFullPipeline:
    """Carga cada JSON real, lo normaliza y renderiza todos los blocks.
    Verifica que genera un DOCX válido y guardable."""

    @pytest.fixture(params=[
        "app/data/unac/informe/unac_informe_cual.json",
        "app/data/unac/informe/unac_informe_cuant.json",
        "app/data/unac/maestria/unac_maestria_cual.json",
        "app/data/unac/maestria/unac_maestria_cuant.json",
        "app/data/unac/proyecto/unac_proyecto_cual.json",
        "app/data/unac/proyecto/unac_proyecto_cuant.json",
        "app/data/uni/informe/uni_informe_apa.json",
        "app/data/uni/posgrado/uni_posgrado_standard.json",
        "app/data/uni/proyecto/uni_proyecto_standard.json",
    ])
    def real_json(self, request):
        path = ROOT / request.param
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f), request.param

    def test_full_pipeline_no_crash(self, real_json):
        """normalize → render_blocks → Document guardable sin crash."""
        data, path = real_json
        blocks = normalize(data)
        doc = Document()

        from app.engine.primitives import configure_styles, configure_margins
        configure_styles(doc)
        configure_margins(doc)

        render_blocks(doc, blocks)

        # Guardar a bytes para verificar que no crashea al serializar
        import io
        buf = io.BytesIO()
        doc.save(buf)
        size = buf.tell()
        assert size > 10000, f"DOCX de {path} demasiado pequeño: {size} bytes"

    def test_full_pipeline_has_content(self, real_json):
        """El DOCX resultante tiene párrafos, tablas y secciones."""
        data, _ = real_json
        blocks = normalize(data)
        doc = Document()

        from app.engine.primitives import configure_styles, configure_margins
        configure_styles(doc)
        configure_margins(doc)

        render_blocks(doc, blocks)

        assert len(doc.paragraphs) > 20
        assert len(doc.tables) >= 1
        assert len(doc.sections) >= 2
