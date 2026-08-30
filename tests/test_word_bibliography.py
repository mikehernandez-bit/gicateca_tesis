"""Regression tests for Word-native simulated sources and citation fields."""

from __future__ import annotations

from zipfile import ZipFile

from docx import Document

from app.engine.word_bibliography import (
    BIBLIOGRAPHY_NS,
    NativeBibliographyValidationError,
    add_bibliography_field,
    embed_word_sources,
    extract_word_sources,
    parse_simulated_source,
    render_text_with_citations,
    validate_native_bibliography_docx,
)


def _data() -> dict:
    return {
        "finales": {
            "referencias": {
                "_ai_content": (
                    "Fuentes simuladas para validación.\n\n"
                    "[[SOURCE:SIM_01_MORALES_2025]] Morales, J., & Quispe, L. (2025). "
                    "Confiabilidad de activos industriales. Revista Técnica Aplicada, 14(2), 45-62. "
                    "Referencia propuesta simulada para validacion del autor.\n\n"
                    "[[SOURCE:SIM_02_ROJAS_2024]] Rojas, M. (2024). Métodos de mantenimiento. "
                    "Editorial Academia Técnica. Referencia propuesta simulada para validacion del autor."
                )
            }
        }
    }


def test_native_sources_fields_and_source_manager_storage(tmp_path) -> None:
    sources = extract_word_sources(_data())
    assert [source["source_type"] for source in sources] == ["JournalArticle", "Book"]
    assert all("FUENTE SIMULADA" in source["comments"] for source in sources)

    document = Document()
    paragraph = document.add_paragraph()
    count = render_text_with_citations(
        paragraph,
        "Sustento del estudio [[CITE:SIM_01_MORALES_2025;SIM_02_ROJAS_2024]].",
        sources,
    )
    add_bibliography_field(document.add_paragraph(), sources)
    output = tmp_path / "native-word-sources.docx"
    document.save(output)
    embed_word_sources(output, sources)

    report = validate_native_bibliography_docx(output, sources)
    assert count == 1
    assert report == {
        "sources": 2,
        "citation_fields": 1,
        "bibliography_fields": 1,
        "bibliography_cached_entries": 2,
    }

    with ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        custom_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("customXml/item") and name.endswith(".xml")
        )
    assert "\\t" in document_xml
    assert BIBLIOGRAPHY_NS in custom_xml
    assert "SIM_01_MORALES_2025" in custom_xml
    assert "FUENTE SIMULADA GENERADA POR GICA" in custom_xml


def test_validation_rejects_uncited_embedded_source(tmp_path) -> None:
    sources = extract_word_sources(_data())
    document = Document()
    paragraph = document.add_paragraph()
    render_text_with_citations(paragraph, "Texto [[CITE:SIM_01_MORALES_2025]].", sources)
    add_bibliography_field(document.add_paragraph(), sources)
    output = tmp_path / "missing-citation.docx"
    document.save(output)
    embed_word_sources(output, sources)

    try:
        validate_native_bibliography_docx(output, sources)
    except NativeBibliographyValidationError as exc:
        assert "sin citar=SIM_02_ROJAS_2024" in str(exc)
    else:
        raise AssertionError("La validación debía rechazar una fuente sin campo CITATION")


def test_validation_rejects_duplicate_author_year_that_expands_word_citation(tmp_path) -> None:
    sources = extract_word_sources(_data())
    duplicate = dict(sources[0])
    duplicate["tag"] = "SIM_03_MORALES_2025"
    duplicate["guid"] = "{22222222-2222-2222-2222-222222222222}"
    duplicate["title"] = "Otro título para el mismo autor y año"
    sources.append(duplicate)

    document = Document()
    paragraph = document.add_paragraph()
    render_text_with_citations(
        paragraph,
        "Texto [[CITE:SIM_01_MORALES_2025;SIM_02_ROJAS_2024;SIM_03_MORALES_2025]].",
        sources,
    )
    add_bibliography_field(document.add_paragraph(), sources)
    output = tmp_path / "duplicate-author-year.docx"
    document.save(output)
    embed_word_sources(output, sources)

    try:
        validate_native_bibliography_docx(output, sources)
    except NativeBibliographyValidationError as exc:
        assert "autor y año duplicados" in str(exc)
        assert "podría mostrar el título" in str(exc)
    else:
        raise AssertionError("La validación debía rechazar autor y año duplicados")


def test_repeated_citation_fields_share_one_source_and_one_bibliography_entry(tmp_path) -> None:
    sources = extract_word_sources(
        {
            "finales": {
                "referencias": {
                    "_ai_content": (
                        "[[SOURCE:SIM_01_MORALES_2025]] Morales, J. (2025). "
                        "Confiabilidad industrial. Revista Técnica, 10(1), 10-20."
                    )
                }
            }
        }
    )
    document = Document()
    for _ in range(4):
        render_text_with_citations(
            document.add_paragraph(),
            "Sustento repetido [[CITE:SIM_01_MORALES_2025]].",
            sources,
        )
    add_bibliography_field(document.add_paragraph(), sources)
    output = tmp_path / "repeated-source.docx"
    document.save(output)
    embed_word_sources(output, sources)

    report = validate_native_bibliography_docx(output, sources)
    assert report["sources"] == 1
    assert report["citation_fields"] == 4
    assert report["bibliography_cached_entries"] == 1


def test_narrative_citation_keeps_author_visible_and_year_native(tmp_path) -> None:
    sources = extract_word_sources(_data())
    document = Document()
    paragraph = document.add_paragraph()
    render_text_with_citations(
        paragraph,
        "El planteamiento de [[CITE_NARRATIVE:SIM_01_MORALES_2025]] sustenta el análisis.",
        sources,
    )
    output = tmp_path / "narrative-citation.docx"
    document.save(output)
    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "Morales &amp; Quispe" in xml
    assert "CITATION SIM_01_MORALES_2025" in xml
    assert "\\n" in xml
    assert "(2025)" in xml


def test_legacy_mil_standard_author_renders_and_validates_without_false_missing_entry(
    tmp_path,
) -> None:
    source = parse_simulated_source(
        "SIM_24_A_1980",
        (
            "MIL-STD-1629A, M. (1980). Fundamentos y evidencia sobre bases teoricas. "
            "Fondo Editorial Tecnico. Referencia propuesta simulada para validacion del autor."
        ),
    )
    assert source is not None
    assert source["authors"] == [{"last": "MIL-STD-1629A", "first": ""}]
    assert source["original_text"].startswith("MIL-STD-1629A (1980).")
    assert "MIL-STD-1629A, M." not in source["original_text"]

    document = Document()
    render_text_with_citations(
        document.add_paragraph(),
        "El AMEF se estructura según la norma [[CITE:SIM_24_A_1980]].",
        [source],
    )
    add_bibliography_field(document.add_paragraph(), [source])
    output = tmp_path / "legacy-mil-standard.docx"
    document.save(output)
    embed_word_sources(output, [source])

    report = validate_native_bibliography_docx(output, [source])
    assert report["sources"] == 1
    assert report["bibliography_cached_entries"] == 1
