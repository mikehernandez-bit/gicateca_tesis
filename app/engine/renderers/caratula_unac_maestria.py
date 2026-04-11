"""Renderer: caratula_unac_maestria."""

from __future__ import annotations

import logging
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.engine.registry import register
from app.engine.primitives import add_p_centered, resolve_logo_path
from app.engine.types import Block

logger = logging.getLogger(__name__)

def _first_nonempty_text(candidates: list) -> str:
    for value in candidates:
        if value is None:
            continue
        text = str(value).strip()
        if text:
            return text
    return ""

def _norm_upper(text: str) -> str:
    import unicodedata
    normalized = unicodedata.normalize("NFKD", text or "")
    stripped = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return stripped.upper().strip()

def _looks_like_cover_title_placeholder(text: str) -> bool:
    value = _norm_upper(text)
    if not value or len(value) < 5:
        return True
    
    # Si tiene brackets pero es largo (> 20 chars), probablemente es un titulo real con algo extra
    if "[" in value and "]" in value and len(value) < 20:
        return True
    if "{{" in value and "}}" in value and len(value) < 20:
        return True
        
    markers = (
        "TITULO DEL PROYECTO",
        "TITULO COMPLETO DEL TRABAJO",
        "TITULO DE LA TESIS",
        "ESCRIBA AQUI",
    )
    # Solo es placeholder si el marcador es casi todo el texto
    for marker in markers:
        if marker in value and len(value) < (len(marker) + 10):
            return True
    return False

@register("caratula_unac_maestria")
def render_caratula_unac_maestria(doc: Document, block: Block) -> None:
    data = block.get("data", {})
    c = block.get("caratula", {})
    
    # Proactive search in multiple pockets (c, data, values)
    def _find_val(key: str, default: str) -> str:
        candidates = [
            c.get(key),
            data.get(key),
            (data.get("values") or {}).get(key)
        ]
        return next((str(v).strip() for v in candidates if v and str(v).strip()), default)

    # Universidad 18pt
    add_p_centered(doc, _find_val("universidad", "UNIVERSIDAD NACIONAL DEL CALLAO"), bold=True, size=18, space_after=6)
    
    # Facultad 16pt
    add_p_centered(doc, _find_val("facultad", "ESCUELA DE POSGRADO"), bold=True, size=16, space_after=6)
    
    # Escuela 16pt
    add_p_centered(doc, _find_val("escuela", "UNIDAD DE POSGRADO"), bold=True, size=16, space_after=24)

    # Logo render sizing exact to 1.5125 as requested
    logo_path = resolve_logo_path(data)
    if logo_path:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(logo_path), width=Inches(1.5125))
        except Exception:
            add_p_centered(doc, "[LOGO]")
    else:
        add_p_centered(doc, "[LOGO]")

    # Tipo Documento 14pt
    tp = _find_val("tipo_documento", "PROYECTO DE INVESTIGACIÓN")
    add_p_centered(doc, tp, bold=True, size=14, space_before=24, space_after=12)

    # Frase Grado 12pt Bold
    frase_grado = _find_val("frase_grado", "")
    if frase_grado:
        add_p_centered(doc, frase_grado, bold=True, size=12, space_after=24)

    # Titulo 14pt Bold UPPER
    titulo = _find_val("titulo", _find_val("titulo_investigacion", "[Título de la tesis]"))
    add_p_centered(doc, titulo.upper(), bold=True, size=14, space_before=24, space_after=30)

    # Autores Block (Centrado)
    label_autor = _find_val("label_autor", "AUTOR(ES):")
    add_p_centered(doc, label_autor, bold=True, size=12)
    
    autores = c.get("autores") or data.get("autores")
    if not autores:
        solo_autor = _find_val("autor_valor", _find_val("autor1_nombres", None))
        if solo_autor:
            autores = [solo_autor]
        else:
            autores = ["[Apellidos y nombres]"]
    
    if isinstance(autores, str):
        autores = [autores]
        
    for autor in autores:
        add_p_centered(doc, autor, bold=False, size=12)

    # Asesor Block (Centrado)
    label_asesor = _find_val("label_asesor", "ASESOR:")
    add_p_centered(doc, label_asesor, bold=True, size=12, space_before=18)
    
    asesor = c.get("asesor") or _find_val("asesor_valor", _find_val("asesor_nombres", "[Apellidos y nombres]"))
    add_p_centered(doc, asesor, bold=False, size=12, space_after=18)

    # Linea de investigacion (Centrado)
    label_linea = _find_val("label_linea", "LÍNEA DE INVESTIGACIÓN:")
    add_p_centered(doc, label_linea, bold=True, size=12)
    
    linea = _find_val("linea_investigacion", _find_val("linea", "[Nombre de la línea]"))
    add_p_centered(doc, linea, bold=False, size=12, space_after=40)

    # Footer
    lugar = _find_val("lugar", "Callao")
    anio = _find_val("anio", "2026")
    
    # Combinar lugar y año si no están ya juntos
    if anio not in lugar:
        lugar_completo = f"{lugar}, {anio}"
    else:
        lugar_completo = lugar
        
    add_p_centered(doc, lugar_completo, bold=True, size=12, space_before=60)
    add_p_centered(doc, "PERÚ", bold=True, size=12)

    doc.add_page_break()
