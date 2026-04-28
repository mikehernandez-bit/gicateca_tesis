"""Renderers: toc_field + index_items.

Agrupados porque ambos generan índices/listas de contenido del documento.
"""

from __future__ import annotations

from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.engine.registry import register
from app.engine.primitives import add_toc_field
from app.engine.types import Block


@register("toc_field")
def render_toc_field(doc: Document, block: Block) -> None:
    """Renderiza un campo TOC de Word (índice auto-actualizable).

    Inserta: heading + campo Word + placeholder + page_break.

    Block keys:
        field_code (str): Código del campo Word (ej: ' TOC \\o "1-3" ').
        heading_text (str): Título del índice.
    """
    add_toc_field(
        doc,
        block.get("field_code", ""),
        block.get("heading_text", ""),
        exclude_from_toc=block.get("exclude_from_toc", False),
    )


@register("index_items")
def render_index_items(doc: Document, block: Block) -> None:
    """Renderiza una lista de índice estática con tab stops y puntos.

    Usado para índices que NO son TOC de Word (abreviaturas, etc.).
    Replica la lógica de render_preliminares para indices tipo list con items.

    Block keys:
        items (list[dict]): Lista de {"texto": str, "pag": str|int, "bold": bool}.
    """
    for item in block.get("items", []):
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.0),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        run = p.add_run(item.get("texto", ""))
        if item.get("bold"):
            run.bold = True
        p.add_run(f"\t{item.get('pag', '')}")


def _add_abbreviation_line(doc: Document, sigla: str, meaning: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(4.0), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(4.5), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)

    run_sigla = paragraph.add_run(sigla)
    run_sigla.bold = True
    run_sigla.font.name = "Arial"
    run_sigla.font.size = Pt(10)

    run_colon = paragraph.add_run("\t:\t")
    run_colon.font.name = "Arial"
    run_colon.font.size = Pt(10)

    run_meaning = paragraph.add_run(meaning)
    run_meaning.font.name = "Arial"
    run_meaning.font.size = Pt(10)


@register("abbreviations_table")
def render_abbreviations_table(doc: Document, block: Block) -> None:
    """Renderiza abreviaturas como lista alineada, sin grid visible."""
    rows = block.get("rows", []) or []

    if not rows:
        note = doc.add_paragraph("No se identificaron abreviaturas relevantes en el documento.")
        note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = (
            note.runs[0]
            if note.runs
            else note.add_run("No se identificaron abreviaturas relevantes en el documento.")
        )
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        return

    for row in rows:
        sigla = str(row.get("sigla", "")).strip()
        meaning = str(row.get("meaning", "")).strip()
        if not sigla or not meaning:
            continue
        _add_abbreviation_line(doc, sigla, meaning)
