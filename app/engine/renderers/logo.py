"""Renderer: logo — logotipo universitario centrado."""
from __future__ import annotations

import logging

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.engine.registry import register
from app.engine.primitives import resolve_logo_path, add_p_centered
from app.engine.types import Block

logger = logging.getLogger(__name__)


@register("logo")
def render_logo(doc: Document, block: Block) -> None:
    """Renderiza el logo universitario centrado.

    Block keys:
        data (dict): JSON completo para resolver la ruta del logo.
        width_inches (float): Ancho en pulgadas. Default 2.0.
    """
    data = block.get("data", {})
    width = block.get("width_inches", 2.0)
    logo_path = resolve_logo_path(data)

    if logo_path:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.add_run().add_picture(str(logo_path), width=Inches(width))
        except Exception as e:
            logger.warning("Could not insert logo %s: %s", logo_path, e)
            add_p_centered(doc, "[LOGO]", space_before=12, space_after=12)
    else:
        add_p_centered(doc, "[LOGO]", space_before=12, space_after=12)
