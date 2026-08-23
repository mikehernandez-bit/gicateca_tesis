"""Renderer: image — figuras con caption SEQ-numbered y fuente."""

from __future__ import annotations

import logging
import re

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.engine.render_state import next_figure_number
from app.engine.registry import register
from app.engine.primitives import resolve_asset, add_seq_field
from app.engine.types import Block

logger = logging.getLogger(__name__)

# Regex defensivas para limpiar títulos de figuras
_RE_CHAPTER_SUFFIX = re.compile(r"\s*[–—-]\s*[IVXLC]+\.\s*.+$")
_RE_FIGURA_PREFIX = re.compile(r"^Figura\s*[\d.]+\s*[:.]*\s*", re.IGNORECASE)
_RE_PROJECT_SUFFIX = re.compile(
    r"\s+(?:aplicad[oa]s?|orientad[oa]s?)\s+(?:a|al)\s+.+$",
    re.IGNORECASE,
)
_CAPTION_MAX_CHARS = 120
_CAPTION_ACRONYMS = (
    "AMEF",
    "CAT",
    "CBM",
    "CMMS",
    "FMEA",
    "GMG",
    "ISO",
    "IoT",
    "MTBF",
    "MTTR",
    "RCM",
    "SAE",
)


def _clean_figure_title(titulo: str) -> str:
    """Normalize a brief academic caption and remove the thesis-title suffix."""
    titulo = re.sub(r"\s+", " ", str(titulo or "").strip())
    titulo = _RE_FIGURA_PREFIX.sub("", titulo)
    titulo = _RE_CHAPTER_SUFFIX.sub("", titulo)
    titulo = _RE_PROJECT_SUFFIX.sub("", titulo).strip(" .;:-")

    if titulo.isupper():
        titulo = titulo.lower()
        titulo = titulo[:1].upper() + titulo[1:]

    for acronym in _CAPTION_ACRONYMS:
        titulo = re.sub(
            rf"(?<![A-Za-z0-9]){re.escape(acronym)}(?![A-Za-z0-9])",
            acronym,
            titulo,
            flags=re.IGNORECASE,
        )

    if len(titulo) > _CAPTION_MAX_CHARS:
        shortened = titulo[: _CAPTION_MAX_CHARS + 1].rsplit(" ", 1)[0].rstrip(" ,;:-")
        titulo = shortened or titulo[:_CAPTION_MAX_CHARS].rstrip(" ,;:-")
    return titulo


def _rgb_from_hex(value: str) -> RGBColor | None:
    text = str(value or "").strip().lstrip("#")
    if len(text) != 6:
        return None
    try:
        red = int(text[0:2], 16)
        green = int(text[2:4], 16)
        blue = int(text[4:6], 16)
    except ValueError:
        return None
    return RGBColor(red, green, blue)


@register("image")
def render_image(doc: Document, block: Block) -> None:
    """Renderiza una imagen con caption numerada y fuente opcional.

    Replica _render_image() del generador:
    1. Caption: "Figura N.N {titulo}" (Arial 10, centrada)
    2. Imagen centrada (12 cm ancho)
    3. Nota o fuente debajo de la figura

    Block keys:
        titulo (str): Título de la figura.
        ruta (str): Ruta del archivo de imagen.
        fuente (str): Fuente de la imagen.
    """
    ruta = block.get("ruta", "")
    titulo = _clean_figure_title(block.get("titulo", ""))
    omit_caption = bool(block.get("omit_caption"))
    nota = str(block.get("nota") or block.get("note") or "").strip()
    nota_color = _rgb_from_hex(str(block.get("nota_color") or block.get("note_color") or "").strip())
    placeholder_text = str(
        block.get("placeholder_text") or block.get("texto_placeholder") or ""
    ).strip()

    # Omit placeholders or missing files: never inject fake "example" figures.
    if not ruta or ruta.lower() == "placeholder":
        return
    path = resolve_asset(ruta)
    if not path:
        return

    try:
        # Caption with SEQ field
        if titulo and not omit_caption:
            chapter_number, figure_number, is_first_in_chapter = next_figure_number()
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.keep_with_next = True
            pc.paragraph_format.space_after = Pt(4)
            prefix = (
                f"Figura {chapter_number}." if chapter_number is not None else "Figura "
            )
            rl = pc.add_run(prefix)
            rl.font.name = "Arial"
            rl.font.size = Pt(10)
            add_seq_field(
                pc,
                "Figura",
                reset_to=1 if is_first_in_chapter else None,
                display_value=figure_number,
                font_name="Arial",
                font_size_pt=10,
                bold=False,
            )
            rt = pc.add_run(f" {titulo}")
            rt.font.name = "Arial"
            rt.font.size = Pt(10)

        if placeholder_text:
            pi = doc.add_paragraph()
            pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pi.paragraph_format.space_before = Pt(12)
            pi.paragraph_format.space_after = Pt(12)
            rp = pi.add_run(placeholder_text)
            rp.italic = True
            rp.font.name = "Arial"
            rp.font.size = Pt(10)
        else:
            pi = doc.add_paragraph()
            pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pi.add_run().add_picture(path, width=Cm(12.0))

        if block.get("fuente"):
            ps = doc.add_paragraph(f"Fuente: {block['fuente']}")
            ps.runs[0].italic = True
            ps.runs[0].font.name = "Arial"
            ps.runs[0].font.size = Pt(10)
            ps.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if nota:
            pn = doc.add_paragraph()
            pn.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if nota_color is not None:
                rn = pn.add_run(nota)
                rn.font.color.rgb = nota_color
                rn.font.name = "Arial"
                rn.font.size = Pt(10)
            else:
                rn = pn.add_run("Nota.")
                rn.italic = True
                rn.font.name = "Arial"
                rn.font.size = Pt(10)
                rt = pn.add_run(f" {nota}")
                rt.italic = True
                rt.font.name = "Arial"
                rt.font.size = Pt(10)
    except Exception as e:
        logger.warning("Image %s: %s", ruta, e)
