"""Renderers: page_break + section_break + section_switch + page_footer.

Agrupados porque todos controlan la estructura de páginas/secciones del documento.
"""

from __future__ import annotations

from docx.document import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

from app.engine.registry import register
from app.engine.primitives import (
    add_fld_page,
    switch_to_landscape,
    switch_to_portrait,
)
from app.engine.types import Block


def _last_content_is_section_break(doc: Document) -> bool:
    """Verifica si el último elemento de contenido del body es un section-break paragraph.

    Cuando ``switch_to_portrait()`` (o ``switch_to_landscape()``) se ejecuta,
    python-docx crea un ``<w:p>`` cuyo ``<w:pPr>`` contiene un ``<w:sectPr>``.
    Ese section break ya inicia una nueva página (NEW_PAGE), por lo que un
    ``page_break`` inmediatamente después sería redundante y generaría una
    hoja en blanco.

    Retorna True si el último contenido del body (ignorando el ``<w:sectPr>``
    a nivel de body) es un section-break paragraph.
    """
    body = doc.element.body
    for child in reversed(list(body)):
        # Saltar el sectPr a nivel de body (propiedades de la última sección)
        if child.tag == qn("w:sectPr"):
            continue
        # ¿Es un párrafo con sectPr embebido? → section break
        if child.tag == qn("w:p"):
            p_pr = child.find(qn("w:pPr"))
            if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
                return True
        return False
    return False


def _last_empty_trailing_paragraph(doc: Document):
    """Return the body's last paragraph if it is a bare, run-less spacer.

    ``python-docx`` (and our own ``add_styled_note``/table helpers) always
    leave an empty ``<w:p>`` right after a table, because OOXML forbids a
    table from being the very last element of the body/section. When a
    ``page_break`` block follows a table, ``doc.add_page_break()`` creates a
    brand-new paragraph *in addition to* that spacer, so two near-empty
    paragraphs stack up right at the page boundary. On some renderers
    (observed via Word/LibreOffice DOCX->PDF conversion) that pairing
    produces one extra blank page. Reusing the existing empty spacer for the
    page-break run avoids ever having two of them back to back.
    """
    body = doc.element.body
    children = list(body)
    if not children:
        return None
    last = children[-1]
    tag = last.tag
    if tag == qn("w:sectPr"):
        if len(children) < 2:
            return None
        last = children[-2]
        tag = last.tag
    if tag != qn("w:p"):
        return None
    if last.find(qn("w:tbl")) is not None:
        return None
    if "".join(last.itertext()).strip():
        return None
    for paragraph in reversed(doc.paragraphs):
        if paragraph._p is last:
            return paragraph
    return None


def _last_content_has_page_break(doc: Document) -> bool:
    """Verifica si el último contenido del body contiene un salto de página explícito."""
    body = doc.element.body
    for child in reversed(list(body)):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p"):
            for br in child.findall(f".//{qn('w:br')}"):
                if br.get(qn("w:type")) == "page":
                    return True
        return False
    return False


def _remove_trailing_page_break(doc: Document) -> bool:
    """Remove trailing page break run (<w:br w:type='page'/>) from body to prevent double page breaks."""
    body = doc.element.body
    for child in reversed(list(body)):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p"):
            for run in list(child.findall(f".//{qn('w:r')}")):
                for br in list(run.findall(f".//{qn('w:br')}")):
                    if br.get(qn("w:type")) == "page":
                        run.remove(br)
                        return True
        return False


@register("page_break")
def render_page_break(doc: Document, block: Block) -> None:
    """Inserta un salto de página.

    Si el último contenido del documento es un section-break paragraph
    (ej: después de ``switch_to_portrait`` al final de una tabla landscape),
    se omite el page_break porque el section break ya inició una nueva página.
    Esto evita hojas en blanco entre tablas landscape y el contenido siguiente.
    Si force=True, el salto se aplica de forma garantizada (ej: hoja de respeto).
    """
    force = bool(block.get("force", False))
    if not force and _last_content_is_section_break(doc):
        return
    # Evita doble page break consecutivo (otra fuente típica de hoja en blanco).
    if not force and _last_content_has_page_break(doc):
        return

    # Limpiar space_after del último párrafo para evitar desbordes antes del salto.
    if doc.paragraphs:
        last_p = doc.paragraphs[-1]
        if last_p.paragraph_format.space_after:
            last_p.paragraph_format.space_after = 0

    # Si el body ya termina en un parrafo vacio (el spacer que OOXML exige
    # despues de una tabla, p.ej. tras una nota en caja azul), reutilizalo
    # para el salto de pagina en vez de crear otro parrafo nuevo justo al
    # lado -- ver docstring de _last_empty_trailing_paragraph.
    trailing = _last_empty_trailing_paragraph(doc)
    if trailing is not None and not trailing.text.strip():
        trailing.add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_page_break()


@register("section_break")
def render_section_break(doc: Document, block: Block) -> None:
    """Inserta un salto de sección (nueva página). Preserva la hoja de respeto inicial e investiga saltos institucionales."""
    force = bool(block.get("force", False))
    if not force and _last_content_is_section_break(doc):
        return
    doc.add_section(WD_SECTION.NEW_PAGE)


@register("section_switch")
def render_section_switch(doc: Document, block: Block) -> None:
    """Cambia orientación de página (landscape/portrait).

    Block keys:
        orientation (str): "landscape" o "portrait".
    """
    orientation = block.get("orientation", "portrait")
    if orientation == "landscape":
        switch_to_landscape(doc)
    else:
        switch_to_portrait(doc)


@register("page_footer")
def render_page_footer(doc: Document, block: Block) -> None:
    """Inserta numeración de páginas en el footer de la última sección.

    Replica la lógica final de generate_document_unified():
    footer alineado a la derecha con campo PAGE.
    """
    section = doc.sections[-1]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_fld_page(p)
