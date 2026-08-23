"""Renderers: table + legacy_table.

Contiene la implementación completa de render_tabla que replica exactamente
la función original del universal_generator.py, incluyendo:
- Orientación landscape/portrait con switch automático de sección
- Títulos con campo SEQ para auto-numeración
- Encabezados con sombreado
- Fusión de celdas
- Notas de pie de tabla

La función _render_tabla_impl es reutilizada por el renderer de matriz.
"""

from __future__ import annotations

import re

from docx.document import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.engine.render_state import next_table_number
from app.engine.registry import register
from app.engine.primitives import (
    DEFAULT_HEADER_COLOR,
    DEFAULT_TABLE_FONT_SIZE,
    LANDSCAPE_FONT_SIZE,
    PORTRAIT_MARGINS,
    LANDSCAPE_MARGINS,
    add_seq_field,
    apply_cell_shading,
    format_cell_text,
    set_cell_vertical_alignment,
    switch_to_landscape,
    switch_to_portrait,
)
from app.engine.types import Block


# ─────────────────────────────────────────────────────────────
# IMPLEMENTACIÓN COMPARTIDA
# ─────────────────────────────────────────────────────────────


def _table_style(tabla_data: dict) -> dict:
    """Return style metadata accepting both GicaTesis and GicaGen key names."""
    style = tabla_data.get("estilo")
    if not isinstance(style, dict):
        style = tabla_data.get("estilos")
    return style if isinstance(style, dict) else {}


def _alignment_from_value(value: str | None, default=WD_ALIGN_PARAGRAPH.LEFT):
    normalized = str(value or "").strip().lower()
    if normalized in {"center", "centrado", "centrada"}:
        return WD_ALIGN_PARAGRAPH.CENTER
    if normalized in {"right", "derecha"}:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if normalized in {"justify", "justificado", "justificada"}:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return default


def _int_field(value, default: int) -> int:
    if value is None or value == "":
        return default
    return int(value)


def _merge_specs(tabla_data: dict) -> list[dict]:
    """Normalize merge metadata from celdas_fusionadas or celdas_combinadas."""
    raw_merges = tabla_data.get("celdas_fusionadas")
    if isinstance(raw_merges, list) and raw_merges:
        return [merge for merge in raw_merges if isinstance(merge, dict)]

    converted: list[dict] = []
    raw_combined = tabla_data.get("celdas_combinadas", [])
    if not isinstance(raw_combined, list):
        return converted

    for merge in raw_combined:
        if not isinstance(merge, dict):
            continue
        start_row = _int_field(merge.get("fila", merge.get("fila_inicio")), 0)
        end_row = _int_field(merge.get("fila_fin"), start_row)
        start_col = _int_field(merge.get("col_inicio", merge.get("col")), 0)
        end_col = _int_field(merge.get("col_fin"), start_col)
        converted.append(
            {
                "fila": start_row,
                "col": start_col,
                "filas_span": max(1, end_row - start_row + 1),
                "cols_span": max(1, end_col - start_col + 1),
                "texto": merge.get("texto", ""),
                "bold": merge.get("bold"),
                "alignment": merge.get("alignment"),
            }
        )
    return converted


def _is_truthy_border(value) -> bool:
    if isinstance(value, str):
        return value.strip().lower() not in {"", "none", "sin bordes", "false", "0"}
    return bool(value)


def _is_schedule_table(tabla_data: dict) -> bool:
    return str(tabla_data.get("subtipo") or "").strip().lower() == "cronograma_actividades"


def _current_section_is_landscape(doc: Document) -> bool:
    section = doc.sections[-1]
    return section.page_width > section.page_height


def _apply_reduced_landscape_margins(doc: Document, style: dict) -> None:
    if not style.get("margenes_reducidos"):
        return
    section = doc.sections[-1]
    # Medidas tomadas de cronograma_actividades.docx: margenes aprox. 0.8 cm.
    section.left_margin = Cm(0.80)
    section.right_margin = Cm(0.80)
    section.top_margin = Cm(1.10)
    section.bottom_margin = Cm(1.00)


def _apply_portrait_page_setup(doc: Document, style: dict) -> None:
    if str(style.get("orientacion_pagina") or "").strip().lower() != "portrait":
        return
    section = doc.sections[-1]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(float(style.get("margen_izquierdo_cm", 2.0)))
    section.right_margin = Cm(float(style.get("margen_derecho_cm", 2.0)))
    section.top_margin = Cm(float(style.get("margen_superior_cm", 2.0)))
    section.bottom_margin = Cm(float(style.get("margen_inferior_cm", 2.0)))


def _set_cell_margins(cell, margin_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(int(margin_twips)))
        node.set(qn("w:type"), "dxa")


def _compact_cell_after_format(cell, style: dict) -> None:
    margin = style.get("celda_margen_twips")
    if margin not in (None, ""):
        _set_cell_margins(cell, int(margin))
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.95


def _apply_table_geometry(table, widths_cm: list[float] | None) -> None:
    if not widths_cm:
        return
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    total_twips = sum(int(Cm(float(width)).twips) for width in widths_cm)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))

    old_grid = table._tbl.tblGrid
    if old_grid is not None:
        table._tbl.remove(old_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(float(width)).twips)))
        tbl_grid.append(grid_col)
    table._tbl.insert(1, tbl_grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths_cm):
                continue
            width_twips = str(int(Cm(float(widths_cm[idx])).twips))
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), width_twips)
            cell.width = Cm(float(widths_cm[idx]))


def _mark_row_as_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def _row_is_special(tabla_data: dict, row_idx: int) -> bool:
    return row_idx in set(tabla_data.get("filas_fase") or []) or row_idx in set(tabla_data.get("filas_categoria") or []) or row_idx == tabla_data.get("fila_total")


def _cell_alignment(tabla_data: dict, row_idx: int, col_idx: int, style: dict):
    subtype = str(tabla_data.get("subtipo") or "").strip().lower()
    if subtype == "presupuesto_investigacion" and _row_is_special(tabla_data, row_idx):
        if col_idx == 4:
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.CENTER
    if _row_is_special(tabla_data, row_idx):
        return WD_ALIGN_PARAGRAPH.CENTER
    if subtype == "cronograma_actividades":
        return WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
    if subtype == "presupuesto_investigacion":
        if col_idx == 1:
            return WD_ALIGN_PARAGRAPH.LEFT
        if col_idx in {3, 4}:
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT


def _cell_bold(tabla_data: dict, row_idx: int) -> bool:
    return _row_is_special(tabla_data, row_idx)


def _cell_font_size(tabla_data: dict, row_idx: int, col_idx: int, text: str, style: dict, default_size):
    subtype = str(tabla_data.get("subtipo") or "").strip().lower()
    if subtype == "presupuesto_investigacion":
        if row_idx < 0:
            return style.get("fuente_encabezado_pt", default_size)
        if row_idx == tabla_data.get("fila_total"):
            return style.get("fuente_total_pt", default_size)
        if row_idx in set(tabla_data.get("filas_categoria") or []):
            return style.get("fuente_categoria_pt", default_size)
        return default_size
    if not _is_schedule_table(tabla_data):
        return default_size
    if _row_is_special(tabla_data, row_idx):
        return style.get("fuente_fases_pt", default_size)
    if row_idx == 0 or col_idx > 0:
        if str(text or "").strip() == str(tabla_data.get("simbolo_marca") or "").strip():
            return style.get("fuente_marcas_pt", default_size)
        return style.get("fuente_meses_pt", default_size)
    return style.get("fuente_actividades_pt", default_size)


def _tighten_previous_heading_spacing(doc: Document, style: dict) -> None:
    if not style.get("compactar_cronograma"):
        return
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip():
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            break


def _render_exact_or_seq_caption(doc: Document, titulo: str, style: dict) -> None:
    if not titulo:
        return
    if style.get("titulo_exacto"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(float(style.get("titulo_space_before_pt", 0)))
        p.paragraph_format.space_after = Pt(float(style.get("titulo_space_after_pt", 6)))
        font_size = float(style.get("titulo_tamano_pt", 10))
        exact_match = re.match(
            r"^\s*Tabla\s+(\d+)\.(\d+)\s*([.:]?)\s*(.*?)\s*$",
            str(titulo),
            flags=re.IGNORECASE,
        )
        if exact_match:
            chapter_number = int(exact_match.group(1))
            table_number = int(exact_match.group(2))
            punctuation = exact_match.group(3)
            clean_title = exact_match.group(4)

            run_label = p.add_run(f"Tabla {chapter_number}.")
            run_label.bold = False
            run_label.font.size = Pt(font_size)
            run_label.font.name = "Arial"
            add_seq_field(
                p,
                "Tabla",
                reset_to=table_number,
                display_value=table_number,
                font_name="Arial",
                font_size_pt=font_size,
                bold=False,
            )
            suffix = f"{punctuation} " if punctuation else " "
            run_title = p.add_run(f"{suffix}{clean_title}" if clean_title else punctuation)
            run_title.bold = False
            run_title.font.size = Pt(font_size)
            run_title.font.name = "Arial"
        else:
            # Compatibilidad con títulos institucionales no numerados. Se
            # conserva el texto, pero estos casos serán marcados por la
            # validación porque un índice automático requiere SEQ Tabla.
            run = p.add_run(str(titulo))
            run.bold = False
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
        return

    clean_title = re.sub(r"^Tabla\s*[\d.]+\s*[:.]*\s*", "", titulo).strip() or titulo
    chapter_number, table_number, is_first_in_chapter = next_table_number()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(6)
    prefix = f"Tabla {chapter_number}." if chapter_number is not None else "Tabla "
    run_label = p.add_run(prefix)
    run_label.bold = False
    run_label.font.size = Pt(11)
    run_label.font.name = "Arial"
    add_seq_field(
        p,
        "Tabla",
        reset_to=1 if is_first_in_chapter else None,
        display_value=table_number,
        font_name="Arial",
        font_size_pt=11,
        bold=False,
    )
    separator = " " if chapter_number is not None else ". "
    run_title = p.add_run(f"{separator}{clean_title}")
    run_title.bold = False
    run_title.font.size = Pt(11)
    run_title.font.name = "Arial"


def _render_tabla_impl(doc: Document, tabla_data: dict) -> None:
    """Renderiza una tabla completa. Replica ``render_tabla()`` del generador.

    Esta función es el corazón del rendering de tablas y es reutilizada
    por el renderer de matriz y legacy_table.

    Parámetros del dict tabla_data:
        encabezados (list[str]): Fila de encabezados.
        filas (list[list[str]]): Filas de datos.
        orientacion (str): "portrait" | "landscape". Default "portrait".
        titulo (str): Caption con SEQ field.
        nota_pie (str): Nota al pie de la tabla.
        estilo (dict): Overrides de estilo.
        celdas_fusionadas (list[dict]): Merges.
    """
    if not tabla_data:
        return

    encabezados = tabla_data.get("encabezados", [])
    filas = tabla_data.get("filas", [])
    if not encabezados:
        return

    orientacion = (tabla_data.get("orientacion") or "portrait").strip().lower()
    is_landscape = orientacion == "landscape"
    restore_portrait = bool(tabla_data.get("restore_portrait", True))
    titulo = tabla_data.get("titulo")
    nota_pie = tabla_data.get("nota_pie")
    estilo = _table_style(tabla_data)

    header_color = estilo.get("encabezado_color", DEFAULT_HEADER_COLOR)
    font_size = estilo.get(
        "fuente_size",
        estilo.get("fuente_tamano_pt", LANDSCAPE_FONT_SIZE if is_landscape else DEFAULT_TABLE_FONT_SIZE),
    )
    ancho_columnas = estilo.get("ancho_columnas")
    show_borders = _is_truthy_border(estilo.get("bordes", True))
    merges = _merge_specs(tabla_data)

    # 1. Switch to landscape if needed
    already_landscape = _current_section_is_landscape(doc)
    if is_landscape and not already_landscape:
        switch_to_landscape(doc)
    if is_landscape:
        _apply_reduced_landscape_margins(doc, estilo)
    else:
        _apply_portrait_page_setup(doc, estilo)

    # 2. Caption / Title
    _tighten_previous_heading_spacing(doc, estilo)
    _render_exact_or_seq_caption(doc, str(titulo or ""), estilo)

    # 3. Create table
    num_cols = len(encabezados)
    num_rows = 1 + len(filas)
    table = doc.add_table(rows=num_rows, cols=num_cols)

    if show_borders:
        table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 4. Column widths
    if ancho_columnas and len(ancho_columnas) == num_cols:
        for i, w in enumerate(ancho_columnas):
            table.columns[i].width = Cm(float(w))
        _apply_table_geometry(table, [float(width) for width in ancho_columnas])
    else:
        if is_landscape:
            available = 29.7 - LANDSCAPE_MARGINS["left"] - LANDSCAPE_MARGINS["right"]
        else:
            available = 21.0 - PORTRAIT_MARGINS["left"] - PORTRAIT_MARGINS["right"]
        col_width = available / num_cols
        for i in range(num_cols):
            table.columns[i].width = Cm(col_width)
        # Sin esto, ``table.columns[i].width`` es solo una sugerencia que
        # Word no respeta de forma confiable -- sobre todo con celdas
        # combinadas verticalmente (celdas_fusionadas) -- y termina
        # autoajustando cada columna al contenido, dejando la tabla mas
        # angosta que el ancho horizontal disponible real. Igual que en la
        # rama de ancho_columnas explicito, hay que fijar el ancho en cada
        # celda (tcW) para que ocupe todo el ancho de la pagina horizontal.
        _apply_table_geometry(table, [col_width] * num_cols)

    # 5. Headers
    for i, header_text in enumerate(encabezados):
        cell = table.rows[0].cells[i]
        format_cell_text(
            cell,
            header_text,
            _cell_font_size(tabla_data, -1, i, header_text, estilo, font_size),
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            word_sources=tabla_data.get("word_sources") or [],
        )
        if _is_schedule_table(tabla_data):
            _compact_cell_after_format(cell, estilo)
        if header_color:
            apply_cell_shading(cell, header_color)
        set_cell_vertical_alignment(cell)

    # 6. Data rows
    for row_idx, fila in enumerate(filas):
        for col_idx, cell_text in enumerate(fila):
            if col_idx >= num_cols:
                break
            cell = table.rows[row_idx + 1].cells[col_idx]
            format_cell_text(
                cell,
                cell_text or "",
                _cell_font_size(tabla_data, row_idx, col_idx, cell_text or "", estilo, font_size),
                bold=_cell_bold(tabla_data, row_idx),
                alignment=_cell_alignment(tabla_data, row_idx, col_idx, estilo),
                word_sources=tabla_data.get("word_sources") or [],
            )
            if _is_schedule_table(tabla_data):
                _compact_cell_after_format(cell, estilo)
            set_cell_vertical_alignment(cell)

    if _is_schedule_table(tabla_data):
        if table.rows:
            _mark_row_as_table_header(table.rows[0])
        if len(table.rows) > 1:
            _mark_row_as_table_header(table.rows[1])
        for row in table.rows:
            _prevent_row_split(row)

    # 7. Cell merges
    for merge in merges:
        start_row = merge.get("fila", 0) + 1  # +1 header
        start_col = int(merge.get("col", 0) or 0)
        row_span = int(merge.get("filas_span", 1) or 1)
        col_span = int(merge.get("cols_span", 1) or 1)

        if start_row < num_rows and start_col < num_cols:
            end_row = min(start_row + row_span - 1, num_rows - 1)
            end_col = min(start_col + col_span - 1, num_cols - 1)
            start_cell = table.cell(start_row, start_col)
            end_cell = table.cell(end_row, end_col)
            merged_cell = start_cell.merge(end_cell)
            if merge.get("texto") not in (None, ""):
                format_cell_text(
                    merged_cell,
                    str(merge.get("texto") or ""),
                    _cell_font_size(tabla_data, int(merge.get("fila", 0) or 0), start_col, str(merge.get("texto") or ""), estilo, font_size),
                    bold=bool(merge.get("bold", True)),
                    alignment=_alignment_from_value(str(merge.get("alignment") or "center"), WD_ALIGN_PARAGRAPH.CENTER),
                    word_sources=tabla_data.get("word_sources") or [],
                )
                if _is_schedule_table(tabla_data):
                    _compact_cell_after_format(merged_cell, estilo)
                set_cell_vertical_alignment(merged_cell)

    # 8. Footer note
    if nota_pie:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(nota_pie)
        run.italic = True
        run.font.size = Pt(9)
        run.font.name = "Arial"

    # 9. Restore portrait
    if is_landscape and restore_portrait:
        switch_to_portrait(doc)


# ─────────────────────────────────────────────────────────────
# RENDERERS REGISTRADOS
# ─────────────────────────────────────────────────────────────


@register("table")
def render_table(doc: Document, block: Block) -> None:
    """Renderiza una tabla canónica (tipo: 'tabla') completa.

    Block keys (heredados del JSON):
        encabezados, filas, orientacion, titulo, nota_pie,
        estilo, celdas_fusionadas.
    """
    _render_tabla_impl(doc, block)


@register("legacy_table")
def render_legacy_table(doc: Document, block: Block) -> None:
    """Renderiza una tabla legacy (headers/rows) convirtiéndola al formato estándar.

    Block keys:
        tabla (dict): {"headers": [...], "rows": [...]}.
        titulo (str): Título de la tabla.
        nota (str): Nota al pie.
    """
    tabla = block.get("tabla", {})
    headers = tabla.get("headers", [])
    rows = tabla.get("rows", [])
    if not headers and not rows:
        return

    tabla_data = {
        "tipo": "tabla",
        "titulo": block.get("titulo") or tabla.get("titulo", ""),
        "orientacion": "landscape" if len(headers) > 5 else "portrait",
        "encabezados": headers,
        "filas": rows,
    }
    nota = block.get("nota")
    if nota:
        tabla_data["nota_pie"] = nota

    _render_tabla_impl(doc, tabla_data)
