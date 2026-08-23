"""Renderers: paragraph + paragraph_centered."""

from __future__ import annotations

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.engine.registry import register
from app.engine.types import Block
from app.engine.word_bibliography import CITATION_MARKER_RE, render_text_with_citations


@register("paragraph")
def render_paragraph(doc: Document, block: Block) -> None:
    """Renderiza un párrafo justificado normal.

    Block keys:
        text (str): Texto del párrafo.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    text = str(block.get("text", "") or "")
    if CITATION_MARKER_RE.search(text):
        render_text_with_citations(
            p,
            text,
            block.get("word_sources") or [],
            font_name="Arial",
            font_size_pt=12,
        )
    else:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)


@register("paragraph_bold")
def render_paragraph_bold(doc: Document, block: Block) -> None:
    """Renderiza una etiqueta o subtitulo interno sin meterlo al TOC."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(block.get("text", ""))
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(block.get("size", 12))


@register("paragraph_centered")
def render_paragraph_centered(doc: Document, block: Block) -> None:
    """Renderiza un párrafo centrado con formato opcional.

    Block keys:
        text (str): Texto del párrafo.
        bold (bool): Negrita. Default False.
        size (int): Tamaño en pt. Default None (hereda estilo Normal).
        space_before (int): Espacio antes en pt. Default 0.
        space_after (int): Espacio después en pt. Default 0.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    space_before = block.get("space_before", 0)
    space_after = block.get("space_after", 0)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)

    run = p.add_run(block.get("text", ""))
    if block.get("bold"):
        run.bold = True
    run.font.size = Pt(block.get("size", 12))
    run.font.name = "Arial"
