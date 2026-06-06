"""Renderer: info_basica_unac_maestria."""

from __future__ import annotations

import logging
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from app.engine.registry import register
from app.engine.primitives import add_p_centered
from app.engine.types import Block

logger = logging.getLogger(__name__)

def _add_p_left(doc: Document, label: str, value: str = None, space_before: int = 0, indent: float = 0.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)

    if label:
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Arial"
    
    if value:
        run_v = p.add_run(" " + value if label else value)
        run_v.bold = False
        run_v.font.size = Pt(12)
        run_v.font.name = "Arial"


@register("info_basica_unac_maestria")
def render_info_basica_unac_maestria(doc: Document, block: Block) -> None:
    data = block.get("data", {})
    info = block.get("info", {})
    
    # 1. INFORMACION BASICA
    titulo = info.get("titulo", "INFORMACIÓN BÁSICA")
    add_p_centered(doc, titulo, bold=True, size=14, space_before=12, space_after=18)
    
    # AI CONTENT INJECTION
    ai_content = info.get("_ai_content")
    if ai_content:
        # If the AI provided content for this section, render it prominently
        from app.engine.primitives import add_ai_paragraph
        if isinstance(ai_content, str):
            add_ai_paragraph(doc, ai_content)
        elif isinstance(ai_content, list):
            # We would need to import the registry or renderers to handle structured content here
            # For simplicity in this specific renderer, we just handle the text
            pass

    # 2-7. FACULTAD, U. INVESTIGACION, TITULO (label one line, value next line)
    # Proactive search in multiple pockets (info, data, values)
    def _find_val(key: str, default: str) -> str:
        candidates = [
            info.get(key),
            data.get(key),
            (data.get("values") or {}).get(key)
        ]
        return next((str(v).strip() for v in candidates if v and str(v).strip()), default)

    _add_p_left(doc, "FACULTAD:", space_before=12)
    _add_p_left(doc, "", _find_val("facultad", "[Nombre de Facultad]"))
    
    _add_p_left(doc, "UNIDAD DE INVESTIGACIÓN:", space_before=12)
    _add_p_left(doc, "", _find_val("unidad_investigacion", "[Nombre Unidad de Investigación]"))
    
    _add_p_left(doc, "TÍTULO:", space_before=12)
    _add_p_left(doc, "", _find_val("titulo_investigacion", _find_val("titulo", "[Título de tesis]")))
    
    # AUTORES
    _add_p_left(doc, "AUTORES:", space_before=12)
    # Deep search for authors
    autores = info.get("autores") or data.get("autores")
    if not autores:
        # Fallback to single author field
        solo_autor = _find_val("autor_valor", None)
        if solo_autor:
            autores = [{"nombre": solo_autor, "dni": "[DNI]", "orcid": "[ORCID]"}]
        else:
            autores = [{"nombre": "Bach. [Apellidos y nombres]", "dni": "[00000000]", "orcid": "[0000-0000-0000-0000]"}]
    for autor in autores:
        p_autor = doc.add_paragraph()
        p_autor.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_autor.paragraph_format.left_indent = Inches(0.4)
        p_autor.paragraph_format.line_spacing = 1.5
        r_nombre = p_autor.add_run(autor.get("nombre", "[Nombre]"))
        r_nombre.font.name = "Arial"
        r_nombre.font.size = Pt(12)
        
        p_dni_orcid = doc.add_paragraph()
        p_dni_orcid.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_dni_orcid.paragraph_format.left_indent = Inches(0.4)
        p_dni_orcid.paragraph_format.line_spacing = 1.5
        p_dni_orcid.paragraph_format.space_after = Pt(6)

        r_dni_l = p_dni_orcid.add_run("DNI: ")
        r_dni_l.bold = True
        r_dni_l.font.name = "Arial"
        r_dni_l.font.size = Pt(12)

        r_dni_v = p_dni_orcid.add_run(f"{autor.get('dni', '[DNI]')} / ")
        r_dni_v.font.name = "Arial"
        r_dni_v.font.size = Pt(12)

        r_orcid_l = p_dni_orcid.add_run("ORCID: ")
        r_orcid_l.bold = True
        r_orcid_l.font.name = "Arial"
        r_orcid_l.font.size = Pt(12)

        r_orcid_v = p_dni_orcid.add_run(autor.get("orcid", "[ORCID]"))
        r_orcid_v.font.name = "Arial"
        r_orcid_v.font.size = Pt(12)

    # ASESOR
    _add_p_left(doc, "ASESOR:")
    asesor = info.get("asesor", {"nombre": "Dr. [Apellidos y nombres]", "dni": "[00000000]", "orcid": "[0000-0000-0000-0000]"})
    
    p_asesor = doc.add_paragraph()
    p_asesor.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_asesor.paragraph_format.left_indent = Inches(0.4)
    p_asesor.paragraph_format.line_spacing = 1.5
    r_asesor_nombre = p_asesor.add_run(asesor.get("nombre", "[Nombre]"))
    r_asesor_nombre.font.name = "Arial"
    r_asesor_nombre.font.size = Pt(12)

    p_asesor_dni = doc.add_paragraph()
    p_asesor_dni.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_asesor_dni.paragraph_format.left_indent = Inches(0.4)
    p_asesor_dni.paragraph_format.line_spacing = 1.5
    p_asesor_dni.paragraph_format.space_after = Pt(12)

    r_adni_l = p_asesor_dni.add_run("DNI: ")
    r_adni_l.bold = True
    r_adni_l.font.name = "Arial"
    r_adni_l.font.size = Pt(12)

    r_adni_v = p_asesor_dni.add_run(f"{asesor.get('dni', '[DNI]')} / ")
    r_adni_v.font.name = "Arial"
    r_adni_v.font.size = Pt(12)

    r_aorcid_l = p_asesor_dni.add_run("ORCID: ")
    r_aorcid_l.bold = True
    r_aorcid_l.font.name = "Arial"
    r_aorcid_l.font.size = Pt(12)

    r_aorcid_v = p_asesor_dni.add_run(asesor.get("orcid", "[ORCID]"))
    r_aorcid_v.font.name = "Arial"
    r_aorcid_v.font.size = Pt(12)

    # LUGAR DE EJECUCION, UNIDAD, TIPO, ENFOQUE, DISENO (label and value on same line)
    _add_p_left(doc, "LUGAR DE EJECUCIÓN:", info.get("lugar_ejecucion", "[Lugar de ejecución]"))
    _add_p_left(doc, "UNIDAD DE ANÁLISIS:", info.get("unidad_analisis", "[Unidad de análisis]"))
    _add_p_left(doc, "TIPO:", info.get("tipo", "Aplicada"))
    _add_p_left(doc, "ENFOQUE:", info.get("enfoque", "Cualitativo/Cuantitativo"))
    _add_p_left(doc, "DISEÑO DE INVESTIGACIÓN:", info.get("diseno_investigacion", "[Diseño]"))
    
    # TEMA OCDE
    ocde = info.get("tema_ocde", [])
    if isinstance(ocde, str):
        ocde = [ocde]

    p_ocde = doc.add_paragraph()
    p_ocde.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_ocde.paragraph_format.space_before = Pt(12)
    p_ocde.paragraph_format.line_spacing = 1.5
    
    r_ocde_label = p_ocde.add_run("TEMA OCDE: ")
    r_ocde_label.bold = True
    r_ocde_label.font.name = "Arial"
    r_ocde_label.font.size = Pt(12)

    if ocde:
        r_ocde_val = p_ocde.add_run(ocde[0])
        r_ocde_val.font.name = "Arial"
        r_ocde_val.font.size = Pt(12)
        
        for idx in range(1, len(ocde)):
            p_o = doc.add_paragraph()
            p_o.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_o.paragraph_format.line_spacing = 1.5
            p_o.paragraph_format.left_indent = Inches(1.1 + (idx * 0.4))
            r_o = p_o.add_run(ocde[idx])
            r_o.font.name = "Arial"
            r_o.font.size = Pt(12)

    # La paginacion de esta seccion debe controlarse en el normalizer.
    # Si este renderer agrega un salto propio, al combinarse con el
    # section_break de preliminares produce una hoja en blanco.
