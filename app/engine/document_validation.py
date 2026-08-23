"""Structural quality gates for generated UNAC project documents."""

from __future__ import annotations

import re
from collections import Counter
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx.document import Document
from docx.oxml.ns import qn

from app.engine.normalizer import _ABBR_USED_TOKEN_RE, _COMMON_ABBREVIATIONS, _norm_upper


class DocumentIndexValidationError(ValueError):
    """Raised when Word fields and generated index sources do not correspond."""


_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def validate_docx_field_safety(path: str | Path) -> dict[str, int]:
    """Validate that cached native fields do not force an update-on-open dialog."""
    with ZipFile(path, "r") as archive:
        names = set(archive.namelist())
        settings_xml = archive.read("word/settings.xml")
        settings = ET.fromstring(settings_xml)
        update_nodes = settings.findall(f".//{{{_WORD_NS}}}updateFields")
        enabled = [
            node
            for node in update_nodes
            if str(node.get(f"{{{_WORD_NS}}}val") or "true").strip().lower()
            not in {"0", "false", "off", "no"}
        ]
        if enabled:
            raise DocumentIndexValidationError("updateFields debe estar ausente o desactivado")

        external_targets: list[str] = []
        for name in names:
            if not name.endswith(".rels"):
                continue
            root = ET.fromstring(archive.read(name))
            for relation in root.findall(f"{{{_PACKAGE_REL_NS}}}Relationship"):
                if str(relation.get("TargetMode") or "").strip().lower() == "external":
                    external_targets.append(str(relation.get("Target") or ""))
        if external_targets:
            raise DocumentIndexValidationError(
                "el DOCX contiene relaciones externas: " + ", ".join(external_targets)
            )

        document = ET.fromstring(archive.read("word/document.xml"))
        instructions = [
            str(node.text or "").strip()
            for node in document.iter(f"{{{_WORD_NS}}}instrText")
            if str(node.text or "").strip()
        ]
        return {
            "native_fields": len(instructions),
            "external_relationships": len(external_targets),
            "update_fields_enabled": len(enabled),
        }


def _field_instructions(doc: Document) -> list[str]:
    return [
        str(element.text or "").strip()
        for element in doc.element.body.iter(qn("w:instrText"))
        if str(element.text or "").strip()
    ]


def _has_field(instructions: Iterable[str], pattern: str) -> bool:
    regex = re.compile(pattern, re.IGNORECASE)
    return any(regex.search(instruction) for instruction in instructions)


def _count_field(instructions: Iterable[str], pattern: str) -> int:
    regex = re.compile(pattern, re.IGNORECASE)
    return sum(1 for instruction in instructions if regex.search(instruction))


def _expected_caption_counts(blocks: list[dict[str, Any]]) -> tuple[int, int]:
    tables = 0
    figures = 0
    for block in blocks:
        if not isinstance(block, dict):
            continue
        block_type = str(block.get("type") or "").strip().lower()
        if block_type == "table":
            if str(block.get("titulo") or "").strip() and block.get("encabezados"):
                tables += 1
        elif block_type == "image":
            if (
                str(block.get("titulo") or "").strip()
                and str(block.get("ruta") or "").strip()
                and not bool(block.get("omit_caption"))
            ):
                figures += 1
    return tables, figures


def _rendered_document_fragments(doc: Document) -> list[str]:
    fragments = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    fragments.append(text)
    return fragments


def _validate_repeated_abbreviations(doc: Document, blocks: list[dict[str, Any]]) -> list[str]:
    indexed = {
        _norm_upper(row.get("sigla", ""))
        for block in blocks
        if isinstance(block, dict)
        and str(block.get("type") or "").strip().lower() == "abbreviations_table"
        for row in block.get("rows", [])
        if isinstance(row, dict)
    }
    if not indexed:
        return []

    token_counts: Counter[str] = Counter()
    for fragment in _rendered_document_fragments(doc):
        for token in _ABBR_USED_TOKEN_RE.findall(fragment):
            normalized = _norm_upper(token)
            if normalized in _COMMON_ABBREVIATIONS:
                token_counts[normalized] += 1

    return sorted(token for token, count in token_counts.items() if count >= 2 and token not in indexed)


def validate_unac_project_document(doc: Document, blocks: list[dict[str, Any]]) -> dict[str, int]:
    """Validate captions, automatic lists, typography and abbreviation coverage."""
    instructions = _field_instructions(doc)
    expected_tables, expected_figures = _expected_caption_counts(blocks)
    seq_tables = _count_field(instructions, r"\bSEQ\s+Tabla\b")
    seq_figures = _count_field(instructions, r"\bSEQ\s+Figura\b")
    has_table_index = _has_field(instructions, r"\bTOC\b.*\\c\s+\"Tabla\"")
    has_figure_index = _has_field(instructions, r"\bTOC\b.*\\c\s+\"Figura\"")

    errors: list[str] = []
    expected_formulas = sum(
        1
        for block in blocks
        if isinstance(block, dict) and str(block.get("type") or "").strip().lower() == "formula"
    )
    rendered_formulas = len(doc.element.body.xpath(".//m:oMath"))
    if rendered_formulas != expected_formulas:
        errors.append(
            f"formulas OMML incompatibles: esperadas={expected_formulas}, renderizadas={rendered_formulas}"
        )
    if seq_tables != expected_tables:
        errors.append(
            f"captions de tabla incompatibles: esperados={expected_tables}, SEQ Tabla={seq_tables}"
        )
    if seq_figures != expected_figures:
        errors.append(
            f"captions de figura incompatibles: esperados={expected_figures}, SEQ Figura={seq_figures}"
        )
    if expected_tables and not has_table_index:
        errors.append("faltante el índice automático TOC \\c \"Tabla\"")
    if expected_figures and not has_figure_index:
        errors.append("faltante el índice automático TOC \\c \"Figura\"")

    automatic_index_count = int(has_table_index) + int(has_figure_index)
    page_headers = sum(1 for paragraph in doc.paragraphs if paragraph.text.strip() == "Pág.")
    if page_headers != automatic_index_count:
        errors.append(
            f"encabezados Pág. incompatibles: esperados={automatic_index_count}, encontrados={page_headers}"
        )

    for paragraph in doc.paragraphs:
        paragraph_instructions = [
            str(element.text or "") for element in paragraph._p.iter(qn("w:instrText"))
        ]
        if not any(re.search(r"\bSEQ\s+(?:Tabla|Figura)\b", item, re.IGNORECASE) for item in paragraph_instructions):
            continue
        if any(run.bold is True for run in paragraph.runs):
            errors.append(f"caption con negrita directa: {paragraph.text.strip()}")

    missing_abbreviations = _validate_repeated_abbreviations(doc, blocks)
    if missing_abbreviations:
        errors.append(
            "siglas técnicas repetidas sin entrada en el índice: " + ", ".join(missing_abbreviations)
        )

    if errors:
        raise DocumentIndexValidationError("; ".join(errors))

    return {
        "table_captions": seq_tables,
        "figure_captions": seq_figures,
        "page_headers": page_headers,
        "formulas": rendered_formulas,
        "indexed_abbreviations": sum(
            len(block.get("rows", []))
            for block in blocks
            if isinstance(block, dict) and block.get("type") == "abbreviations_table"
        ),
    }
