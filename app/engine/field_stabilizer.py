"""Stabilize cached TOC results so DOCX fields are useful without an open prompt."""

from __future__ import annotations

import os
import re
import tempfile
import unicodedata
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pypdf import PdfReader

from app.core.pdf_converter import convert_docx_to_pdf
from app.engine.primitives import disable_update_fields


class FieldStabilizationError(RuntimeError):
    """Raised when cached page results do not converge."""


def _uses_word_com() -> bool:
    return os.name == "nt"


def _norm(value: str) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(character for character in text if not unicodedata.combining(character))
    return " ".join(re.sub(r"[^a-z0-9]+", " ", text.lower()).split())


def _pdf_pages(path: Path) -> list[str]:
    return [_norm(page.extract_text() or "") for page in PdfReader(str(path)).pages]


def _instruction(paragraph) -> str:
    return " ".join(
        str(node.text or "").strip()
        for node in paragraph._p.iter(qn("w:instrText"))
        if str(node.text or "").strip()
    )


def _cached_lines(paragraph) -> list[str]:
    text = str(paragraph.text or "").replace("\r", "\n")
    return [line.strip() for line in text.split("\n") if line.strip()]


def _targets(lines: Iterable[str]) -> list[str]:
    targets: list[str] = []
    for line in lines:
        candidate = re.sub(r"\s+\d+\s*$", "", line.replace("\t", " ")).strip()
        if candidate and candidate != "Sin entradas aplicables":
            targets.append(candidate)
    return targets


def _last_matching_page(pages: list[str], target: str) -> int | None:
    needle = _norm(target)
    if not needle:
        return None
    candidates = [index + 1 for index, page in enumerate(pages) if needle in page]
    if not candidates and len(needle) > 70:
        prefix = " ".join(needle.split()[:12])
        candidates = [index + 1 for index, page in enumerate(pages) if prefix in page]
    return max(candidates) if candidates else None


def _result_run(text: str, *, tab: bool = False, break_after: bool = False):
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for name in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{name}"), "Arial")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    props.extend([fonts, size])
    run.append(props)
    value = OxmlElement("w:t")
    value.set(qn("xml:space"), "preserve")
    value.text = text
    run.append(value)
    if tab:
        run.append(OxmlElement("w:tab"))
    if break_after:
        run.append(OxmlElement("w:br"))
    return run


def _replace_field_result(paragraph, entries: list[tuple[str, int]]) -> None:
    children = list(paragraph._p)
    separate_index = end_index = None
    for index, child in enumerate(children):
        field_types = [
            str(node.get(qn("w:fldCharType")) or "").lower()
            for node in child.iter(qn("w:fldChar"))
        ]
        if "separate" in field_types:
            separate_index = index
        if "end" in field_types and separate_index is not None:
            end_index = index
            break
    if separate_index is None or end_index is None:
        raise FieldStabilizationError("campo TOC sin separador o cierre")
    for child in children[separate_index + 1 : end_index]:
        paragraph._p.remove(child)
    end_node = list(paragraph._p)[separate_index + 1]
    insertion = paragraph._p.index(end_node)
    for index, (title, page) in enumerate(entries):
        paragraph._p.insert(insertion, _result_run(title, tab=True))
        insertion += 1
        paragraph._p.insert(
            insertion,
            _result_run(str(page), break_after=index < len(entries) - 1),
        )
        insertion += 1


def _update_cached_results(docx_path: Path, pdf_path: Path) -> tuple[tuple[str, int], ...]:
    pages = _pdf_pages(pdf_path)
    doc = Document(str(docx_path))
    signature: list[tuple[str, int]] = []
    changed = False
    for paragraph in doc.paragraphs:
        if not re.search(r"\bTOC\b", _instruction(paragraph), re.IGNORECASE):
            continue
        targets = _targets(_cached_lines(paragraph))
        entries: list[tuple[str, int]] = []
        for target in targets:
            page = _last_matching_page(pages, target)
            if page is None:
                raise FieldStabilizationError(f"no se encontro en el PDF la entrada de indice: {target}")
            entries.append((target, page))
            signature.append((target, page))
        if entries:
            _replace_field_result(paragraph, entries)
            changed = True
    if changed:
        disable_update_fields(doc)
        doc.save(str(docx_path))
    return tuple(signature)


def stabilize_docx_fields(docx_path: str | Path, *, max_cycles: int = 3) -> dict[str, int]:
    """Render and rewrite cached index pages until their mapping is stable."""
    source = Path(docx_path).resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    previous: tuple[tuple[str, int], ...] | None = None
    with tempfile.TemporaryDirectory(prefix="gicatesis-fields-") as directory:
        pdf = Path(directory) / "provisional.pdf"
        for cycle in range(1, max_cycles + 1):
            convert_docx_to_pdf(str(source), str(pdf), timeout=180.0)
            # Word COM updates and saves native fields itself. The custom
            # rewrite is needed only for LibreOffice/Docker, which exports a
            # PDF without persisting those calculated results in the DOCX.
            if _uses_word_com():
                doc = Document(str(source))
                disable_update_fields(doc)
                doc.save(str(source))
                return {"cycles": cycle, "entries": 0}
            current = _update_cached_results(source, pdf)
            if previous == current:
                return {"cycles": cycle, "entries": len(current)}
            previous = current
    raise FieldStabilizationError(
        f"la paginacion de los indices no convergio despues de {max_cycles} ciclos"
    )
