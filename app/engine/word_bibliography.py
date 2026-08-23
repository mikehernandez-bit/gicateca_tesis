"""Native Microsoft Word citations, bibliography fields and source storage."""

from __future__ import annotations

import os
import re
import unicodedata
import uuid
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BIBLIOGRAPHY_NS = "http://schemas.openxmlformats.org/officeDocument/2006/bibliography"
CUSTOM_XML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/customXml"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
WORD_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WORDPROCESSING_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
SPANISH_PERU_LCID = 3082

SOURCE_MARKER_RE = re.compile(r"^\s*\[\[SOURCE:([A-Z0-9_-]+)\]\]\s*(.+?)\s*$", re.DOTALL)
CITATION_MARKER_RE = re.compile(r"\[\[CITE:([A-Z0-9_-]+(?:;[A-Z0-9_-]+)*)\]\]")
_AUTHOR_YEAR_RE = re.compile(r"^(.+?)\s+\(((?:19|20)\d{2})\)\.\s+(.+)$", re.DOTALL)
_AUTHOR_RE = re.compile(
    r"(?<![A-Za-z0-9])([A-Za-zÁÉÍÓÚÜÑáéíóúüñ'-]+),\s*"
    r"([A-ZÁÉÍÓÚÜÑ](?:\.[A-ZÁÉÍÓÚÜÑ])?\.)"
)
_TECHNICAL_AUTHOR_RE = re.compile(
    r"^(?P<name>(?:MIL-STD-\d+[A-Z]?|"
    r"ISO(?:\s+[A-Z0-9][A-Z0-9.:/\-]*)?|"
    r"IEC(?:\s+[A-Z0-9][A-Z0-9.:/\-]*)?|"
    r"EN(?:\s+[A-Z0-9][A-Z0-9.:/\-]*)?|"
    r"SAE(?:\s+[A-Z0-9][A-Z0-9.:/\-]*)?|GMG))"
    r"(?:,\s*[A-Z0-9]\.?)?$",
    re.IGNORECASE,
)
_JOURNAL_RE = re.compile(
    r"^(.+?),\s*(\d+)\((\d+)\),\s*([A-Za-z0-9–—-]+)\.?$",
    re.DOTALL,
)
_SIMULATION_NOTE_RE = re.compile(
    r"\s*Referencia propuesta simulada para validacion del autor\.?\s*$",
    re.IGNORECASE,
)


class NativeBibliographyValidationError(ValueError):
    """Raised when a generated DOCX has incomplete native bibliography wiring."""


def _reference_paragraphs(value: Any) -> list[str]:
    if isinstance(value, str):
        normalized = value.replace("\r\n", "\n").replace("\r", "\n")
        return [part.strip() for part in re.split(r"\n\s*\n", normalized) if part.strip()]
    if isinstance(value, list):
        paragraphs: list[str] = []
        for item in value:
            if isinstance(item, str) and item.strip():
                paragraphs.extend(_reference_paragraphs(item))
            elif isinstance(item, dict):
                for key in ("texto", "text", "contenido"):
                    text = item.get(key)
                    if isinstance(text, str) and text.strip():
                        paragraphs.extend(_reference_paragraphs(text))
                        break
        return paragraphs
    return []


def _parse_authors(author_text: str) -> list[dict[str, str]]:
    normalized = " ".join(str(author_text or "").split()).strip(" ,.&")
    technical = _TECHNICAL_AUTHOR_RE.fullmatch(normalized)
    if technical:
        # A technical standard is a corporate author. Older GICA payloads may
        # contain a synthetic initial (``MIL-STD-1629A, M.``); strip only that
        # suffix so retries of already-saved projects remain renderable.
        return [{"last": technical.group("name").upper(), "first": ""}]
    authors = [
        {"last": last.strip(), "first": first.strip()}
        for last, first in _AUTHOR_RE.findall(author_text)
    ]
    if authors:
        return authors
    fallback = normalized
    return [{"last": fallback or "Autor simulado", "first": ""}]


def _split_reference_body(body: str) -> tuple[str, str]:
    clean = _SIMULATION_NOTE_RE.sub("", body).strip()
    match = re.match(r"^(.+?)\.\s+(.+)$", clean, re.DOTALL)
    if not match:
        return clean.rstrip("."), "Editorial académica simulada"
    return match.group(1).strip(), match.group(2).strip()


def parse_simulated_source(tag: str, reference_text: str) -> dict[str, Any] | None:
    """Parse GICA's controlled simulated reference into Word source metadata."""
    normalized_reference = " ".join(str(reference_text or "").split())
    match = _AUTHOR_YEAR_RE.match(normalized_reference)
    if not match:
        return None

    author_text, year, body = match.groups()
    authors = _parse_authors(author_text)
    visible_author = author_text
    if (
        len(authors) == 1
        and not str(authors[0].get("first") or "").strip()
        and _TECHNICAL_AUTHOR_RE.fullmatch(" ".join(author_text.split()).strip(" ,.&"))
    ):
        visible_author = str(authors[0].get("last") or author_text)
    title, publication = _split_reference_body(body)
    journal_match = _JOURNAL_RE.match(publication)
    source: dict[str, Any] = {
        "tag": tag,
        "guid": "{" + str(uuid.uuid5(uuid.NAMESPACE_URL, f"gica:{tag}")).upper() + "}",
        "lcid": SPANISH_PERU_LCID,
        "authors": authors,
        "title": title,
        "year": year,
        "comments": (
            "FUENTE SIMULADA GENERADA POR GICA PARA PRUEBAS. "
            "Debe validarse, corregirse o reemplazarse antes de una entrega académica."
        ),
        "original_text": f"{visible_author} ({year}). {body}",
    }
    if journal_match:
        journal, volume, issue, pages = journal_match.groups()
        source.update(
            {
                "source_type": "JournalArticle",
                "journal_name": journal.strip(),
                "volume": volume,
                "issue": issue,
                "pages": pages.replace("–", "-").replace("—", "-"),
            }
        )
    else:
        source.update(
            {
                "source_type": "Book",
                "city": "Lima",
                "publisher": publication.rstrip("."),
            }
        )
    return source


def extract_word_sources_from_finales(finales: Any) -> list[dict[str, Any]]:
    if not isinstance(finales, dict):
        return []
    referencias = finales.get("referencias")
    if isinstance(referencias, dict):
        content = referencias.get("_ai_content")
    else:
        content = referencias

    sources: list[dict[str, Any]] = []
    seen: set[str] = set()
    for paragraph in _reference_paragraphs(content):
        marker = SOURCE_MARKER_RE.match(paragraph)
        if not marker:
            continue
        tag, reference_text = marker.groups()
        if tag in seen:
            continue
        source = parse_simulated_source(tag, reference_text)
        if source is None:
            continue
        seen.add(tag)
        sources.append(source)
    return sources


def extract_word_sources(data: Any) -> list[dict[str, Any]]:
    if not isinstance(data, dict):
        return []
    return extract_word_sources_from_finales(data.get("finales"))


def strip_source_markers(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    return re.sub(r"\[\[SOURCE:[A-Z0-9_-]+\]\]\s*", "", value)


def extract_reference_preamble(value: Any) -> list[str]:
    """Return explanatory paragraphs, excluding machine-readable sources."""
    return [
        paragraph
        for paragraph in _reference_paragraphs(value)
        if not SOURCE_MARKER_RE.match(paragraph)
    ]


def _author_citation(source: dict[str, Any]) -> str:
    authors = source.get("authors") or []
    last_names = [str(author.get("last") or "").strip() for author in authors if isinstance(author, dict)]
    last_names = [name for name in last_names if name]
    if not last_names:
        return "Autor simulado"
    if len(last_names) == 1:
        return last_names[0]
    if len(last_names) == 2:
        return f"{last_names[0]} & {last_names[1]}"
    return f"{last_names[0]} et al."


def citation_display(tags: Iterable[str], source_map: dict[str, dict[str, Any]]) -> str:
    entries: list[str] = []
    for tag in tags:
        source = source_map.get(tag)
        if source is None:
            entries.append(tag)
            continue
        entries.append(f"{_author_citation(source)}, {source.get('year') or 's. f.'}")
    return "(" + "; ".join(entries) + ")"


def _style_run(run: Any, *, font_name: str, font_size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _append_field(
    paragraph: Any,
    instruction: str,
    display: str,
    *,
    font_name: str,
    font_size_pt: float,
) -> None:
    begin_run = paragraph.add_run()
    _style_run(begin_run, font_name=font_name, font_size_pt=font_size_pt)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    _style_run(instruction_run, font_name=font_name, font_size_pt=font_size_pt)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instruction_run._r.append(instr)

    separate_run = paragraph.add_run()
    _style_run(separate_run, font_name=font_name, font_size_pt=font_size_pt)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    display_run = paragraph.add_run(display)
    _style_run(display_run, font_name=font_name, font_size_pt=font_size_pt)

    end_run = paragraph.add_run()
    _style_run(end_run, font_name=font_name, font_size_pt=font_size_pt)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def render_text_with_citations(
    paragraph: Any,
    text: str,
    sources: Iterable[dict[str, Any]],
    *,
    font_name: str = "Arial",
    font_size_pt: float = 12,
) -> int:
    source_map = {str(source.get("tag") or ""): source for source in sources}
    position = 0
    citation_count = 0
    for match in CITATION_MARKER_RE.finditer(str(text or "")):
        prefix = str(text or "")[position : match.start()]
        if prefix:
            run = paragraph.add_run(prefix)
            _style_run(run, font_name=font_name, font_size_pt=font_size_pt)
        tags = [tag for tag in match.group(1).split(";") if tag]
        if tags:
            instruction = f" CITATION {tags[0]} \\l {SPANISH_PERU_LCID}"
            instruction += "".join(f" \\m {tag}" for tag in tags[1:])
            # Word may add the source title when an author occurs in more than
            # one entry, even when the years differ.  The native \t switch
            # suppresses only that title and preserves the expected APA form:
            # (Autor, año) or (Autor, año; Autor, año).
            instruction += " \\t "
            _append_field(
                paragraph,
                instruction,
                citation_display(tags, source_map),
                font_name=font_name,
                font_size_pt=font_size_pt,
            )
            citation_count += 1
        position = match.end()

    suffix = str(text or "")[position:]
    if suffix:
        run = paragraph.add_run(suffix)
        _style_run(run, font_name=font_name, font_size_pt=font_size_pt)
    return citation_count


def bibliography_display_entries(sources: Iterable[dict[str, Any]]) -> list[str]:
    entries: list[str] = []
    for source in sources:
        original = " ".join(str(source.get("original_text") or "").split())
        if original:
            entry = _SIMULATION_NOTE_RE.sub("", original).strip()
        else:
            authors = _author_citation(source)
            year = str(source.get("year") or "s. f.")
            title = str(source.get("title") or "Fuente académica simulada").strip()
            if source.get("source_type") == "JournalArticle":
                publication = str(source.get("journal_name") or "Revista académica simulada").strip()
                volume = str(source.get("volume") or "").strip()
                issue = str(source.get("issue") or "").strip()
                pages = str(source.get("pages") or "").strip()
                detail = f"{publication}, {volume}({issue}), {pages}".strip(" ,")
            else:
                detail = str(source.get("publisher") or "Editorial académica simulada").strip()
            entry = f"{authors} ({year}). {title}. {detail}."
        if entry:
            entries.append(entry)
    return sorted(dict.fromkeys(entries), key=lambda item: item.casefold())


def add_bibliography_field(paragraph: Any, sources: Iterable[dict[str, Any]] = ()) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    entries = bibliography_display_entries(sources)
    fallback = "\n\n".join(entries) or "Bibliografía nativa pendiente de actualización en Microsoft Word."
    _append_field(
        paragraph,
        f" BIBLIOGRAPHY \\l {SPANISH_PERU_LCID} ",
        fallback,
        font_name="Arial",
        font_size_pt=12,
    )


def _subelement(parent: ET.Element, name: str, text: Any = None) -> ET.Element:
    element = ET.SubElement(parent, f"{{{BIBLIOGRAPHY_NS}}}{name}")
    if text not in (None, ""):
        element.text = str(text)
    return element


def build_sources_xml(sources: list[dict[str, Any]]) -> bytes:
    ET.register_namespace("b", BIBLIOGRAPHY_NS)
    root = ET.Element(
        f"{{{BIBLIOGRAPHY_NS}}}Sources",
        {
            "SelectedStyle": "\\APA.XSL",
            "StyleName": "APA",
        },
    )
    for order, source in enumerate(sources, start=1):
        node = _subelement(root, "Source")
        _subelement(node, "Tag", source.get("tag"))
        _subelement(node, "SourceType", source.get("source_type") or "Book")
        _subelement(node, "Guid", source.get("guid"))
        _subelement(node, "LCID", source.get("lcid") or SPANISH_PERU_LCID)

        author_outer = _subelement(node, "Author")
        author_inner = _subelement(author_outer, "Author")
        name_list = _subelement(author_inner, "NameList")
        for author in source.get("authors") or []:
            if not isinstance(author, dict):
                continue
            person = _subelement(name_list, "Person")
            _subelement(person, "Last", author.get("last"))
            _subelement(person, "First", author.get("first"))

        _subelement(node, "Title", source.get("title"))
        _subelement(node, "Year", source.get("year"))
        if source.get("source_type") == "JournalArticle":
            _subelement(node, "JournalName", source.get("journal_name"))
            _subelement(node, "Volume", source.get("volume"))
            _subelement(node, "Issue", source.get("issue"))
            _subelement(node, "Pages", source.get("pages"))
        else:
            _subelement(node, "City", source.get("city") or "Lima")
            _subelement(node, "Publisher", source.get("publisher"))
        _subelement(node, "Comments", source.get("comments"))
        _subelement(node, "RefOrder", order)

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _find_sources_part(entries: dict[str, bytes]) -> str | None:
    for name, payload in entries.items():
        if not name.startswith("customXml/item") or not name.endswith(".xml"):
            continue
        if "itemProps" in name:
            continue
        try:
            root = ET.fromstring(payload)
        except ET.ParseError:
            continue
        if _local_name(root.tag) == "Sources" and "bibliography" in root.tag:
            return name
    return None


def _next_relationship_id(root: ET.Element) -> str:
    used = {str(item.get("Id") or "") for item in root}
    number = 1
    while f"rId{number}" in used:
        number += 1
    return f"rId{number}"


def _add_sources_part(entries: dict[str, bytes], sources_xml: bytes) -> str:
    indices = [
        int(match.group(1))
        for name in entries
        if (match := re.fullmatch(r"customXml/item(\d+)\.xml", name))
    ]
    index = max(indices, default=0) + 1
    source_name = f"customXml/item{index}.xml"
    props_name = f"customXml/itemProps{index}.xml"
    rels_name = f"customXml/_rels/item{index}.xml.rels"
    entries[source_name] = sources_xml

    ET.register_namespace("ds", CUSTOM_XML_NS)
    props = ET.Element(
        f"{{{CUSTOM_XML_NS}}}datastoreItem",
        {f"{{{CUSTOM_XML_NS}}}itemID": "{" + str(uuid.uuid4()).upper() + "}"},
    )
    refs = ET.SubElement(props, f"{{{CUSTOM_XML_NS}}}schemaRefs")
    ET.SubElement(
        refs,
        f"{{{CUSTOM_XML_NS}}}schemaRef",
        {f"{{{CUSTOM_XML_NS}}}uri": BIBLIOGRAPHY_NS},
    )
    entries[props_name] = ET.tostring(props, encoding="utf-8", xml_declaration=True)

    ET.register_namespace("", PACKAGE_REL_NS)
    rels = ET.Element(f"{{{PACKAGE_REL_NS}}}Relationships")
    ET.SubElement(
        rels,
        f"{{{PACKAGE_REL_NS}}}Relationship",
        {
            "Id": "rId1",
            "Type": f"{WORD_REL_NS}/customXmlProps",
            "Target": f"itemProps{index}.xml",
        },
    )
    entries[rels_name] = ET.tostring(rels, encoding="utf-8", xml_declaration=True)

    document_rels_name = "word/_rels/document.xml.rels"
    document_rels = ET.fromstring(entries[document_rels_name])
    ET.SubElement(
        document_rels,
        f"{{{PACKAGE_REL_NS}}}Relationship",
        {
            "Id": _next_relationship_id(document_rels),
            "Type": f"{WORD_REL_NS}/customXml",
            "Target": f"../{source_name}",
        },
    )
    entries[document_rels_name] = ET.tostring(
        document_rels, encoding="utf-8", xml_declaration=True
    )

    content_types_name = "[Content_Types].xml"
    content_types = ET.fromstring(entries[content_types_name])
    ET.SubElement(
        content_types,
        f"{{{CONTENT_TYPES_NS}}}Override",
        {
            "PartName": f"/{props_name}",
            "ContentType": "application/vnd.openxmlformats-officedocument.customXmlProperties+xml",
        },
    )
    entries[content_types_name] = ET.tostring(
        content_types, encoding="utf-8", xml_declaration=True
    )
    return source_name


def embed_word_sources(docx_path: str | Path, sources: list[dict[str, Any]]) -> str | None:
    """Embed sources in the DOCX current list used by Word's Source Manager."""
    if not sources:
        return None
    path = Path(docx_path)
    with ZipFile(path, "r") as archive:
        infos = archive.infolist()
        entries = {info.filename: archive.read(info.filename) for info in infos}

    sources_xml = build_sources_xml(sources)
    sources_part = _find_sources_part(entries)
    if sources_part:
        entries[sources_part] = sources_xml
    else:
        sources_part = _add_sources_part(entries, sources_xml)

    temporary = path.with_name(f".{path.name}.bibliography.tmp")
    try:
        with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as archive:
            written: set[str] = set()
            for info in infos:
                archive.writestr(info, entries[info.filename])
                written.add(info.filename)
            for name, payload in entries.items():
                if name not in written:
                    archive.writestr(name, payload)
        os.replace(temporary, path)
    finally:
        if temporary.exists():
            temporary.unlink()
    return sources_part


def _field_instructions(document_xml: bytes) -> list[str]:
    root = ET.fromstring(document_xml)
    return [
        str(element.text or "").strip()
        for element in root.iter(f"{{{WORDPROCESSING_NS}}}instrText")
        if str(element.text or "").strip()
    ]


def _bibliography_field_result_text(document_xml: bytes) -> str:
    """Extract only the cached/result text of the native BIBLIOGRAPHY field.

    Word rewrites the cached result when the user presses F9: line breaks may
    become paragraphs and punctuation/style runs may change. Reading the field
    result itself lets validation remain strict about source coverage without
    depending on one exact serialization of APA display text.
    """
    root = ET.fromstring(document_xml)
    field_stack: list[dict[str, Any]] = []
    completed_results: list[str] = []
    fld_char_attr = f"{{{WORDPROCESSING_NS}}}fldCharType"

    for element in root.iter():
        if element.tag == f"{{{WORDPROCESSING_NS}}}fldChar":
            field_type = str(element.get(fld_char_attr) or "").strip().lower()
            if field_type == "begin":
                field_stack.append({"instruction": "", "in_result": False, "text": []})
            elif field_type == "separate" and field_stack:
                field_stack[-1]["in_result"] = True
            elif field_type == "end" and field_stack:
                completed = field_stack.pop()
                if re.search(r"\bBIBLIOGRAPHY\b", completed["instruction"], re.IGNORECASE):
                    completed_results.append(" ".join(completed["text"]))
            continue

        if element.tag == f"{{{WORDPROCESSING_NS}}}instrText" and field_stack:
            field_stack[-1]["instruction"] += str(element.text or "")
            continue

        if element.tag == f"{{{WORDPROCESSING_NS}}}t" and element.text:
            for field in field_stack:
                if field["in_result"]:
                    field["text"].append(str(element.text))

    return " ".join(completed_results)


def _normalize_visible_bibliography_text(value: Any) -> str:
    decomposed = unicodedata.normalize("NFKD", str(value or ""))
    ascii_text = "".join(ch for ch in decomposed if not unicodedata.combining(ch))
    return " ".join(re.sub(r"[^A-Za-z0-9]+", " ", ascii_text).casefold().split())


def _source_signature_visible(result_text: str, source: dict[str, Any]) -> bool:
    """Match a Word-updated entry by its unique first-author/year signature."""
    authors = source.get("authors") or []
    first_author = next(
        (
            str(author.get("last") or "").strip()
            for author in authors
            if isinstance(author, dict) and str(author.get("last") or "").strip()
        ),
        "",
    )
    year = str(source.get("year") or "").strip()
    if not first_author or not year:
        return False
    normalized_result = _normalize_visible_bibliography_text(result_text)
    author_token = _normalize_visible_bibliography_text(first_author)
    for match in re.finditer(rf"(?<![a-z0-9]){re.escape(author_token)}(?![a-z0-9])", normalized_result):
        if re.search(rf"(?<!\d){re.escape(year)}(?!\d)", normalized_result[match.end():match.end() + 120]):
            return True
    return False


def validate_native_bibliography_docx(
    docx_path: str | Path,
    expected_sources: list[dict[str, Any]],
) -> dict[str, int]:
    """Verify source-manager entries and their CITATION/BIBLIOGRAPHY fields."""
    citation_signatures: dict[str, str] = {}
    duplicate_signatures: list[str] = []
    for source in expected_sources:
        authors = source.get("authors") or []
        author_key = "|".join(
            str(author.get("last") or "").strip().casefold()
            for author in authors
            if isinstance(author, dict) and str(author.get("last") or "").strip()
        )
        signature = f"{author_key}|{str(source.get('year') or '').strip()}"
        tag = str(source.get("tag") or "")
        previous = citation_signatures.get(signature)
        if previous and previous != tag:
            duplicate_signatures.append(f"{previous}/{tag}")
        else:
            citation_signatures[signature] = tag
    if duplicate_signatures:
        raise NativeBibliographyValidationError(
            "fuentes Word con autor y año duplicados; Word podría mostrar el título en la cita: "
            + ", ".join(sorted(duplicate_signatures))
        )

    path = Path(docx_path)
    with ZipFile(path, "r") as archive:
        entries = {name: archive.read(name) for name in archive.namelist()}
    sources_part = _find_sources_part(entries)
    if not sources_part:
        raise NativeBibliographyValidationError("el DOCX no contiene la lista actual de fuentes de Word")

    root = ET.fromstring(entries[sources_part])
    embedded_tags = {
        str(element.text or "").strip()
        for element in root.iter(f"{{{BIBLIOGRAPHY_NS}}}Tag")
        if str(element.text or "").strip()
    }
    expected_tags = {str(source.get("tag") or "") for source in expected_sources}
    if embedded_tags != expected_tags:
        raise NativeBibliographyValidationError(
            f"fuentes Word incompatibles: esperadas={len(expected_tags)}, encontradas={len(embedded_tags)}"
        )

    instructions = _field_instructions(entries["word/document.xml"])
    citation_fields = [item for item in instructions if re.search(r"\bCITATION\b", item, re.IGNORECASE)]
    bibliography_fields = [
        item for item in instructions if re.search(r"\bBIBLIOGRAPHY\b", item, re.IGNORECASE)
    ]
    citations_with_visible_title = [
        item for item in citation_fields if not re.search(r"\\t(?:\s|$)", item)
    ]
    if citations_with_visible_title:
        raise NativeBibliographyValidationError(
            "campos CITATION sin el modificador que suprime títulos largos"
        )
    cited_tags: set[str] = set()
    for instruction in citation_fields:
        primary = re.search(r"\bCITATION\s+([A-Z0-9_-]+)", instruction, re.IGNORECASE)
        if primary:
            cited_tags.add(primary.group(1))
        cited_tags.update(re.findall(r"\\m\s+([A-Z0-9_-]+)", instruction, re.IGNORECASE))

    if cited_tags != expected_tags:
        missing = sorted(expected_tags - cited_tags)
        unknown = sorted(cited_tags - expected_tags)
        raise NativeBibliographyValidationError(
            "correspondencia CITATION/fuentes inválida"
            + (f"; sin citar={', '.join(missing)}" if missing else "")
            + (f"; desconocidas={', '.join(unknown)}" if unknown else "")
        )
    if len(bibliography_fields) != 1:
        raise NativeBibliographyValidationError(
            f"se esperaba un campo BIBLIOGRAPHY y se encontraron {len(bibliography_fields)}"
        )
    document_xml = entries["word/document.xml"]
    visible_document_text = " ".join(
        str(element.text or "")
        for element in ET.fromstring(document_xml).iter()
        if str(element.text or "")
    )
    pending_text = "Bibliografía nativa pendiente de actualización en Microsoft Word."
    if pending_text in visible_document_text:
        raise NativeBibliographyValidationError(
            "la bibliografía conserva el texto pendiente en lugar de resultados visibles"
        )
    cached_entries = bibliography_display_entries(expected_sources)
    bibliography_result_text = _bibliography_field_result_text(document_xml)
    visible_source_count = sum(
        1 for source in expected_sources
        if _source_signature_visible(bibliography_result_text, source)
    )
    if visible_source_count < len(expected_sources):
        raise NativeBibliographyValidationError(
            "bibliografía visible incompleta; "
            f"entradas={visible_source_count}, fuentes citadas={len(expected_sources)}"
        )
    return {
        "sources": len(embedded_tags),
        "citation_fields": len(citation_fields),
        "bibliography_fields": len(bibliography_fields),
        "bibliography_cached_entries": len(cached_entries),
    }
