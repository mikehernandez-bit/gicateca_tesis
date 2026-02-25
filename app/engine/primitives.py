"""
Archivo: app/engine/primitives.py
Proposito:
- Funciones DOCX puras reutilizables por los renderers del Block Engine.

Responsabilidades:
- Proveer helpers atómicos para manipular documentos python-docx.
- Cada función realiza UNA sola operación sobre el documento.
No hace:
- No lee JSON ni decide qué contenido renderizar (eso es del normalizer).
- No registra renderers (eso es del registry).

Entradas/Salidas:
- Entradas: documento python-docx + parámetros de formato.
- Salidas: documento modificado in-place (sin retorno).

Dependencias:
- python-docx (docx).

Puntos de extension:
- Agregar nuevos helpers atómicos según se necesiten.

Donde tocar si falla:
- Verificar imports de docx y constantes de márgenes/fuentes.

NOTA: Estas funciones son COPIAS de las que existen en universal_generator.py.
      El generador viejo sigue usando sus propias copias hasta la Fase 5.
      Los renderers nuevos (Fase 4) importarán de aquí.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from docx.document import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════
# CONSTANTES
# ═══════════════════════════════════════════════════════════════

DEFAULT_HEADER_COLOR = "D9D9D9"
DEFAULT_TABLE_FONT_SIZE = 10
LANDSCAPE_FONT_SIZE = 9

PORTRAIT_MARGINS = {"left": 3.5, "right": 2.5, "top": 3.0, "bottom": 3.0}
LANDSCAPE_MARGINS = {"left": 2.0, "right": 2.0, "top": 2.0, "bottom": 2.0}

# Raíz del proyecto: engine/ → app/ → gicateca_tesis/
_ENGINE_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _ENGINE_DIR.parent.parent

# Rutas de búsqueda de assets (logos, imágenes de ejemplo)
ASSETS_PATHS = [
    _PROJECT_ROOT / "app" / "static" / "assets",
    _PROJECT_ROOT / "app" / "universities" / "unac" / "assets",
    _PROJECT_ROOT / "app" / "universities" / "uni" / "assets",
]


# ═══════════════════════════════════════════════════════════════
# RESOLUCIÓN DE ASSETS
# ═══════════════════════════════════════════════════════════════

def resolve_asset(filename: str) -> Optional[str]:
    """Busca un archivo de asset en las ubicaciones conocidas.

    Retorna la ruta absoluta como string, o None si no se encuentra.
    """
    if not filename:
        return None
    # Ruta absoluta directa
    if Path(filename).exists():
        return str(Path(filename).resolve())
    # Buscar en rutas de assets
    for p in ASSETS_PATHS:
        candidate = p / filename
        if candidate.exists():
            return str(candidate)
    return None


def resolve_logo_path(data: dict) -> Optional[str]:
    """Resuelve la ruta del logo universitario desde el JSON.

    Estrategia:
    1. Ruta explícita en configuracion.ruta_logo
    2. Convención Logo{CODE}.png desde _meta.university
    3. Fallback genérico LogoGeneric.png
    """
    config = data.get("configuracion", {})
    if "ruta_logo" in config:
        candidate = _PROJECT_ROOT / config["ruta_logo"]
        if candidate.exists():
            return str(candidate)
        name = Path(config["ruta_logo"]).name
        found = resolve_asset(name)
        if found:
            return found

    meta = data.get("_meta", {})
    uni = (meta.get("university") or "").strip().lower()
    if uni:
        logo_name = f"Logo{uni.upper()}.png"
        found = resolve_asset(logo_name)
        if found:
            return found

    return resolve_asset("LogoGeneric.png")


# ═══════════════════════════════════════════════════════════════
# CONFIGURACIÓN DEL DOCUMENTO
# ═══════════════════════════════════════════════════════════════

def configure_styles(doc: Document, font_name: str = "Arial", font_size: int = 11) -> None:
    """Configura el estilo Normal del documento."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Evita páginas en blanco por "page break before" heredado en estilos heading.
    for heading_name in ("Heading 1", "Heading 2"):
        try:
            heading_style = doc.styles[heading_name]
            heading_style.paragraph_format.page_break_before = False
        except KeyError:
            continue


def configure_margins(doc: Document) -> None:
    """Configura márgenes A4 estándar UNAC/UNI para todas las secciones."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(PORTRAIT_MARGINS["left"])
        section.right_margin = Cm(PORTRAIT_MARGINS["right"])
        section.top_margin = Cm(PORTRAIT_MARGINS["top"])
        section.bottom_margin = Cm(PORTRAIT_MARGINS["bottom"])


# ═══════════════════════════════════════════════════════════════
# PÁRRAFOS Y ENCABEZADOS
# ═══════════════════════════════════════════════════════════════

def add_p_centered(
    doc: Document,
    text: str,
    bold: bool = False,
    size: int = 12,
    space_before: int = 0,
    space_after: int = 0,
    italic: bool = False,
) -> None:
    """Agrega un párrafo centrado con formato personalizado."""
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.italic = italic
    run.font.name = "Arial"


def add_heading_formal(
    doc: Document,
    text: str,
    level: int = 1,
    space_before: int = 12,
    space_after: int = 12,
    centered: bool = False,
) -> None:
    """Agrega un encabezado formal con fuente Arial negra."""
    if not text:
        return
    h = doc.add_heading(level=level)
    # Hardening: nunca permitir salto de página implícito en títulos.
    h.paragraph_format.page_break_before = False
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT

    run = h.runs[0] if h.runs else h.add_run(text)
    if run.text != text:
        run.text = text
    run.font.name = "Arial"
    run.font.size = Pt(14 if level == 1 else 12)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_black_heading(
    doc: Document,
    text: str,
    level: int = 2,
    size: int = 13,
    centered: bool = True,
) -> None:
    """Agrega un encabezado con fuente Arial negra, opcionalmente centrado."""
    h = doc.add_heading(text, level=level)
    # Hardening: nunca permitir salto de página implícito en subtítulos.
    h.paragraph_format.page_break_before = False
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Arial"
        run.font.size = Pt(size)
    if centered:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ═══════════════════════════════════════════════════════════════
# NOTAS ESTILIZADAS (BLUE BOX)
# ═══════════════════════════════════════════════════════════════

def add_styled_note(doc: Document, text: str) -> None:
    """Agrega una nota en caja azul (instrucciones/validación)."""
    if not text:
        return
    try:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Cm(15.0)
        cell = table.cell(0, 0)

        # Fondo F2F8FD
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F2F8FD")
        tc_pr.append(shd)

        # Bordes 8DB3E2
        tblBorders = OxmlElement("w:tblBorders")
        for b in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{b}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "8DB3E2")
            tblBorders.append(border)
        table._tbl.tblPr.append(tblBorders)

        # Contenido
        cell.paragraphs[0].text = ""
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)

            if line.startswith(("- ", "• ", "* ")):
                p.paragraph_format.left_indent = Cm(0.75)
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.add_run("• ").font.size = Pt(10)
                p.add_run(line[2:].strip()).font.size = Pt(10)
            elif ":" in line and len(line.split(":", 1)[0]) < 60:
                parts = line.split(":", 1)
                r = p.add_run(parts[0] + ":")
                r.bold = True
                r.font.size = Pt(10)
                p.add_run(parts[1]).font.size = Pt(10)
            else:
                p.add_run(line).font.size = Pt(10)

        doc.add_paragraph()  # Spacer
    except Exception:
        logger.warning("Error rendering styled note: %s", text[:80])
        doc.add_paragraph(f"[Nota: {text}]")


# ═══════════════════════════════════════════════════════════════
# CAMPOS WORD (TOC, SEQ, PAGE)
# ═══════════════════════════════════════════════════════════════

def add_fld_page(paragraph) -> None:
    """Agrega un campo PAGE para numeración de páginas."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc_field(
    doc: Document,
    field_code: str,
    heading_text: str,
    exclude_from_toc: bool = False,
) -> None:
    """Inserta un campo Word real (TOC, lista de tablas/figuras).

    Agrega encabezado + campo + placeholder + salto de página.
    Si exclude_from_toc=True, el título usa estilo Normal (no aparece en TOC).
    Si exclude_from_toc=False, usa Heading 1 (aparece en TOC).
    """
    if exclude_from_toc:
        # Párrafo Normal — no aparece en el TOC
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(12)
        run_title = p_title.add_run(heading_text)
        run_title.font.name = "Arial"
        run_title.font.size = Pt(14)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(0, 0, 0)
    else:
        # Heading 1 — aparece en el TOC
        add_heading_formal(doc, heading_text, centered=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)

    placeholder_run = paragraph.add_run(
        "Actualice este campo: clic derecho → Actualizar campo"
    )
    placeholder_run.font.size = Pt(9)
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    placeholder_run.italic = True

    end_run = paragraph.add_run()
    end_run._r.append(fld_end)

    doc.add_page_break()


def add_seq_field(paragraph, seq_name: str) -> None:
    """Inserta un campo SEQ para auto-numeración (tablas, figuras).

    Permite que Word construya índices de tablas/figuras automáticamente.
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    # Placeholder number
    num_run = paragraph.add_run("#")
    num_run.font.size = Pt(10)
    num_run.italic = True
    end_run = paragraph.add_run()
    end_run._r.append(fld_end)


# ═══════════════════════════════════════════════════════════════
# TABLA: HELPERS DE CELDA
# ═══════════════════════════════════════════════════════════════

def apply_cell_shading(cell, color_hex: str) -> None:
    """Aplica sombreado de fondo a una celda de tabla."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_vertical_alignment(
    cell, alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
) -> None:
    """Alineación vertical de una celda (por defecto centrada)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    val_align = OxmlElement("w:vAlign")
    _ALIGN_MAP = {
        WD_CELL_VERTICAL_ALIGNMENT.TOP: "top",
        WD_CELL_VERTICAL_ALIGNMENT.CENTER: "center",
        WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: "bottom",
    }
    val_align.set(qn("w:val"), _ALIGN_MAP.get(alignment, "center"))
    tc_pr.append(val_align)


def format_cell_text(
    cell,
    text: str,
    font_size: int,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Escribe texto formateado en una celda, soportando multi-línea."""
    cell.text = ""
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.size = Pt(font_size)
        run.font.name = "Arial"
        run.bold = bold


# ═══════════════════════════════════════════════════════════════
# SECCIONES: ORIENTACIÓN LANDSCAPE / PORTRAIT
# ═══════════════════════════════════════════════════════════════

def _shrink_section_break_paragraph(doc: Document) -> None:
    """Minimiza el párrafo que contiene el section-break para evitar páginas en blanco.

    Cuando ``doc.add_section(NEW_PAGE)`` crea un quiebre de sección, python-docx
    inserta un ``<w:p>`` vacío cuyo ``<w:pPr>`` contiene el ``<w:sectPr>``.
    Si la tabla landscape llenó la página, ese párrafo vacío se desborda a la
    siguiente página generando una hoja en blanco.

    Solución: recorrer hacia atrás los ``<w:p>`` del body y, al encontrar el
    que tiene ``<w:sectPr>`` en su ``<w:pPr>``, forzar fuente a 2 half-points
    (1 pt) y spacing a 0 para que no ocupe espacio visible.
    """
    body = doc.element.body
    for p_elem in reversed(list(body.iterchildren(qn("w:p")))):
        p_pr = p_elem.find(qn("w:pPr"))
        if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
            # --- spacing 0 ---
            spacing = p_pr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                p_pr.insert(0, spacing)
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:line"), "240")
            spacing.set(qn("w:lineRule"), "auto")

            # --- rPr con tamaño mínimo para cualquier run implícito ---
            r_pr = p_pr.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                p_pr.append(r_pr)
            for tag in ("w:sz", "w:szCs"):
                el = r_pr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    r_pr.append(el)
                el.set(qn("w:val"), "2")  # 2 half-points = 1 pt
            break


def switch_to_landscape(doc: Document) -> None:
    """Agrega nueva sección en orientación landscape."""
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.left_margin = Cm(LANDSCAPE_MARGINS["left"])
    new_section.right_margin = Cm(LANDSCAPE_MARGINS["right"])
    new_section.top_margin = Cm(LANDSCAPE_MARGINS["top"])
    new_section.bottom_margin = Cm(LANDSCAPE_MARGINS["bottom"])
    _shrink_section_break_paragraph(doc)


def switch_to_portrait(doc: Document) -> None:
    """Agrega nueva sección para restaurar orientación portrait."""
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.left_margin = Cm(PORTRAIT_MARGINS["left"])
    new_section.right_margin = Cm(PORTRAIT_MARGINS["right"])
    new_section.top_margin = Cm(PORTRAIT_MARGINS["top"])
    new_section.bottom_margin = Cm(PORTRAIT_MARGINS["bottom"])
    _shrink_section_break_paragraph(doc)
