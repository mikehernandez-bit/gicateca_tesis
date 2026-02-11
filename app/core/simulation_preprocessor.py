"""
Simulation preprocessor for render/docx and render/pdf simulation mode.

Responsibilities:
1. Sanitize definition (remove notes/guides keys recursively)
2. Compile section index for stable path/sectionId mapping
3. Inject simulation AI content under real headings in rendered DOCX
4. Remove fixed template guide paragraphs in simulation mode only
"""
from __future__ import annotations

import copy
import json
import tempfile
import unicodedata
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


# Keys that must be removed in simulation mode (case-insensitive).
EXCLUDED_KEYS = frozenset(
    {
        "nota",
        "notas",
        "guia",
        "guias",
        "ejemplo",
        "ejemplos",
        "instruccion",
        "instrucciones",
        "instruccion_detallada",
        "comentario",
        "comentarios",
        "observacion",
        "observaciones",
        # Extra legacy guidance keys commonly used in current definitions.
        "nota_capitulo",
        "nota_general",
    }
)

SECTION_TITLE_KEYS = ("titulo", "title", "titulo_seccion", "texto")
SECTION_CONTAINER_KEYS = frozenset(
    {
        "preliminares",
        "cuerpo",
        "finales",
        "capitulos",
        "contenido",
        "items",
        "secciones",
        "subsecciones",
        "lista",
        "anexos",
        "indices",
    }
)

# Required fallback text when a section has no aiResult content.
DEFAULT_AI_PLACEHOLDER = "Contenido IA simulado"

# Fixed template guide tokens to purge only in simulation outputs.
GUIDE_TOKENS = (
    "NOTA:",
    "GUIA:",
    "EJEMPLO:",
    "INSTRUCCION:",
    "[ESCRIBA",
    "[COLOQUE",
)


def _normalize_text(value: Any) -> str:
    if not isinstance(value, str):
        return ""
    return " ".join(value.strip().split())


def _normalize_for_match(value: Any) -> str:
    text = _normalize_text(value)
    if not text:
        return ""
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return text.lower()


def _should_exclude_key(key: str) -> bool:
    return key.strip().lower() in EXCLUDED_KEYS


def sanitize_definition(obj: Any) -> Any:
    """Recursively remove notes/guides keys from a format definition."""
    if isinstance(obj, dict):
        sanitized: Dict[str, Any] = {}
        for key, value in obj.items():
            if _should_exclude_key(key):
                continue
            sanitized[key] = sanitize_definition(value)
        return sanitized
    if isinstance(obj, list):
        return [sanitize_definition(item) for item in obj]
    return obj


def _extract_section_title(node: Dict[str, Any]) -> str:
    for key in SECTION_TITLE_KEYS:
        title = _normalize_text(node.get(key))
        if title:
            return title
    return ""


def _build_section_index_recursive(
    obj: Any,
    out: List[Dict[str, Any]],
    path_stack: List[str],
    level: int,
    in_structure: bool,
) -> None:
    if isinstance(obj, list):
        for item in obj:
            _build_section_index_recursive(
                item,
                out=out,
                path_stack=path_stack,
                level=level,
                in_structure=in_structure,
            )
        return

    if not isinstance(obj, dict):
        return

    title = _extract_section_title(obj) if in_structure else ""
    next_stack = path_stack
    next_level = level

    if title:
        next_stack = path_stack + [title]
        out.append(
            {
                "sectionId": f"sec-{len(out) + 1:04d}",
                "path": "/".join(next_stack),
                "level": max(1, min(level, 6)),
                "kind": "heading",
                "title": title,
            }
        )
        next_level = min(level + 1, 6)

    for key, value in obj.items():
        key_lower = key.lower()
        if _should_exclude_key(key) or key_lower in SECTION_TITLE_KEYS:
            continue
        if not isinstance(value, (dict, list)):
            continue

        child_in_structure = in_structure or key_lower in SECTION_CONTAINER_KEYS
        _build_section_index_recursive(
            value,
            out=out,
            path_stack=next_stack,
            level=next_level if child_in_structure else level,
            in_structure=child_in_structure,
        )


def build_section_index(definition: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Compile ordered sectionIndex entries from a sanitized definition."""
    out: List[Dict[str, Any]] = []
    _build_section_index_recursive(
        definition,
        out=out,
        path_stack=[],
        level=1,
        in_structure=False,
    )
    return out


def _build_ai_maps(ai_sections: Iterable[Dict[str, Any]]) -> Tuple[Dict[str, str], Dict[str, str], Dict[str, str]]:
    by_section_id: Dict[str, str] = {}
    by_path: Dict[str, str] = {}
    by_path_normalized: Dict[str, str] = {}

    for section in ai_sections:
        if not isinstance(section, dict):
            continue
        content = _normalize_text(section.get("content"))
        if not content:
            continue

        section_id = _normalize_text(section.get("sectionId"))
        if section_id:
            by_section_id[section_id.lower()] = content

        path = _normalize_text(section.get("path"))
        if path:
            by_path[path] = content
            by_path_normalized[_normalize_for_match(path)] = content

    return by_section_id, by_path, by_path_normalized


def _resolve_ai_content(
    section: Dict[str, Any],
    by_section_id: Dict[str, str],
    by_path: Dict[str, str],
    by_path_normalized: Dict[str, str],
) -> str:
    section_id = _normalize_text(section.get("sectionId")).lower()
    if section_id and section_id in by_section_id:
        return by_section_id[section_id]

    path = _normalize_text(section.get("path"))
    if path and path in by_path:
        return by_path[path]

    path_normalized = _normalize_for_match(path)
    if path_normalized and path_normalized in by_path_normalized:
        return by_path_normalized[path_normalized]

    return DEFAULT_AI_PLACEHOLDER


def _looks_like_heading(paragraph) -> bool:
    text = _normalize_text(paragraph.text)
    if not text:
        return False

    style_name = ""
    try:
        style_name = paragraph.style.name or ""
    except Exception:
        style_name = ""
    style_norm = _normalize_for_match(style_name)
    if "heading" in style_norm or "titulo" in style_norm:
        return True

    # Some generators render subtitles as bold paragraphs instead of heading style.
    has_non_empty_runs = False
    for run in paragraph.runs:
        run_text = _normalize_text(run.text)
        if not run_text:
            continue
        has_non_empty_runs = True
        if run.bold is not True:
            return False
    return has_non_empty_runs


def _find_section_paragraphs(doc: Document, section_index: List[Dict[str, Any]]) -> List[Tuple[Dict[str, Any], Any]]:
    paragraphs = [p for p in doc.paragraphs if _normalize_text(p.text)]
    matches: List[Tuple[Dict[str, Any], Any]] = []
    cursor = 0

    for section in section_index:
        expected_title = _normalize_for_match(section.get("title"))
        found_index = -1

        if expected_title:
            for idx in range(cursor, len(paragraphs)):
                actual = _normalize_for_match(paragraphs[idx].text)
                if not actual:
                    continue
                if expected_title in actual or actual in expected_title:
                    found_index = idx
                    break

        if found_index < 0:
            for idx in range(cursor, len(paragraphs)):
                if _looks_like_heading(paragraphs[idx]):
                    found_index = idx
                    break

        if found_index < 0:
            continue

        matches.append((section, paragraphs[found_index]))
        cursor = found_index + 1

    return matches


def _insert_paragraph_after(paragraph, text: str) -> Any:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    try:
        new_paragraph.style = "Normal"
    except Exception:
        pass
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _is_guide_paragraph(text: str) -> bool:
    normalized = _normalize_for_match(text)
    if not normalized:
        return False
    for token in GUIDE_TOKENS:
        normalized_token = _normalize_for_match(token)
        if normalized_token and normalized_token in normalized:
            return True
    return False


def postprocess_simulation_docx(
    docx_path: Path,
    section_index: List[Dict[str, Any]],
    ai_sections: Iterable[Dict[str, Any]],
) -> Dict[str, int]:
    """
    Post-process simulation DOCX before returning it:
    - Insert AI content below each heading found from sectionIndex
    - Remove fixed guide paragraphs (NOTA/Guia/Ejemplo/Instruccion placeholders)
    """
    doc = Document(str(docx_path))

    by_section_id, by_path, by_path_normalized = _build_ai_maps(ai_sections)
    section_paragraphs = _find_section_paragraphs(doc, section_index)

    inserted = 0
    for section, paragraph in section_paragraphs:
        content = _resolve_ai_content(section, by_section_id, by_path, by_path_normalized)
        lines = [line for line in content.splitlines() if _normalize_text(line)]
        if not lines:
            lines = [DEFAULT_AI_PLACEHOLDER]

        cursor = paragraph
        for line in lines:
            cursor = _insert_paragraph_after(cursor, _normalize_text(line))
            inserted += 1

    removed_guides = 0
    for paragraph in list(_iter_all_paragraphs(doc)):
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        if _is_guide_paragraph(text):
            _delete_paragraph(paragraph)
            removed_guides += 1

    doc.save(str(docx_path))
    return {
        "inserted": inserted,
        "removed_guides": removed_guides,
    }


def prepare_simulation_json(
    format_data: Dict[str, Any],
    values: Optional[Dict[str, Any]] = None,
    ai_result: Optional[Dict[str, Any]] = None,
) -> Path:
    """
    Build a sanitized temporary JSON for simulation mode generator execution.

    ai_result is accepted for backwards compatibility but intentionally ignored here,
    because AI content is injected post-render in DOCX/PDF simulation flow.
    """
    _ = ai_result
    data = copy.deepcopy(format_data)
    sanitized = sanitize_definition(data)

    if values and isinstance(sanitized.get("caratula"), dict):
        for key, val in values.items():
            sanitized["caratula"][key] = val

    tmp_file = tempfile.NamedTemporaryFile(
        mode="w",
        suffix=".json",
        prefix="sim_",
        delete=False,
        encoding="utf-8",
    )
    json.dump(sanitized, tmp_file, ensure_ascii=False, indent=2)
    tmp_file.close()
    return Path(tmp_file.name)


def cleanup_temp_json(path: Path) -> None:
    """Delete temporary JSON file created for simulation rendering."""
    try:
        path.unlink()
    except Exception:
        pass
