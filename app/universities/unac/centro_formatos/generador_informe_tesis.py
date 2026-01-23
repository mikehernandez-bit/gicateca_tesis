"""
Archivo: app/universities/unac/centro_formatos/generador_informe_tesis.py
Proposito:
- Genera documentos DOCX para el formato "informe" de UNAC.

Responsabilidades:
- Cargar JSON legacy y construir el documento Word con estilos oficiales.
- Aplicar estructura (caratula, preliminares, cuerpo, finales).
No hace:
- No descubre formatos ni valida reglas del catalogo.

Entradas/Salidas:
- Entradas: ruta JSON y ruta de salida (via CLI o funciones internas).
- Salidas: archivo DOCX generado en la ruta indicada.

Dependencias:
- python-docx, json, os, sys.

Puntos de extension:
- Ajustar estilos/estructura si cambia el formato institucional.

Donde tocar si falla:
- Revisar cargar_contenido y generar_documento_core.
"""

import json
import os
import sys
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ==========================================
# CONFIGURACIÓN
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(BASE_DIR, "..", "..", "..", ".."))
FORMATS_DIR = os.path.join(BASE_DIR, "formats")
STATIC_ASSETS_DIR = os.path.join(PROJECT_ROOT, "app", "static", "assets")

def cargar_contenido(path_archivo):
    """Carga el JSON base desde la ruta indicada o fallback."""
    if not os.path.exists(path_archivo):
        nombre = os.path.basename(path_archivo)
        path_archivo = os.path.join(FORMATS_DIR, nombre)
    if not os.path.exists(path_archivo):
        raise FileNotFoundError(f"No se encontró el JSON en: {path_archivo}")
    try:
        with open(path_archivo, 'r', encoding='utf-8') as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        raise ValueError(f"Error de sintaxis en el archivo JSON: {e}")

# ==========================================
# UTILIDADES DE FORMATO Y XML
# ==========================================

def configurar_seccion_unac(section):
    """Establece tamaño A4 y márgenes oficiales UNAC."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)

def configurar_estilos_base(doc):
    """Configura estilos base del documento (fuente e interlineado)."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'; font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def _add_fldSimple(paragraph, instr: str):
    """Agrega campos dinámicos (TOC/PAGE) de Word."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    paragraph._p.append(fld)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "
    r.append(t)
    fld.append(r)

def agregar_bloque(doc, texto, negrita=False, tamano=12, antes=0, despues=0):
    """Agrega un bloque centrado con estilo basico."""
    if not texto: return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(despues)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(texto)
    run.bold = negrita; run.font.size = Pt(tamano)
    return p

def agregar_titulo_formal(doc, texto, espaciado_antes=0):
    """Agrega un titulo formal nivel 1."""
    if not texto: return
    h = doc.add_heading(texto, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    h.paragraph_format.space_before = Pt(espaciado_antes); h.paragraph_format.space_after = Pt(12)
    run = h.runs[0]
    run.font.name = 'Arial'; run.font.size = Pt(14); run.bold = True; run.font.color.rgb = RGBColor(0, 0, 0)

def agregar_nota_estilizada(doc, texto):
    """Renderiza una nota destacada en una tabla sombreada."""
    if not texto: return
    try:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False; table.columns[0].width = Cm(15.0)
        cell = table.cell(0, 0)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F2F8FD') 
        tc_pr.append(shd)
        tblBorders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:color'), '8DB3E2')
            tblBorders.append(border)
        table._tbl.tblPr.append(tblBorders)
        for p in cell.paragraphs: p._element.getparent().remove(p._element)
        lineas = texto.split('\n')
        for linea in lineas:
            linea = linea.strip()
            if not linea: continue 
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Pt(6); p.paragraph_format.space_after = Pt(6)
            if linea.startswith("- ") or linea.startswith("• ") or linea.startswith("* "):
                p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.5)
                p.add_run("•\t").font.size = Pt(10)
                p.add_run(linea[2:].strip()).font.size = Pt(10)
            elif ":" in linea and len(linea.split(":", 1)[0]) < 65:
                parts = linea.split(":", 1)
                r_t = p.add_run(parts[0] + ":"); r_t.bold = True; r_t.font.size = Pt(10)
                if len(parts) > 1: p.add_run(parts[1]).font.size = Pt(10)
            else:
                p.add_run(linea).font.size = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
    except: pass

# --- REFERENCIAS APA 7 ---
def imprimir_ejemplos_apa(doc, lista_ejemplos):
    """Imprime referencias APA con sangria francesa."""
    if not lista_ejemplos: return
    p_head = doc.add_paragraph()
    p_head.add_run("Referencias Bibliográficas (Formato APA 7ma Ed.):").bold = True
    for ej in lista_ejemplos:
        p = doc.add_paragraph(ej)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Sangría Francesa (Hanging Indent) de 1.27 cm
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(12)
        for r in p.runs:
            r.font.name = 'Arial'; r.font.size = Pt(11)

# ==========================================
# UTILIDADES DE MATRIZ
# ==========================================

def add_compact_p(cell, text, bold=False, color=None, bullet=False):
    """Agrega un parrafo compacto dentro de una celda."""
    p = cell.add_paragraph(); p.paragraph_format.line_spacing = 1.0; p.paragraph_format.space_after = Pt(2)
    if bullet: p.style = 'List Bullet'; p.paragraph_format.left_indent = Pt(10)
    run = p.add_run(str(text)); run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = bold
    if color: run.font.color.rgb = color
    return p

def crear_tabla_matriz_consistencia(doc, matriz_data, es_cualitativo=False):
    """Construye la tabla de matriz de consistencia."""
    doc.add_paragraph()
    p_t = doc.add_paragraph(); p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo = "Anexo 1: Matriz de Categorización" if es_cualitativo else "Anexo 1: Matriz de Consistencia"
    r_t = p_t.add_run(titulo); r_t.bold = True; r_t.font.size = Pt(12)
    table = doc.add_table(rows=3, cols=5); table.autofit = False
    for row in table.rows:
        for cell in row.cells: cell.width = Cm(3.0)
    headers = ["PROBLEMAS", "OBJETIVOS", "SUPUESTOS" if es_cualitativo else "HIPÓTESIS", "CATEGORÍAS" if es_cualitativo else "VARIABLES", "METODOLOGÍA"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]; tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "D9D9D9"); tcPr.append(shd)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = True; r.font.size = Pt(8)
    prob = matriz_data.get('problemas', {}); obj = matriz_data.get('objetivos', {}); hip = matriz_data.get('hipotesis', {}); vars_d = matriz_data.get('variables', {}); met = matriz_data.get('metodologia', {})
    for i, d in enumerate([prob, obj, hip]):
        add_compact_p(table.rows[1].cells[i], "General:", bold=True); add_compact_p(table.rows[1].cells[i], d.get('general', ''))
        add_compact_p(table.rows[2].cells[i], "Específicos:", bold=True)
        for item in d.get('especificos', []): add_compact_p(table.rows[2].cells[i], item, bullet=True)
    c_v = table.rows[1].cells[3].merge(table.rows[2].cells[3])
    vi = vars_d.get('independiente', {}); add_compact_p(c_v, vi.get('nombre', ''), bold=True, color=RGBColor(255, 0, 0))
    for d in vi.get('dimensiones', []): add_compact_p(c_v, d, bullet=True)
    add_compact_p(c_v, ""); vd = vars_d.get('dependiente', {}); add_compact_p(c_v, vd.get('nombre', ''), bold=True, color=RGBColor(0, 0, 255))
    for d in vd.get('dimensiones', []): add_compact_p(c_v, d, bullet=True)
    c_m = table.rows[1].cells[4].merge(table.rows[2].cells[4])
    for k, v in met.items():
        p = c_m.add_paragraph(); r_k = p.add_run(f"{k.capitalize()}: "); r_k.font.size = Pt(8); r_k.bold = True
        p.add_run(str(v)).font.size = Pt(8)

# ==========================================
# GENERACIÓN DE SECCIONES
# ==========================================

def crear_caratula_dinamica(doc, data):
    """Carátula distribuida en toda la hoja."""
    c = data.get('caratula', {})
    agregar_bloque(doc, c.get('universidad', ''), negrita=True, tamano=16, despues=4)
    agregar_bloque(doc, c.get('facultad', ''), negrita=True, tamano=13, despues=4)
    agregar_bloque(doc, c.get('escuela', ''), negrita=True, tamano=13, despues=25)
    
    r_l = os.path.join(STATIC_ASSETS_DIR, "LogoUNAC.png")
    if os.path.exists(r_l):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(r_l, width=Inches(2.4))
    
    agregar_bloque(doc, c.get('tipo_documento', ''), negrita=True, tamano=15, antes=40)
    agregar_bloque(doc, c.get('titulo_placeholder', ''), negrita=True, tamano=14, antes=30, despues=40)
    agregar_bloque(doc, "PARA OPTAR EL TITULO PROFESIONAL DE:", tamano=11, antes=20)
    agregar_bloque(doc, c.get('grado_objetivo', ''), negrita=True, tamano=12, despues=50)
    agregar_bloque(doc, c.get('label_autor', ''), negrita=True, tamano=11, despues=5)
    agregar_bloque(doc, c.get('label_asesor', ''), negrita=True, tamano=11, despues=30)
    agregar_bloque(doc, c.get('label_linea', ''), tamano=10, despues=60)
    
    # Empujar pie al final
    
    agregar_bloque(doc, c.get('fecha', ''), tamano=11)
    agregar_bloque(doc, c.get('pais', ''), negrita=True, tamano=11)

def agregar_indice_automatico(doc):
    """Inserta el indice de contenido automatico."""
    agregar_bloque(doc, "ÍNDICE DE CONTENIDO", negrita=True, tamano=14, despues=12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.add_run("Pág.")
    p2 = doc.add_paragraph()
    _add_fldSimple(p2, 'TOC \\o "1-3" \\h \\z \\u')

def agregar_preliminares_dinamico(doc, data):
    """Construye las secciones preliminares desde JSON."""
    p = data.get('preliminares', {})
    for sec in ['dedicatoria', 'resumen']:
        if sec in p:
            agregar_bloque(doc, p[sec].get('titulo', ''), negrita=True, tamano=14, despues=12)
            doc.add_paragraph(p[sec].get('texto', ''))
            doc.add_page_break()
    if 'indices' in p:
        agregar_indice_automatico(doc); doc.add_page_break()
    if 'introduccion' in p:
        agregar_titulo_formal(doc, p['introduccion'].get('titulo', ''))
        doc.add_paragraph(p['introduccion'].get('texto', ''))

def agregar_cuerpo_dinamico(doc, data):
    """Construye el cuerpo principal y capitulos."""
    for cap in data.get('cuerpo', []):
        agregar_titulo_formal(doc, cap.get('titulo', ''), espaciado_antes=10)
        if 'nota_capitulo' in cap: agregar_nota_estilizada(doc, cap['nota_capitulo'])
        # Ejemplos APA dentro de capítulos si existen
        if 'ejemplos_apa' in cap: imprimir_ejemplos_apa(doc, cap['ejemplos_apa'])
        for item in cap.get('contenido', []):
            sub = doc.add_heading(item.get('texto', ''), level=2)
            sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sub.runs[0].font.name = 'Arial'; sub.runs[0].font.size = Pt(12); sub.runs[0].font.color.rgb = RGBColor(0,0,0)
            if 'instruccion_detallada' in item: agregar_nota_estilizada(doc, item['instruccion_detallada'])
        doc.add_page_break()

def agregar_finales_dinamico(doc, data):
    """Construye las secciones finales y anexos."""
    fin = data.get('finales', {})
    if 'referencias' in fin:
        agregar_titulo_formal(doc, fin['referencias'].get('titulo', 'REFERENCIAS BIBLIOGRÁFICAS'))
        if 'ejemplos_apa' in fin['referencias']:
            imprimir_ejemplos_apa(doc, fin['referencias']['ejemplos_apa'])
        doc.add_page_break()
    if 'anexos' in fin:
        agregar_titulo_formal(doc, fin['anexos'].get('titulo_seccion', 'ANEXOS'))
        if 'matriz_consistencia' in data:
            es_cuali = "CUALITATIVO" in data.get('caratula', {}).get('tipo_documento', '').upper()
            crear_tabla_matriz_consistencia(doc, data['matriz_consistencia'], es_cuali)
            doc.add_page_break()
        for anexo in fin['anexos'].get('lista', []):
            if "Matriz" in anexo.get('titulo', ''): continue
            p = doc.add_paragraph(anexo.get('titulo', '')); p.runs[0].bold = True; doc.add_page_break()

# ==========================================
# NUMERACIÓN Y EJECUCIÓN
# ==========================================

def configurar_numeracion_tesis(doc):
    """Configura numeracion romana y arabiga."""
    doc.sections[0].different_first_page_header_footer = True
    # Carátula y Hoja Respeto no llevan número visible
    for i in range(min(len(doc.sections), 2)):
        doc.sections[i].footer.is_linked_to_previous = False
    # Romanos desde la sección 3 (Dedicatoria)
    if len(doc.sections) > 2:
        s = doc.sections[2]; s.footer.is_linked_to_previous = False
        pg = OxmlElement('w:pgNumType'); pg.set(qn('w:fmt'), 'lowerRoman'); pg.set(qn('w:start'), '1')
        s._sectPr.append(pg); _insertar_n(s.footer)
    # Arábigos desde el Cuerpo
    if len(doc.sections) > 3:
        s = doc.sections[3]; s.footer.is_linked_to_previous = False
        pg = OxmlElement('w:pgNumType'); pg.set(qn('w:fmt'), 'decimal'); pg.set(qn('w:start'), '1')
        s._sectPr.append(pg); _insertar_n(s.footer)

def _insertar_n(footer):
    """Inserta el campo PAGE en el footer."""
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = "PAGE"
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.extend([f1, it, f2])

def generar_documento_core(ruta_json, ruta_salida):
    """Orquesta la generacion completa del DOCX."""
    data = cargar_contenido(ruta_json); doc = Document(); configurar_estilos_base(doc)
    
    # 1. CARÁTULA
    configurar_seccion_unac(doc.sections[0]); crear_caratula_dinamica(doc, data)
    # 2. HOJA DE RESPETO
    doc.add_section(WD_SECTION.NEW_PAGE); configurar_seccion_unac(doc.sections[-1]); doc.add_paragraph("")
    # 3. PRELIMINARES
    doc.add_section(WD_SECTION.NEW_PAGE); configurar_seccion_unac(doc.sections[-1]); agregar_preliminares_dinamico(doc, data)
    # 4. CUERPO
    doc.add_section(WD_SECTION.NEW_PAGE); configurar_seccion_unac(doc.sections[-1]); agregar_cuerpo_dinamico(doc, data)
    # 5. FINALES
    agregar_finales_dinamico(doc, data)
    # 6. CONFIGURACIÓN FINAL
    configurar_numeracion_tesis(doc)
    doc.settings.element.append(OxmlElement('w:updateFields'))
    doc.save(ruta_salida); return ruta_salida

if __name__ == "__main__":
    if len(sys.argv) > 2: generar_documento_core(sys.argv[1], sys.argv[2])
