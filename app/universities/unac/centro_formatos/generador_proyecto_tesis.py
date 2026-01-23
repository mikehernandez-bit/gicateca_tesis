"""
Archivo: app/universities/unac/centro_formatos/generador_proyecto_tesis.py
Proposito:
- Genera documentos DOCX para el formato "proyecto" de UNAC.

Responsabilidades:
- Cargar JSON legacy y construir el documento Word con estilos oficiales.
- Renderizar caratula, preliminares, cuerpo y finales.
No hace:
- No descubre formatos ni valida reglas del catalogo.

Entradas/Salidas:
- Entradas: ruta JSON y ruta de salida (via CLI o funciones internas).
- Salidas: archivo DOCX generado en la ruta indicada.

Dependencias:
- python-docx, json, os, sys.

Puntos de extension:
- Ajustar estilos/estructura si cambia el formato institucional.

Donde tocar si falla:
- Revisar cargar_contenido y generar_documento_core.
"""

import json
import os
import sys
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

# ==========================================
# 1. CONFIGURACIÓN Y RUTAS
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(BASE_DIR, "..", "..", "..", ".."))
FORMATS_DIR = os.path.join(BASE_DIR, "formats")
STATIC_ASSETS_DIR = os.path.join(PROJECT_ROOT, "app", "static", "assets")

def encontrar_recurso(nombre_archivo):
    """Busca imágenes o logos en varias carpetas comunes."""
    rutas_a_probar = [
        nombre_archivo, # En la misma carpeta
        os.path.join(BASE_DIR, nombre_archivo),
        os.path.join(BASE_DIR, "assets", nombre_archivo),
        os.path.join(BASE_DIR, "static", "assets", nombre_archivo),
        os.path.join(STATIC_ASSETS_DIR, nombre_archivo),
        # Subir un nivel
        os.path.join(BASE_DIR, "..", nombre_archivo),
        os.path.join(BASE_DIR, "..", "assets", nombre_archivo),
        # Estructura típica de proyectos web
        os.path.join(PROJECT_ROOT, "app", "static", "assets", nombre_archivo)
    ]
    
    for ruta in rutas_a_probar:
        if os.path.exists(ruta):
            print(f"[DEBUG] Archivo encontrado: {ruta}")
            return ruta
    return None

def cargar_contenido(path_archivo):
    """Carga el JSON base desde la ruta indicada."""
    if not os.path.exists(path_archivo):
        path_archivo = os.path.join(BASE_DIR, path_archivo)
    if not os.path.exists(path_archivo):
        raise FileNotFoundError(f"No se encontró el JSON en: {path_archivo}")

    with open(path_archivo, 'r', encoding='utf-8') as f:
        return json.load(f)

# ==========================================
# 2. FORMATO Y ESTILOS
# ==========================================

def configurar_formato_unac(doc):
    """Configura el formato de pagina y estilos base."""
    for section in doc.sections:
        section.page_width = Cm(21.0); section.page_height = Cm(29.7)
        section.left_margin = Cm(3.5); section.right_margin = Cm(2.5)
        section.top_margin = Cm(3.0); section.bottom_margin = Cm(3.0)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # Opción B: WD_ALIGN_PARAGRAPH.LEFT (Todo pegado a la izquierda, borde irregular a la derecha - Lo estándar al quitar justificado)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def agregar_bloque(doc, texto, negrita=False, tamano=12, antes=0, despues=0, cursiva=False):
    """Agrega un bloque centrado con estilo basico."""
    if not texto: return
    p = doc.add_paragraph()
    
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE 
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(despues)
    run = p.add_run(texto)
    run.bold = negrita; run.italic = cursiva; run.font.size = Pt(tamano)

def agregar_titulo_formal(doc, texto, espaciado_antes=0):
    """Agrega un titulo formal con estilo institucional."""
    if not texto: return
    h = doc.add_heading(level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(texto)
    run.font.name = 'Arial'; run.font.size = Pt(14); run.bold = True; run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(espaciado_antes)
    h.paragraph_format.space_after = Pt(12)

def agregar_nota_guia(doc, texto):
    """Agrega una nota guia en cursiva."""
    if not texto: return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Nota: {texto}")
    run.font.name = 'Arial'; run.font.size = Pt(10); run.italic = True; run.font.color.rgb = RGBColor(89, 89, 89) 
    p.paragraph_format.space_after = Pt(12)

def agregar_nota_estilizada(doc, texto):
    """Crea un cuadro azul claro solo con el texto de la nota."""
    if not texto: return
    try:
        # Crear tabla de 1 celda para el fondo
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False 
        table.columns[0].width = Cm(15.0) # Ancho fijo para que no se deforme
        
        cell = table.cell(0, 0)
        
        # Color de fondo (Azul muy suave F2F8FD)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F8FD')
        tc_pr.append(shd)
        
        # Bordes (Azul intermedio 8DB3E2)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '8DB3E2')
            tblBorders.append(border)
        table._tbl.tblPr.append(tblBorders)

        # Contenido (Sin prefijos, solo el texto)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(texto)
        run.font.name = 'Arial'
        run.font.size = Pt(9) # Letra un poco más pequeña para diferenciar
        
        # Espacio después del cuadro
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
    except Exception as e:
        # Fallback por si algo falla en el XML
        print(f"[Warn] No se pudo estilizar nota: {e}")
        p = doc.add_paragraph(texto)
        p.runs[0].font.size = Pt(9)
        p.runs[0].italic = True

def agregar_tabla_simple(doc, data_tabla):
    """Agrega una tabla simple tipo grid."""
    if not data_tabla: return
    headers = data_tabla.get('headers', [])
    rows = data_tabla.get('rows', [])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        
    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            if i < len(row_cells): row_cells[i].text = str(val)
    doc.add_paragraph()

def crear_tabla_estilizada(doc, data_tabla):
    """Crea una tabla con estilos y anchos fijos."""
    if not data_tabla: return
    headers = data_tabla.get('headers', [])
    rows = data_tabla.get('rows', [])
    
    # Crear tabla con bordes definidos
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False # Usamos ancho fijo para mejor distribución
    
    # Ancho de columnas (Ajuste para que quepa en la hoja)
    ancho_total = 26.0 # cm (considerando hoja horizontal o márgenes estrechos)
    ancho_col = ancho_total / len(headers)
    
    # --- ENCABEZADOS ---
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9) # Letra un poco más pequeña para títulos
        run.font.name = 'Arial'
        
        # Color de Fondo Gris (D9D9D9) igual a la imagen
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # --- FILAS DE CONTENIDO ---
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                cell = row_cells[i]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centrado como en la imagen
                
                # Procesar saltos de línea y negritas simuladas
                texto_str = str(cell_text)
                run = p.add_run(texto_str)
                run.font.size = Pt(8) # Letra pequeña (8pt) para que entre todo
                run.font.name = 'Arial Narrow' # Fuente más compacta
                
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph() # Espacio final

# ==========================================
# 3. LÓGICA DE MATRIZ (COMPACTA)
# ==========================================
def establecer_bordes_horizontales(table):
    """Configura bordes horizontales de la tabla."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'bottom', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:space'), '0'); border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    for border_name in ['left', 'right', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    if tblPr.find(qn('w:tblBorders')) is not None: tbl.tblPr.remove(tbl.tblPr.find(qn('w:tblBorders')))
    tblPr.append(tblBorders)

def add_compact_p(cell, text, bold=False, italic=False, color=None, bullet=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Agrega un parrafo compacto dentro de una celda."""
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.space_before = Pt(0)
    if bullet:
        p.style = 'List Bullet'
        p.paragraph_format.left_indent = Pt(10)
    run = p.add_run(str(text))
    run.font.name = 'Arial'; run.font.size = Pt(8); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color
    return p

def crear_tabla_matriz_consistencia(doc, matriz_data, titulo="Figura 2.1 Matriz de consistencia"):
    """Construye la matriz de consistencia compacta."""
    p_titulo = doc.add_paragraph(titulo)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.runs[0].bold = True; p_titulo.runs[0].font.size = Pt(11)
    
    # Subtítulo con colores
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vi = matriz_data.get('variables', {}).get('independiente', {}).get('nombre', 'VI')
    vd = matriz_data.get('variables', {}).get('dependiente', {}).get('nombre', 'VD')
    
    r1 = p2.add_run(f"{vi} "); r1.bold = True; r1.font.color.rgb = RGBColor(255, 0, 0); r1.font.size = Pt(9)
    p2.add_run("PARA REDUCIR EL ").font.size = Pt(9)
    r2 = p2.add_run(f"{vd} "); r2.bold = True; r2.font.color.rgb = RGBColor(0, 112, 192); r2.font.size = Pt(9)
    
    table = doc.add_table(rows=3, cols=5)
    establecer_bordes_horizontales(table)
    # Ajuste de anchos manual
    anchos = [Cm(3.5), Cm(3.5), Cm(3.5), Cm(3.0), Cm(3.0)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells): cell.width = anchos[idx]

    headers = ["PROBLEMAS", "OBJETIVOS", "HIPÓTESIS", "VARIABLES", "METODOLOGÍA"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "D9D9D9"); tcPr.append(shd)
        cell.text = ""; p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True; run.font.size = Pt(8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    prob = matriz_data.get('problemas', {})
    obj = matriz_data.get('objetivos', {})
    hip = matriz_data.get('hipotesis', {})
    vars_data = matriz_data.get('variables', {})
    met = matriz_data.get('metodologia', {})

    # Fila 1: General
    add_compact_p(table.rows[1].cells[0], "General:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_compact_p(table.rows[1].cells[0], prob.get('general', ''), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_compact_p(table.rows[1].cells[1], "General:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_compact_p(table.rows[1].cells[1], obj.get('general', ''), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_compact_p(table.rows[1].cells[2], "General:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_compact_p(table.rows[1].cells[2], hip.get('general', ''), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Fila 2: Específicos
    add_compact_p(table.rows[2].cells[0], "Específicos:", bold=True)
    for t in prob.get('especificos', []): add_compact_p(table.rows[2].cells[0], t, bullet=True)
    add_compact_p(table.rows[2].cells[1], "Específicos:", bold=True)
    for t in obj.get('especificos', []): add_compact_p(table.rows[2].cells[1], t, bullet=True)
    add_compact_p(table.rows[2].cells[2], "Específicos:", bold=True)
    for t in hip.get('especificos', []): add_compact_p(table.rows[2].cells[2], t, bullet=True)

    # Variables (Fusionada)
    cell_var = table.rows[1].cells[3].merge(table.rows[2].cells[3])
    add_compact_p(cell_var, "V. INDEPENDIENTE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_compact_p(cell_var, vars_data.get('independiente', {}).get('nombre', ''), bold=True, color=RGBColor(255,0,0), align=WD_ALIGN_PARAGRAPH.CENTER)
    for d in vars_data.get('independiente', {}).get('dimensiones', []): add_compact_p(cell_var, d, bullet=True)
    add_compact_p(cell_var, "")
    add_compact_p(cell_var, "V. DEPENDIENTE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_compact_p(cell_var, vars_data.get('dependiente', {}).get('nombre', ''), bold=True, color=RGBColor(0,112,192), align=WD_ALIGN_PARAGRAPH.CENTER)
    for d in vars_data.get('dependiente', {}).get('dimensiones', []): add_compact_p(cell_var, d, bullet=True)

    # Metodología (Fusionada)
    cell_met = table.rows[1].cells[4].merge(table.rows[2].cells[4])
    cell_met.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for k, v in met.items():
        if k == 'procesamiento': continue
        add_compact_p(cell_met, k.capitalize(), bold=True)
        add_compact_p(cell_met, str(v))
    
    doc.add_page_break()

# ==========================================
# 4. GENERADORES DE SECCIONES
# ==========================================

def crear_caratula_dinamica(doc, data):
    """Construye la caratula del proyecto desde el JSON."""
    c = data.get('caratula', {})
    agregar_bloque(doc, c.get('universidad', ''), negrita=True, tamano=18, despues=4)
    agregar_bloque(doc, c.get('facultad', ''), negrita=True, tamano=14, despues=4)
    agregar_bloque(doc, c.get('escuela', ''), negrita=True, tamano=14, despues=25)
    
    ruta_logo = encontrar_recurso("LogoUNAC.png")
    if ruta_logo:
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(ruta_logo, width=Inches(3.2))
        except: pass
    else:
        agregar_bloque(doc, "[LOGO UNAC]", tamano=10, antes=40, despues=40)
        
    agregar_bloque(doc, c.get('tipo_documento', ''), negrita=True, tamano=16, antes=30)
    agregar_bloque(doc, c.get('titulo_placeholder', ''), negrita=True, tamano=14, antes=30, despues=30)
    agregar_bloque(doc, c.get('frase_grado', ''), tamano=12, antes=10)
    agregar_bloque(doc, c.get('grado_objetivo', ''), negrita=True, tamano=13, despues=35)
    agregar_bloque(doc, c.get('label_autor', ''), negrita=True, tamano=12, antes=5)
    agregar_bloque(doc, c.get('label_asesor', ''), negrita=True, tamano=12, antes=5, despues=20)
    agregar_bloque(doc, c.get('label_linea', ''), tamano=11, cursiva=True, despues=40)
    agregar_bloque(doc, c.get('fecha', ''), tamano=12)
    agregar_bloque(doc, c.get('pais', ''), negrita=True, tamano=12)

def generar_pagina_respeto(doc, data):
    """Genera la pagina de respeto con notas."""
    respeto = data.get('pagina_respeto', {})
    if not respeto: return
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'; table.autofit = False; table.cell(0, 0).width = Cm(15.5)
    cell = table.cell(0, 0); cell._element.clear_content()
    for nota in respeto.get('notas', []):
        p = cell.add_paragraph(nota.get('texto', ''))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(12)
    doc.add_page_break()

def agregar_preliminares_dinamico(doc, data):
    """Genera indices y preliminares desde JSON."""
    p = data.get('preliminares', {})
    
    # 1. GENERACIÓN DE ÍNDICES (General, Tablas, Figuras, Abreviaturas)
    if 'indices' in p:
        lista_indices = p['indices']
        # Validamos si es una lista (Tu JSON actual)
        if isinstance(lista_indices, list):
            for bloque_indice in lista_indices:
                # Título del índice (Ej: ÍNDICE DE TABLAS)
                agregar_titulo_formal(doc, bloque_indice.get('titulo', ''))
                
                # Nota opcional (Ej: para Abreviaturas)
                if 'nota' in bloque_indice:
                    p_nota = doc.add_paragraph(bloque_indice['nota'])
                    p_nota.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_nota.runs[0].italic = True
                    p_nota.paragraph_format.space_after = Pt(12)

                # Generar cada línea del índice
                for item in bloque_indice.get('items', []):
                    paragraph = doc.add_paragraph()
                    paragraph.paragraph_format.space_after = Pt(3) # Espacio pequeño entre líneas
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT # Importante: Alinear izquierda para que los puntos funcionen
                    
                    # Configurar el tabulador derecho (Puntos suspensivos hasta el final)
                    # Ancho A4 (21cm) - Márgenes (2.5+3.5=6cm) = 15cm ancho útil aprox.
                    tab_stops = paragraph.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                    
                    # Escribir Texto (Ej: "I. PLANTEAMIENTO...")
                    run = paragraph.add_run(item.get('texto', ''))
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    if item.get('bold'): run.bold = True
                    
                    # Escribir Número de Página (Ej: "6")
                    if 'pag' in item:
                        run_pag = paragraph.add_run(f"\t{item['pag']}")
                        run_pag.font.name = 'Arial'
                        run_pag.font.size = Pt(11)
                
                doc.add_page_break()

def agregar_cuerpo_dinamico(doc, data):
    """Construye el cuerpo principal desde el JSON."""
    for cap in data.get('cuerpo', []):
        agregar_titulo_formal(doc, cap.get('titulo', ''), espaciado_antes=24)
        
        if 'nota_capitulo' in cap: 
            agregar_nota_estilizada(doc, cap['nota_capitulo'])
        
        for item in cap.get('contenido', []):
            # 1. Subtítulo
            if 'texto' in item:
                sub = doc.add_paragraph()
                run = sub.add_run(item.get('texto', ''))
                run.font.name = 'Arial'; run.font.size = Pt(12); run.bold = True
            
            # 2. Nota corta
            if 'nota' in item: 
                agregar_nota_guia(doc, item['nota'])
            
            # 3. Instrucción Detallada
            if 'instruccion_detallada' in item:
                agregar_nota_estilizada(doc, item['instruccion_detallada'])
            
            # 4. MATRICES AUTOMÁTICAS (2.1)
            if item.get('mostrar_matriz') == True and 'matriz_consistencia' in data:
                crear_tabla_matriz_consistencia(doc, data['matriz_consistencia'])

            # 5. MATRICES ESPECIALES (2.2, 2.3) - NUEVO BLOQUE
            if 'tablas_especiales' in item:
                for t_esp in item['tablas_especiales']:
                    p_tit = doc.add_paragraph(t_esp.get('titulo', ''))
                    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_tit.runs[0].bold = True
                    p_tit.runs[0].font.size = Pt(11)
                    crear_tabla_estilizada(doc, t_esp)

            # 6. IMÁGENES (Solo si existen)
            if 'imagenes' in item:
                for img in item['imagenes']:
                    if 'titulo' in img:
                        doc.add_paragraph(img['titulo']).paragraph_format.space_after = Pt(6)
                    
                    ruta_img = encontrar_recurso(img['ruta'])
                    if ruta_img:
                        p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.add_run().add_picture(ruta_img, width=Cm(12.0))
                    else:
                        doc.add_paragraph(f"[ERROR: Imagen no encontrada {img['ruta']}]").runs[0].font.color.rgb = RGBColor(255,0,0)
                    
                    if 'fuente' in img:
                        doc.add_paragraph(img['fuente']).paragraph_format.space_after = Pt(12)

            # 7. TABLAS NORMALES (3.2)
            if 'tabla' in item: 
                crear_tabla_estilizada(doc, item['tabla'])
            
            doc.add_paragraph("") 

        doc.add_page_break()

def agregar_finales_dinamico(doc, data):
    """Construye la seccion de finales y anexos."""
    fin = data.get('finales', {})
    
    # --- SECCIÓN REFERENCIAS ---
    if 'referencias' in fin:
        ref_data = fin['referencias']
        agregar_titulo_formal(doc, "REFERENCIAS BIBLIOGRÁFICAS")
        
        # 1. Agregar la nota instructiva (cuadro azul)
        if 'nota' in ref_data:
            agregar_nota_estilizada(doc, ref_data['nota'])
        
        # 2. Agregar los ejemplos de referencias (ESTO ES LO QUE FALTABA)
        if 'ejemplos' in ref_data:
            p_ej = doc.add_paragraph()
            p_ej.add_run("Ejemplos de formato:").bold = True
            p_ej.paragraph_format.space_after = Pt(6)
            
            for ejemplo in ref_data['ejemplos']:
                p = doc.add_paragraph(ejemplo)
                p.style = 'List Bullet' # Ponerlos con viñeta para que se vea ordenado
                p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    # --- SECCIÓN ANEXOS ---
    if 'anexos' in fin:
        agregar_titulo_formal(doc, "ANEXOS")
        for anexo in fin['anexos'].get('lista', []):
            # Título del anexo en negrita
            doc.add_paragraph(anexo.get('titulo', '')).runs[0].bold = True
            
            # Nota del anexo (si existe)
            if 'nota' in anexo:
                doc.add_paragraph(anexo['nota']).paragraph_format.space_after = Pt(12)
            else:
                doc.add_paragraph("")

def agregar_numeracion_paginas(doc):
    """Inserta numeracion de paginas en footers."""
    try:
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.clear(); run = p.add_run()
            fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    except: pass

def generar_documento_core(ruta_json, ruta_salida):
    """Orquesta la generacion completa del DOCX."""
    data = cargar_contenido(ruta_json)
    doc = Document()
    configurar_formato_unac(doc)
    crear_caratula_dinamica(doc, data)
    generar_pagina_respeto(doc, data)
    agregar_preliminares_dinamico(doc, data)
    agregar_cuerpo_dinamico(doc, data)
    agregar_finales_dinamico(doc, data)
    agregar_numeracion_paginas(doc)
    doc.save(ruta_salida)
    print(f"[OK] Generado: {ruta_salida}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        try: generar_documento_core(sys.argv[1], sys.argv[2])
        except Exception as e: print(f"[ERROR] {str(e)}"); sys.exit(1)
