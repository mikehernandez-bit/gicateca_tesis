#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GENERADOR MAESTRIA UNAC (compat + standalone)
- Un solo archivo (NO depende de generador_unac_maestria.py).
- Soporta JSON NUEVO (SPEC) y JSON VIEJO (legacy) con conversión automática.
- Numeración por secciones:
  * Carátula: sin número visible (puede "contar" para romanos según flag)
  * Hoja de respeto: blank (primera hoja de preliminares), romanos visibles
  * Índices: arábigos desde 1
  * Cuerpo: continúa arábigos

USO:
1) Modo servidor (compat antiguo):
   python generador_maestria.py <json_path> <output.docx>

2) Modo moderno:
   python generador_maestria.py --spec <json_path> --out <output.docx>

3) Modo manual:
   python generador_maestria.py

NOTA (Word):
- Para que Word calcule páginas reales en TOC/PAGEREF:
  Abrir el DOCX -> Ctrl+A -> F9 (Actualizar campos)
"""

from __future__ import annotations

import argparse
import json
import os
import platform
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# -----------------------------------------------------------------------------
# Placeholder PNG fijo para figuras "placeholder"
# -----------------------------------------------------------------------------
_PLACEHOLDER_IMAGE_PATH = Path(
    r"C:\Users\jhoan\Downloads\FORMATOTECA_UNAC_V4\FORMATOTECA_UNAC\gicateca_tesis\app\data\unac\maestria\.tmp\figura_ejemplo.png"
)

_NOTE_FILL_COLOR = "F2F8FD"
_NOTE_BORDER_COLOR = "8DB3E2"


# -----------------------------------------------------------------------------
# Config dataclasses
# -----------------------------------------------------------------------------
@dataclass
class PageSetup:
    page_width_cm: float = 21.0
    page_height_cm: float = 29.7
    margin_top_cm: float = 3.0
    margin_bottom_cm: float = 3.0
    margin_left_cm: float = 3.5
    margin_right_cm: float = 2.5


@dataclass
class FontSetup:
    name: str = "Arial"
    size_pt: int = 12


# -----------------------------------------------------------------------------
# Utilidades generales
# -----------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent


def load_json(path: Path) -> dict:
    if not path.exists():
        raise FileNotFoundError(f"JSON no encontrado: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def resolve_spec_path(spec_arg: str) -> Path:
    """
    Resuelve ruta al JSON:
    - Si existe como ruta absoluta o relativa al cwd -> ok
    - Si no, intenta ./formats/<basename> relativo al script
    """
    p = Path(spec_arg)

    if p.is_absolute():
        if p.exists():
            return p.resolve()
        raise FileNotFoundError(f"JSON no encontrado: {p}")

    if p.exists():
        return p.resolve()

    candidate = (BASE_DIR / "formats" / p.name)
    if candidate.exists():
        return candidate.resolve()

    raise FileNotFoundError(f"JSON no encontrado: {spec_arg} (ni en ./formats)")


def open_document(path: Path):
    """Solo para modo manual."""
    try:
        if platform.system() == "Windows":
            os.startfile(str(path))  # type: ignore[attr-defined]
        elif platform.system() == "Darwin":
            subprocess.run(["open", str(path)], check=False)
        else:
            subprocess.run(["xdg-open", str(path)], check=False)
    except Exception as exc:
        print(f"[WARN] No se pudo abrir el documento: {exc}")


def default_output_path(spec_path: Path, spec: dict) -> Path:
    """
    Output por defecto:
    - Si el spec trae "output_name", lo usa.
    - Si no, usa <spec_stem>.docx al lado del script.
    """
    output_name = str(spec.get("output_name", "")).strip()
    if output_name:
        out = Path(output_name)
        if out.is_absolute():
            return out
        return (BASE_DIR / out).resolve()

    return (BASE_DIR / f"{spec_path.stem}.docx").resolve()


# -----------------------------------------------------------------------------
# Detectar LEGACY JSON (viejo) y convertir a SPEC (nuevo)
# -----------------------------------------------------------------------------
def _is_legacy_json(d: dict) -> bool:
    # Heurística: el viejo tenía "structure" (lista) + "cover" con claves tipo universidad_linea
    if isinstance(d.get("structure"), list) and isinstance(d.get("cover"), dict):
        cov = d.get("cover") or {}
        return any(k in cov for k in ("universidad_linea", "grado_maestria", "asesor", "autor"))
    return False


def _legacy_to_spec(legacy: dict, legacy_path: Path) -> dict:
    """
    Convierte el JSON viejo al SPEC nuevo (mínimo viable + compat):
    - page_setup/márgenes -> page_setup nuevo
    - cover viejo -> cover.blocks nuevo
    - pre_pages -> preliminaries custom_page (romanos)
    - structure -> content.chapters/sections con {{COMPLETAR}}
    - include_list_of_tables/figures -> índices por campo (TOC \\c) para emular viejo
    """
    spec: Dict[str, Any] = {}

    ps_old = legacy.get("page_setup", {}) or {}
    margins = ps_old.get("margins_cm", {}) or {}
    font_old = ps_old.get("font", {}) or {}

    spec["page_setup"] = {
        "page_width_cm": 21.0,
        "page_height_cm": 29.7,
        "margin_top_cm": float(margins.get("top", 3.0)),
        "margin_bottom_cm": float(margins.get("bottom", 3.0)),
        "margin_left_cm": float(margins.get("left", 3.5)),
        "margin_right_cm": float(margins.get("right", 2.5)),
    }
    spec["font"] = {
        "name": str(font_old.get("name", "Arial")),
        "size_pt": int(float(font_old.get("size_pt", 12))),
    }

    # assets (logo)
    logo_path = str(legacy.get("logo_path", "")).strip()
    spec["assets"] = {"logo_path": (logo_path or "LogoUNAC.png")}

    # cover
    cov = legacy.get("cover", {}) or {}
    blocks: List[Dict[str, Any]] = []
    blocks.append({"type": "logo", "width_cm": float(cov.get("logo_width_cm", 5.5))})
    blocks.append({"type": "spacer", "lines": 1})

    def _add_cov_text(txt: str, bold=False, size_pt=12):
        if txt:
            blocks.append({"type": "text", "text": txt, "bold": bool(bold), "size_pt": int(size_pt)})

    text_size = int(cov.get("text_size_pt", 12))
    title_size = int(cov.get("title_size_pt", 14))

    _add_cov_text(str(cov.get("universidad_linea", "")).upper(), bold=True, size_pt=text_size)
    _add_cov_text(str(cov.get("unidad", "")).upper(), bold=False, size_pt=text_size)

    titulo = str(cov.get("titulo", "")).strip()
    if titulo:
        _add_cov_text(f"\"{titulo}\"".upper(), bold=True, size_pt=title_size)

    grado = str(cov.get("grado_maestria", "")).strip()
    if grado:
        _add_cov_text(f"TESIS PARA OPTAR EL GRADO ACADEMICO DE {grado}".upper(), bold=True, size_pt=text_size)

    autor = str(cov.get("autor", "")).strip()
    asesor = str(cov.get("asesor", "")).strip()
    linea = str(cov.get("linea", "")).strip()
    ciudad = str(cov.get("ciudad", "")).strip()
    anio = str(cov.get("anio", "")).strip()
    pais = str(cov.get("pais", "PERU")).strip()

    _add_cov_text(autor.upper(), bold=False, size_pt=text_size)
    _add_cov_text(asesor.upper(), bold=False, size_pt=text_size)
    if linea:
        _add_cov_text(f"LINEA DE INVESTIGACION: {linea}".upper(), bold=False, size_pt=text_size)
    if ciudad or anio:
        _add_cov_text(f"{ciudad}, {anio}".upper(), bold=False, size_pt=text_size)
    if pais:
        _add_cov_text(pais.upper(), bold=False, size_pt=text_size)

    spec["cover"] = {"blocks": blocks}

    # preliminaries
    prelims: List[Dict[str, Any]] = []

    # pre_pages viejo -> custom_page (romanos)
    for blk in (legacy.get("pre_pages", []) or []):
        prelims.append(
            {
                "type": "custom_page",
                "title": str(blk.get("title", "")).strip(),
                "title_level": int(blk.get("title_level", 4)),
                "lines": [str(x) for x in (blk.get("lines", []) or [])],
                "page_break_after": bool(blk.get("page_break_after", True)),
            }
        )

    toc_cfg = legacy.get("toc", {}) or {}
    prelims.append(
        {
            "type": "toc",
            "title": "ÍNDICE DE CONTENIDO",
            "min_level": int(toc_cfg.get("min_level", 1)),
            "max_level": int(toc_cfg.get("max_level", 3)),
        }
    )

    if bool(legacy.get("include_list_of_tables", False)):
        prelims.append({"type": "index_tables", "title": "ÍNDICE DE TABLAS", "auto_field": True})

    if bool(legacy.get("include_list_of_figures", False)):
        prelims.append({"type": "index_figures", "title": "ÍNDICE DE FIGURAS", "auto_field": True})

    spec["preliminaries"] = prelims

    # content desde structure viejo
    structure = legacy.get("structure", []) or []
    rules = legacy.get("structure_rules", {}) or {}
    add_placeholder = bool(rules.get("add_placeholder_after_heading", True))

    chapters: List[Dict[str, Any]] = []
    current_ch: Optional[Dict[str, Any]] = None

    for item in structure:
        lvl = int(item.get("level", 1))
        title = str(item.get("title", "")).strip()
        if not title:
            continue

        want_placeholder = bool(item.get("placeholder", True))
        extra_lines = [str(x) for x in (item.get("lines", []) or [])]

        if lvl == 1:
            current_ch = {"title": title, "level": 1, "guide": "", "elements": [], "sections": []}
            if add_placeholder and want_placeholder:
                current_ch["sections"].append({"title": "", "level": 2, "paragraphs": ["{{COMPLETAR}}"], "elements": []})
            if extra_lines:
                current_ch["sections"].append({"title": "", "level": 2, "paragraphs": extra_lines, "elements": []})
            chapters.append(current_ch)
        else:
            if current_ch is None:
                current_ch = {"title": "CAPÍTULO", "level": 1, "guide": "", "elements": [], "sections": []}
                chapters.append(current_ch)

            section = {"title": title, "level": lvl, "paragraphs": [], "elements": []}
            if add_placeholder and want_placeholder:
                section["paragraphs"].append("{{COMPLETAR}}")
            if extra_lines:
                section["paragraphs"].extend(extra_lines)
            current_ch["sections"].append(section)

    spec["content"] = {"chapters": chapters}

    if legacy.get("output_name"):
        spec["output_name"] = legacy.get("output_name")

    # paginación: por defecto, emulamos el comportamiento del generador grande (carátula "cuenta")
    spec.setdefault("pagination", {})
    spec["pagination"].setdefault("cover_counts_for_romans", True)
    spec["pagination"].setdefault("roman_format", "upperRoman")

    return spec


# -----------------------------------------------------------------------------
# XML helpers (fields, bordes, bookmarks, numeración)
# -----------------------------------------------------------------------------
def _add_fldSimple(paragraph, instr: str):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    fld.set(qn("w:dirty"), "true")
    paragraph._p.append(fld)

    # run dummy para que Word no “ignore” el campo en algunos casos
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "
    r.append(t)
    fld.append(r)
    return fld


def _set_table_borders(table, visible: bool):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for el in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(el)

    borders = OxmlElement("w:tblBorders")
    if not visible:
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "nil")
            borders.append(node)
        tblPr.append(borders)
        return

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), "000000")
        borders.append(node)
    tblPr.append(borders)


def _set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for mname, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{mname}"))
        if node is None:
            node = OxmlElement(f"w:{mname}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


class BookmarkManager:
    def __init__(self):
        self._next_id = 1

    def add_bookmark_around_paragraph(self, paragraph, name: str):
        bid = str(self._next_id)
        self._next_id += 1

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bid)
        start.set(qn("w:name"), name)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bid)

        p = paragraph._p
        p.insert(0, start)
        p.append(end)


def _set_section_page_numbering(section, fmt: str, start: Optional[int]):
    """fmt: 'upperRoman','lowerRoman','decimal'. start=None => continuar."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)

    fmt_lower = fmt.lower()
    if fmt_lower in ("roman", "upperroman", "upperromanformat", "upperroman_fmt", "upperroman "):
        fmt_value = "upperRoman"
    elif fmt_lower in ("lowerroman", "lowerromanformat", "lowerroman_fmt"):
        fmt_value = "lowerRoman"
    elif fmt_lower == "decimal":
        fmt_value = "decimal"
    else:
        # permite pasar "upperRoman" / "lowerRoman" directo
        fmt_value = fmt

    pgNumType.set(qn("w:fmt"), fmt_value)

    if start is None:
        pgNumType.attrib.pop(qn("w:start"), None)
    else:
        pgNumType.set(qn("w:start"), str(start))


def _clear_paragraph_runs(p):
    # Quita texto de runs existentes (y evita errores de métodos no disponibles).
    for r in p.runs:
        r.text = ""


def _add_center_page_number(section, font: FontSetup):
    footer = section.footer
    # Asegura un párrafo “limpio”
    if footer.paragraphs:
        for p in footer.paragraphs:
            _clear_paragraph_runs(p)
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.font.name = font.name
    run.font.size = Pt(font.size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)

    _add_fldSimple(p, "PAGE")


# -----------------------------------------------------------------------------
# Formatting helpers
# -----------------------------------------------------------------------------
def _apply_page_setup(doc: Document, setup: PageSetup):
    sec = doc.sections[0]
    sec.page_width = Cm(setup.page_width_cm)
    sec.page_height = Cm(setup.page_height_cm)
    sec.top_margin = Cm(setup.margin_top_cm)
    sec.bottom_margin = Cm(setup.margin_bottom_cm)
    sec.left_margin = Cm(setup.margin_left_cm)
    sec.right_margin = Cm(setup.margin_right_cm)


def _apply_styles(doc: Document, font: FontSetup):
    normal = doc.styles["Normal"]
    normal.font.name = font.name
    normal.font.size = Pt(font.size_pt)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    # Paridad con tu legacy (Heading 1..5)
    for sname, size, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 12, True),
        ("Heading 4", 12, True),
        ("Heading 5", 12, False),
    ]:
        if sname in doc.styles:
            st = doc.styles[sname]
            st.font.name = font.name
            st.font.size = Pt(size)
            st.font.bold = bold
            st.font.color.rgb = RGBColor(0, 0, 0)

    # Evita saltos automáticos raros
    if "Heading 1" in doc.styles:
        h1_pf = doc.styles["Heading 1"].paragraph_format
        h1_pf.page_break_before = False
        h1_pf.keep_with_next = False
        h1_pf.keep_together = False


def _set_paragraph_base_format(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5


def _add_spacer_lines(doc: Document, n: int = 1):
    for _ in range(n):
        p = doc.add_paragraph()
        _set_paragraph_base_format(p)


def _set_paragraph_black(p):
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)


def _add_center_title(doc: Document, text: str, size_pt: int = 12, bold: bool = True):
    p = doc.add_paragraph()
    _set_paragraph_base_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _add_heading_center_p(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    _set_paragraph_base_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_black(p)
    return p


def _paragraph_has_page_break(p) -> bool:
    for br in p._p.findall(".//w:br", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
        if br.get(qn("w:type")) in (None, "page"):
            return True
    return False


def _paragraph_has_section_break(p) -> bool:
    return p._p.find(".//w:sectPr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}) is not None


def _paragraph_has_field(p) -> bool:
    return p._p.find(".//w:fldSimple", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}) is not None


def _trim_trailing_blank_paragraphs(doc: Document):
    # Elimina párrafos en blanco al final para evitar páginas vacías.
    while doc.paragraphs:
        p = doc.paragraphs[-1]
        if p.text.strip():
            break
        if _paragraph_has_page_break(p) or _paragraph_has_section_break(p) or _paragraph_has_field(p):
            break
        p._element.getparent().remove(p._element)


def _add_page_break_if_needed(doc: Document):
    _trim_trailing_blank_paragraphs(doc)
    if doc.paragraphs and _paragraph_has_page_break(doc.paragraphs[-1]):
        return
    doc.add_page_break()


def _add_boxed_note(doc: Document, text: str, font: FontSetup):
    """Cuadro con borde, texto JUSTIFICADO."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(15.0)

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), _NOTE_FILL_COLOR)
    tc_pr.append(shd)
    tblBorders = OxmlElement("w:tblBorders")
    for b in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{b}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), _NOTE_BORDER_COLOR)
        tblBorders.append(border)
    table._tbl.tblPr.append(tblBorders)
    _set_cell_margins(cell, top=160, start=180, bottom=160, end=180)

    cell.text = ""
    paragraphs = [p.strip() for p in (text or "").split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [(text or "").strip()]

    for i, para_text in enumerate(paragraphs):
        p = cell.add_paragraph() if i else cell.paragraphs[0]
        _set_paragraph_base_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(para_text)
        run.font.name = font.name
        run.font.size = Pt(font.size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)

    _add_spacer_lines(doc, 1)


def _add_cover_boxed_note(doc: Document, text: str, font: FontSetup, width_cm: float = 15.0, size_pt: Optional[int] = None):
    """Cuadro con borde para CARÁTULA. Centrado y ancho controlado."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(width_cm)

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), _NOTE_FILL_COLOR)
    tc_pr.append(shd)
    tblBorders = OxmlElement("w:tblBorders")
    for b in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{b}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), _NOTE_BORDER_COLOR)
        tblBorders.append(border)
    table._tbl.tblPr.append(tblBorders)
    _set_cell_margins(cell, top=120, start=160, bottom=120, end=160)

    cell.text = ""
    paragraphs = [p.strip() for p in (text or "").split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [(text or "").strip()]

    fsize = int(size_pt) if size_pt is not None else int(font.size_pt)

    for i, para_text in enumerate(paragraphs):
        p = cell.add_paragraph() if i else cell.paragraphs[0]
        _set_paragraph_base_format(p)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(para_text)
        run.font.name = font.name
        run.font.size = Pt(fsize)
        run.font.color.rgb = RGBColor(0, 0, 0)

    _add_spacer_lines(doc, 1)


def _resolve_asset(spec_path: Path, rel: str) -> Path:
    rel_str = str(rel or "").strip()
    if not rel_str:
        return spec_path.parent.resolve()

    direct = Path(rel_str)
    if direct.is_absolute() and direct.exists():
        return direct

    parts = [p for p in rel_str.replace("\\", "/").split("/") if p]
    if parts:
        bases = [spec_path.parent] + list(spec_path.parents) + [BASE_DIR] + list(BASE_DIR.parents)
        for base in bases:
            if base.name == parts[0]:
                candidate = base / Path(*parts[1:])
                if candidate.exists():
                    return candidate.resolve()

    if direct.is_absolute():
        return direct
    return (spec_path.parent / rel_str).resolve()


# -----------------------------------------------------------------------------
# Cover blocks
# -----------------------------------------------------------------------------
def _is_cover_title_text(text: str) -> bool:
    if not text:
        return False
    t = text.strip().upper()
    if "TRABAJO" in t and "[" in t and "]" in t:
        return True
    if t.startswith('"[T') or t.startswith("[T"):
        return True
    return False


def _add_cover(doc: Document, spec_path: Path, cover: Dict[str, Any], assets: Dict[str, Any], font: FontSetup):
    blocks = cover.get("blocks", []) or []
    guide_text = (cover.get("guide") or cover.get("box_note_text") or "").strip()
    has_boxed_block = any(b.get("type") == "boxed_note" for b in blocks)
    auto_box = bool(guide_text) and not has_boxed_block
    box_inserted = False

    for b in blocks:
        btype = b.get("type")
        if btype == "text":
            p = doc.add_paragraph()
            _set_paragraph_base_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(b.get("text", ""))
            run.bold = bool(b.get("bold", False))
            run.font.size = Pt(int(b.get("size_pt", font.size_pt)))
            run.font.color.rgb = RGBColor(0, 0, 0)

            if auto_box and not box_inserted and _is_cover_title_text(b.get("text", "")):
                _add_cover_boxed_note(doc, guide_text, font, width_cm=float(b.get("box_width_cm", 15.0)))
                box_inserted = True

        elif btype == "spacer":
            _add_spacer_lines(doc, int(b.get("lines", 1)))

        elif btype == "logo":
            p = doc.add_paragraph()
            _set_paragraph_base_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            width_cm = float(b.get("width_cm", 5.5))
            logo_path = str(assets.get("logo_path", "LogoUNAC.png") or "LogoUNAC.png")
            logo_abs = _resolve_asset(spec_path, logo_path)
            if not logo_abs.exists():
                logo_abs = (BASE_DIR / "assets" / "LogoUNAC.png").resolve()
            try:
                p.add_run().add_picture(str(logo_abs), width=Cm(width_cm))
            except Exception:
                r = p.add_run("[LOGO UNAC]")
                r.italic = True
                r.font.color.rgb = RGBColor(0, 0, 0)

        elif btype == "boxed_note":
            note_text = b.get("text", "")
            width_cm = float(b.get("width_cm", 15.0))
            size_pt = b.get("size_pt", None)
            _add_cover_boxed_note(doc, note_text, font, width_cm=width_cm, size_pt=size_pt)

    if auto_box and not box_inserted:
        _add_cover_boxed_note(doc, guide_text, font)


# -----------------------------------------------------------------------------
# Preliminaries blocks
# -----------------------------------------------------------------------------
def _add_basic_info(doc: Document, title: str, fields: List[Dict[str, str]], font: FontSetup):
    _add_center_title(doc, title, size_pt=12, bold=True)
    _add_spacer_lines(doc, 1)

    table = doc.add_table(rows=len(fields), cols=2)
    table.autofit = False
    _set_table_borders(table, visible=False)

    for row in table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(10)

    for i, f in enumerate(fields):
        label = (f.get("label") or "").strip()
        value = (f.get("value") or "").strip()

        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c0.text = ""
        c1.text = ""

        p0 = c0.paragraphs[0]
        _set_paragraph_base_format(p0)
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(f"{label}:")
        r0.bold = True
        r0.font.name = font.name
        r0.font.size = Pt(font.size_pt)
        r0.font.color.rgb = RGBColor(0, 0, 0)

        parts = value.split("\n") if value else [""]
        for j, line in enumerate(parts):
            p1 = c1.add_paragraph() if j else c1.paragraphs[0]
            _set_paragraph_base_format(p1)
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r1 = p1.add_run(line)
            r1.font.name = font.name
            r1.font.size = Pt(font.size_pt)
            r1.font.color.rgb = RGBColor(0, 0, 0)


def _add_jury_sheet(doc: Document, title: str, jury: Dict[str, Any], acta: Dict[str, Any], font: FontSetup):
    _add_center_title(doc, title, size_pt=12, bold=True)
    _add_spacer_lines(doc, 1)

    p = doc.add_paragraph()
    _set_paragraph_base_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("MIEMBROS DE JURADO DE SUSTENTACIÓN:")
    r.bold = True
    r.font.name = font.name
    r.font.size = Pt(font.size_pt)
    r.font.color.rgb = RGBColor(0, 0, 0)

    _add_spacer_lines(doc, 1)

    members = jury.get("members", []) or [
        {"name": "[Nombre]", "role": "PRESIDENTE"},
        {"name": "[Nombre]", "role": "SECRETARIO(A)"},
        {"name": "[Nombre]", "role": "MIEMBRO"},
        {"name": "[Nombre]", "role": "MIEMBRO"},
    ]

    for m in members:
        line = f"- {m.get('name','[Nombre]')}  :  {m.get('role','MIEMBRO')}"
        p = doc.add_paragraph()
        _set_paragraph_base_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = font.name
        run.font.size = Pt(font.size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)

    _add_spacer_lines(doc, 1)

    p = doc.add_paragraph()
    _set_paragraph_base_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"ASESOR: {jury.get('advisor','[ASESOR]')}")
    run.bold = True
    run.font.name = font.name
    run.font.size = Pt(font.size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)

    _add_spacer_lines(doc, 2)

    acta_lines = [
        ("N° LIBRO", acta.get("libro", "[N°]")),
        ("N° ACTAS", acta.get("actas", "[N°]")),
        ("N° FOLIO", acta.get("folio", "[N°]")),
        ("FECHA DE APROBACIÓN DE LA TESIS", acta.get("fecha", "[dd de mes de aaaa]")),
        ("RESOLUCIÓN", acta.get("resolucion", "[Código]")),
    ]
    for k, v in acta_lines:
        p = doc.add_paragraph()
        _set_paragraph_base_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = p.add_run(f"{k}: {v}")
        rr.font.name = font.name
        rr.font.size = Pt(font.size_pt)
        rr.font.color.rgb = RGBColor(0, 0, 0)


def _add_centered_page_with_placeholder(doc: Document, title: str, placeholder: str, font: FontSetup):
    _add_center_title(doc, title, size_pt=12, bold=True)
    _add_spacer_lines(doc, 6)

    p = doc.add_paragraph()
    _set_paragraph_base_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(placeholder)
    r.italic = True
    r.font.name = font.name
    r.font.size = Pt(font.size_pt)
    r.font.color.rgb = RGBColor(0, 0, 0)


def _add_custom_page(doc: Document, title: str, title_level: int, lines: List[str], font: FontSetup):
    # Para legacy pre_pages / custom_page
    if title:
        if 1 <= title_level <= 5:
            p = doc.add_heading(title, level=title_level)
            _set_paragraph_base_format(p)
            _set_paragraph_black(p)
        else:
            _add_center_title(doc, title, size_pt=12, bold=True)

    for ln in lines:
        p = doc.add_paragraph()
        _set_paragraph_base_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(ln)
        r.font.name = font.name
        r.font.size = Pt(font.size_pt)
        r.font.color.rgb = RGBColor(0, 0, 0)


# -----------------------------------------------------------------------------
# Índices (PAGEREF + auto field)
# -----------------------------------------------------------------------------
def _add_dotline_pageref(doc: Document, label: str, bookmark: str):
    """Línea con puntos + número de página (PAGEREF)."""
    p = doc.add_paragraph()
    _set_paragraph_base_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Cm(16), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
    p.add_run(label)
    p.add_run("\t")
    _add_fldSimple(p, f"PAGEREF {bookmark} \\h")


def _add_index_with_pageref(doc: Document, entries: List[Dict[str, str]]):
    """Índice tipo UNAC: 'Pág.' + líneas con PAGEREF."""
    p = doc.add_paragraph()
    _set_paragraph_base_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Pág.")

    for e in entries or []:
        txt = str(e.get("text", "")).strip()
        bm = str(e.get("bookmark", "")).strip()
        if not (txt and bm):
            continue
        _add_dotline_pageref(doc, txt, bm)


def _add_index_auto_field(doc: Document, caption_label: str):
    """
    Lista de tablas/figuras usando campo Word:
      TOC \\c "Tabla"   o  TOC \\c "Figura"
    """
    p = doc.add_paragraph()
    _set_paragraph_base_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Pág.")

    p2 = doc.add_paragraph()
    _set_paragraph_base_format(p2)
    _add_fldSimple(p2, f'TOC \\h \\z \\c "{caption_label}"')


# -----------------------------------------------------------------------------
# Content rendering (capítulos / secciones / tablas / figuras / párrafos)
# -----------------------------------------------------------------------------
def _write_temp_placeholder_png(tmp_dir: Path) -> Path:
    return _PLACEHOLDER_IMAGE_PATH


def _add_caption_with_bookmark(doc: Document, bm: BookmarkManager, caption: str, bookmark: str, font: FontSetup):
    pcap = doc.add_paragraph()
    _set_paragraph_base_format(pcap)
    pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pcap.add_run(caption)
    run.font.name = font.name
    run.font.size = Pt(font.size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    bm.add_bookmark_around_paragraph(pcap, bookmark)


def _add_table_element(doc: Document, bm: BookmarkManager, el: Dict[str, Any], font: FontSetup):
    _add_caption_with_bookmark(doc, bm, el.get("caption", "Tabla (ejemplo)"), el.get("bookmark", "tab_x"), font)

    headers = el.get("headers") or ["Columna 1", "Columna 2"]
    rows = el.get("rows") or [["Dato", "Dato"]]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = str(h)
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.cell(i, j).text = str(val)

    note = (el.get("note") or "").strip()
    if note:
        _add_boxed_note(doc, note, font)
    else:
        _add_spacer_lines(doc, 1)


def _add_figure_element(doc: Document, bm: BookmarkManager, el: Dict[str, Any], spec_path: Path, font: FontSetup):
    _add_caption_with_bookmark(doc, bm, el.get("caption", "Figura (ejemplo)"), el.get("bookmark", "fig_x"), font)

    img = el.get("image", "placeholder")
    if img == "placeholder" or not img:
        img_path = _write_temp_placeholder_png(spec_path.parent / ".tmp")
    else:
        img_path = _resolve_asset(spec_path, str(img))

    pimg = doc.add_paragraph()
    _set_paragraph_base_format(pimg)
    pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        pimg.add_run().add_picture(str(img_path), width=Cm(14))
    except Exception:
        r = pimg.add_run("[IMAGEN]")
        r.italic = True

    note = (el.get("note") or "").strip()
    if note:
        _add_boxed_note(doc, note, font)
    else:
        _add_spacer_lines(doc, 1)


def _add_paragraphs(doc: Document, paras: List[str], font: FontSetup, boxed: bool = False):
    if boxed:
        _add_boxed_note(doc, "\n\n".join(paras), font)
        return
    for t in paras:
        p = doc.add_paragraph()
        _set_paragraph_base_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(t)
        r.font.name = font.name
        r.font.size = Pt(font.size_pt)
        r.font.color.rgb = RGBColor(0, 0, 0)


def _render_elements(doc: Document, bm: BookmarkManager, elements: List[Dict[str, Any]], spec_path: Path, font: FontSetup):
    for el in elements or []:
        et = el.get("type")
        if et == "table":
            _add_table_element(doc, bm, el, font)
        elif et == "figure":
            _add_figure_element(doc, bm, el, spec_path, font)
        elif et == "paragraphs":
            _add_paragraphs(doc, el.get("paragraphs") or [], font, boxed=False)


def _add_content(doc: Document, spec_path: Path, content: Dict[str, Any], font: FontSetup):
    bm = BookmarkManager()
    chapters = content.get("chapters", []) or []
    first_level1 = True

    for ch in chapters:
        title = (ch.get("title") or "").strip()
        if not title:
            continue
        level = int(ch.get("level", 1))

        if level == 1 and not first_level1:
            _trim_trailing_blank_paragraphs(doc)
        p = doc.add_heading(title, level=level)

        if level == 1:
            if not first_level1:
                p.paragraph_format.page_break_before = True
            p.paragraph_format.keep_with_next = False
            p.paragraph_format.keep_together = False
            first_level1 = False

        guide = (ch.get("guide") or "").strip()
        if guide:
            _add_boxed_note(doc, guide, font)

        _render_elements(doc, bm, ch.get("elements") or [], spec_path, font)

        for sec in ch.get("sections", []) or []:
            st = (sec.get("title") or "").strip()
            slevel = int(sec.get("level", level + 1))
            if st:
                doc.add_heading(st, level=slevel)

            sguide = (sec.get("guide") or "").strip()
            if sguide:
                _add_boxed_note(doc, sguide, font)

            if sec.get("paragraphs"):
                _add_paragraphs(doc, sec.get("paragraphs") or [], font, boxed=False)

            _render_elements(doc, bm, sec.get("elements") or [], spec_path, font)


def _enable_update_fields(doc: Document):
    settings = doc.settings.element
    u = settings.find(qn("w:updateFields"))
    if u is None:
        u = OxmlElement("w:updateFields")
        settings.append(u)
    u.set(qn("w:val"), "true")


# -----------------------------------------------------------------------------
# Motor principal (build_docx)
# -----------------------------------------------------------------------------
def _compute_pagination_flags(spec: Dict[str, Any]) -> Tuple[bool, str]:
    """
    cover_counts_for_romans:
      - True  => carátula "cuenta" como I (pero no se imprime). Hoja de respeto normalmente quedaría II.
      - False => romanos visibles empiezan en hoja de respeto (i). Carátula no cuenta.
    roman_format:
      - "upperRoman" o "lowerRoman"
    """
    pag = spec.get("pagination", {}) or {}
    cover_counts = bool(pag.get("cover_counts_for_romans", True))
    roman_fmt = str(pag.get("roman_format", "upperRoman")).strip() or "upperRoman"
    return cover_counts, roman_fmt


def build_docx(spec: Dict[str, Any], spec_path: Path, out_path: Path):
    doc = Document()

    ps = spec.get("page_setup", {}) or {}
    setup = PageSetup(
        page_width_cm=float(ps.get("page_width_cm", 21.0)),
        page_height_cm=float(ps.get("page_height_cm", 29.7)),
        margin_top_cm=float(ps.get("margin_top_cm", 3.0)),
        margin_bottom_cm=float(ps.get("margin_bottom_cm", 3.0)),
        margin_left_cm=float(ps.get("margin_left_cm", 3.5)),
        margin_right_cm=float(ps.get("margin_right_cm", 2.5)),
    )

    fs = spec.get("font", {}) or {}
    font = FontSetup(
        name=str(fs.get("name", "Arial")),
        size_pt=int(fs.get("size_pt", 12)),
    )

    _apply_page_setup(doc, setup)
    _apply_styles(doc, font)

    assets = spec.get("assets", {}) or {}
    prelims = spec.get("preliminaries", []) or []

    cover_counts_for_romans, roman_fmt = _compute_pagination_flags(spec)

    # ----------------
    # Sección 0: Carátula (SIN número visible)
    # ----------------
    cover_section = doc.sections[0]
    cover_section.footer.is_linked_to_previous = False

    # Si "cuenta", ponemos start=1 aquí y seguimos en prelim. Si no cuenta, igual dejamos start=1 (no se imprime).
    _set_section_page_numbering(cover_section, fmt=roman_fmt, start=1)
    _add_cover(doc, spec_path, spec.get("cover", {}) or {}, assets, font)

    # ------------------------------------------------------------
    # Sección 1: Preliminares (romanos visibles) - inicia con hoja de respeto
    # ------------------------------------------------------------
    prelim_section = doc.add_section(WD_SECTION.NEW_PAGE)
    prelim_section.footer.is_linked_to_previous = False

    if cover_counts_for_romans:
        # Continuar numeración desde carátula (I no visible, hoja de respeto sería II)
        _set_section_page_numbering(prelim_section, fmt=roman_fmt, start=None)
    else:
        # Reiniciar en hoja de respeto (i)
        _set_section_page_numbering(prelim_section, fmt=roman_fmt, start=1)

    _add_center_page_number(prelim_section, font)

    # Hoja de respeto (blank):
    # - Si el JSON trae blank_page al inicio => no duplicar
    inserted_blank = False
    if prelims and (prelims[0].get("type") == "blank_page"):
        inserted_blank = True
        # ya “existe” en el spec: creamos una página en blanco y luego seguimos
        p_blank = doc.add_paragraph("")
        _set_paragraph_base_format(p_blank)
        _add_page_break_if_needed(doc)
        prelims_iter = prelims[1:]
    else:
        # autoinsert
        p_blank = doc.add_paragraph("")
        _set_paragraph_base_format(p_blank)
        _add_page_break_if_needed(doc)
        prelims_iter = prelims

    # Separar: romanos vs índices (arábigos)
    roman_types = {"basic_info", "jury_sheet", "dedication", "acknowledgements", "custom_page"}
    index_types = {"toc", "index_tables", "index_figures", "abbreviations"}

    roman_items = [it for it in prelims_iter if it.get("type") in roman_types]
    index_items = [it for it in prelims_iter if it.get("type") in index_types]

    # Render preliminares romanos (cada uno en hoja distinta)
    for idx, item in enumerate(roman_items):
        it = item.get("type")
        if it == "basic_info":
            _add_basic_info(doc, item.get("title", "INFORMACIÓN BÁSICA"), item.get("fields", []) or [], font)
        elif it == "jury_sheet":
            _add_jury_sheet(
                doc,
                item.get("title", "HOJA DE REFERENCIA DEL JURADO Y APROBACIÓN"),
                item.get("jury", {}) or {},
                item.get("acta", {}) or {},
                font,
            )
        elif it == "dedication":
            _add_centered_page_with_placeholder(
                doc, item.get("title", "DEDICATORIA"), item.get("placeholder", "[Escriba aquí su dedicatoria...]"), font
            )
        elif it == "acknowledgements":
            _add_centered_page_with_placeholder(
                doc,
                item.get("title", "AGRADECIMIENTO"),
                item.get("placeholder", "[Escriba aquí su agradecimiento...]"),
                font,
            )
        elif it == "custom_page":
            _add_custom_page(
                doc,
                title=str(item.get("title", "")).strip(),
                title_level=int(item.get("title_level", 4)),
                lines=[str(x) for x in (item.get("lines", []) or [])],
                font=font,
            )

        # Respetar page_break_after si existe (default True para legacy)
        page_break_after = bool(item.get("page_break_after", True))
        if idx != len(roman_items) - 1 and page_break_after:
            _add_page_break_if_needed(doc)

    # ----------------
    # Sección 2: Índices (arábigos desde 1)
    # ----------------
    indices_section = doc.add_section(WD_SECTION.NEW_PAGE)
    indices_section.footer.is_linked_to_previous = False
    _set_section_page_numbering(indices_section, fmt="decimal", start=1)
    _add_center_page_number(indices_section, font)

    # Bookmarks (para líneas con PAGEREF dentro del TOC)
    bm = BookmarkManager()
    BM_TOC = "bm_indice_contenido"
    BM_LOT = "bm_indice_tablas"
    BM_LOF = "bm_indice_figuras"
    BM_ABBR = "bm_indice_abreviaturas"

    # Identifica items
    toc_item = next((x for x in index_items if x.get("type") == "toc"), {"type": "toc"})
    lot_item = next((x for x in index_items if x.get("type") == "index_tables"), None)
    lof_item = next((x for x in index_items if x.get("type") == "index_figures"), None)
    abbr_item = next((x for x in index_items if x.get("type") == "abbreviations"), None)

    # 2.1 Índice de contenido (Heading 1 para que Word lo incluya)
    toc_title = toc_item.get("title", "ÍNDICE DE CONTENIDO")
    ptitle = _add_heading_center_p(doc, toc_title, level=1)
    bm.add_bookmark_around_paragraph(ptitle, BM_TOC)

    _add_spacer_lines(doc, 1)
    p = doc.add_paragraph()
    _set_paragraph_base_format(p)
    _add_fldSimple(p, 'TOC \\o "1-3" \\h \\z \\u')
    _add_page_break_if_needed(doc)

    # 2.2 Índice de tablas
    if lot_item is not None:
        pt = _add_heading_center_p(doc, lot_item.get("title", "ÍNDICE DE TABLAS"), level=1)
        bm.add_bookmark_around_paragraph(pt, BM_LOT)

        entries = lot_item.get("entries", []) or []
        auto_field = bool(lot_item.get("auto_field", False))

        if entries:
            _add_index_with_pageref(doc, entries)
        elif auto_field:
            _add_index_auto_field(doc, "Tabla")
        else:
            # fallback seguro (si no viene nada)
            _add_index_auto_field(doc, "Tabla")

        _add_page_break_if_needed(doc)

    # 2.3 Índice de figuras
    if lof_item is not None:
        pt = _add_heading_center_p(doc, lof_item.get("title", "ÍNDICE DE FIGURAS"), level=1)
        bm.add_bookmark_around_paragraph(pt, BM_LOF)

        entries = lof_item.get("entries", []) or []
        auto_field = bool(lof_item.get("auto_field", False))

        if entries:
            _add_index_with_pageref(doc, entries)
        elif auto_field:
            _add_index_auto_field(doc, "Figura")
        else:
            _add_index_auto_field(doc, "Figura")

        _add_page_break_if_needed(doc)

    # 2.4 Abreviaturas (después de figuras)
    if abbr_item is not None:
        pt = _add_heading_center_p(doc, abbr_item.get("title", "ÍNDICE DE ABREVIATURAS"), level=1)
        bm.add_bookmark_around_paragraph(pt, BM_ABBR)

        ex = abbr_item.get("example_line", "ORGANIZACIÓN MUNDIAL DE LA SALUD (OMS)")
        p = doc.add_paragraph()
        _set_paragraph_base_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(str(ex))

        guide = abbr_item.get("guide", "Colocar el significado de las abreviaturas utilizadas en la investigación.")
        _add_boxed_note(doc, str(guide), font)

    # ----------------
    # Sección 3: Cuerpo (arábigos, continúa)
    # ----------------
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.footer.is_linked_to_previous = False
    _set_section_page_numbering(body_section, fmt="decimal", start=None)
    _add_center_page_number(body_section, font)

    _add_content(doc, spec_path, spec.get("content", {}) or {}, font)

    _enable_update_fields(doc)
    _trim_trailing_blank_paragraphs(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# -----------------------------------------------------------------------------
# generate() (compat con tu proyecto viejo)
# -----------------------------------------------------------------------------
def generate(json_path: str, output_path_override: Optional[str] = None, open_after: bool = False):
    json_file = resolve_spec_path(json_path)
    raw = load_json(json_file)

    # legacy -> spec
    if _is_legacy_json(raw):
        spec = _legacy_to_spec(raw, json_file)
    else:
        spec = raw

    if output_path_override:
        out_path = Path(output_path_override).resolve()
    else:
        out_path = default_output_path(json_file, spec)

    build_docx(spec, json_file, out_path)

    print(f"[OK] Documento guardado en: {out_path}")
    print("Abra el archivo en Word y use Ctrl+A y F9 para actualizar índices (TOC/PAGEREF).")

    if open_after:
        open_document(out_path)


# -----------------------------------------------------------------------------
# Modo manual (menú como el viejo)
# -----------------------------------------------------------------------------
def manual_cli() -> int:
    formats_dir = BASE_DIR / "formats"

    print("=" * 44)
    print("   GENERADOR MAESTRIA UNAC (CLI)")
    print("=" * 44)
    print("1. Maestría Cualitativa")
    print("2. Maestría Cuantitativa")

    try:
        op = input(">> Opcion (1/2): ").strip()
    except Exception:
        return 2

    if op == "1":
        candidates = [
            "maestria_tesis_cualitativa.json",
            "maestria_informe_cualitativa.json",
            "maestria_proyecto_cualitativa.json",
            "unac_maestria_cual.json",
            "unac_maestria_cualitativa.json",
        ]
    elif op == "2":
        candidates = [
            "maestria_tesis_cuantitativa.json",
            "maestria_informe_cuantitativa.json",
            "maestria_proyecto_cuantitativa.json",
            "unac_maestria_cuant.json",
            "unac_maestria_cuantitativa.json",
        ]
    else:
        print("Opcion no valida")
        return 2

    chosen: Optional[Path] = None
    for name in candidates:
        p = formats_dir / name
        if p.exists():
            chosen = p.resolve()
            break

    if not chosen:
        print("[ERROR] No se encontró un JSON en ./formats con nombres esperados.")
        if formats_dir.exists():
            print(f"Busca en: {formats_dir}")
            files = sorted(formats_dir.glob("*.json"))
            if files:
                print("JSON disponibles:")
                for f in files:
                    print(f" - {f.name}")
        return 1

    generate(str(chosen), output_path_override=None, open_after=True)
    return 0


# -----------------------------------------------------------------------------
# MAIN
# -----------------------------------------------------------------------------
def parse_args():
    ap = argparse.ArgumentParser(add_help=True)
    ap.add_argument("--spec", help="Ruta al JSON (SPEC nuevo o legacy)")
    ap.add_argument("--out", help="Ruta de salida .docx")
    return ap.parse_args()


def main() -> int:
    # MODO SERVIDOR (compat antiguo): script.py <json> <out>
    if len(sys.argv) > 2 and not any(a.startswith("-") for a in sys.argv[1:]):
        json_arg = sys.argv[1]
        out_arg = sys.argv[2]
        try:
            generate(json_arg, output_path_override=out_arg, open_after=False)
            return 0
        except Exception as e:
            print(f"[ERROR] Fallo critico: {e}")
            return 1

    args = parse_args()

    # MODO NUEVO (flags)
    if args.spec:
        try:
            generate(args.spec, output_path_override=args.out, open_after=False)
            return 0
        except Exception as e:
            print(f"[ERROR] Fallo critico: {e}")
            return 1

    # MODO MANUAL
    return manual_cli()


if __name__ == "__main__":
    raise SystemExit(main())
