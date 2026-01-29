
import json
import os
import sys
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Importamos las herramientas compartidas
# ADAPTACIÓN: Mapeamos las funciones requeridas a los archivos existentes
from generador_informe_tesis import (
    configurar_seccion_uni,
    configurar_estilos_base,
    agregar_bloque,
    agregar_titulo_formal,
    agregar_nota_estilizada,
    _add_fldSimple,
    _insertar_n,
    encontrar_recurso
)

def crear_caratula_posgrado_oficial(doc, data):
    """Genera la carátula según el documento 'formato posgrado.docx'"""
    c = data.get("caratula", {})
    
    # 1. Logo
    ruta_logo = encontrar_recurso("LogoUNI.png")
    if ruta_logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Ancho 2.0 inches según código usuario
        p.add_run().add_picture(ruta_logo, width=Inches(2.0))
    else:
        agregar_bloque(doc, "[LOGO UNI]", despues=12)

    # 2. Universidad y Facultad
    agregar_bloque(doc, "UNIVERSIDAD NACIONAL DE INGENIERÍA", negrita=True, tamano=16, despues=6)
    agregar_bloque(doc, c.get("facultad", "").upper(), negrita=True, tamano=14, despues=30)

    # 3. Datos Tesis
    agregar_bloque(doc, c.get("tipo_documento", ""), negrita=True, tamano=18, despues=18)
    agregar_bloque(doc, c.get("titulo", ""), negrita=True, tamano=16, despues=24)

    # 4. Grado
    agregar_bloque(doc, c.get("frase_grado", ""), tamano=12, despues=6)
    agregar_bloque(doc, c.get("grado", ""), negrita=True, tamano=13, despues=30)

    # 5. Autor y Asesor
    agregar_bloque(doc, c.get("label_autor", ""), negrita=True, tamano=12, despues=6)
    agregar_bloque(doc, c.get("autor", ""), tamano=12, despues=18)

    agregar_bloque(doc, c.get("label_asesor", ""), negrita=True, tamano=12, despues=6)
    agregar_bloque(doc, c.get("asesor", ""), tamano=12, despues=40)

    # 6. Pie
    agregar_bloque(doc, c.get("lugar_fecha", ""), negrita=True, tamano=12)
    doc.add_page_break()

def agregar_seccion_preliminar(doc, titulo, texto):
    if not titulo: return
    # Título nivel 1 pero centrado o izquierda según gusto, el doc muestra izquierda simple
    h = doc.add_heading(titulo, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.runs[0].font.name = 'Arial'; h.runs[0].font.size = Pt(14); h.runs[0].font.color.rgb = RGBColor(0,0,0)
    
    if texto:
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Si el texto está entre paréntesis (instrucción), lo ponemos en cursiva
        if texto.strip().startswith("(") and texto.strip().endswith(")"):
            p.runs[0].italic = True
    
    doc.add_page_break()

def agregar_indices_unificados(doc, data):
    indices = data.get("preliminares", {}).get("indices", {})
    
    # Índice de Contenidos
    if "contenido" in indices:
        h = doc.add_heading(indices["contenido"], level=1)
        h.runs[0].font.name = 'Arial'; h.runs[0].font.size = Pt(14); h.runs[0].font.color.rgb = RGBColor(0,0,0)
        p = doc.add_paragraph()
        _add_fldSimple(p, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_page_break()

    # Lista de Tablas/Figuras
    if "tablas" in indices:
        h = doc.add_heading(indices["tablas"], level=1)
        h.runs[0].font.name = 'Arial'; h.runs[0].font.size = Pt(14); h.runs[0].font.color.rgb = RGBColor(0,0,0)
        
        # Nota explicativa
        doc.add_paragraph("(Índice específico para ubicar rápidamente los recursos gráficos y estadísticos).").runs[0].italic = True
        
        p = doc.add_paragraph("Tablas:")
        _add_fldSimple(doc.add_paragraph(), 'TOC \\h \\z \\c "Tabla"')
        
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        p = doc.add_paragraph("Figuras/Ilustraciones:")
        _add_fldSimple(doc.add_paragraph(), 'TOC \\h \\z \\c "Figura"')
        doc.add_page_break()

def procesar_cuerpo_jerarquico(doc, cuerpo):
    for capitulo in cuerpo:
        # Título del Capítulo (Nivel 1)
        tit_cap = capitulo.get("titulo", "")
        h1 = doc.add_heading(tit_cap, level=1)
        # Nivel 1: Arial 14, Bold, Color Negro
        if h1.runs:
            h1.runs[0].font.name = 'Arial'
            h1.runs[0].font.size = Pt(14)
            h1.runs[0].font.color.rgb = RGBColor(0,0,0)
            h1.runs[0].bold = True
        
        for item in capitulo.get("items", []):
            # Subtítulos (Nivel 2 - Ej: 1.1. Identificación...)
            txt_item = item.get("titulo", "")
            h2 = doc.add_heading(txt_item, level=2)
            # Nivel 2: Arial 12, Color Negro, Regular (no bold según esquema)
            if h2.runs:
                h2.runs[0].font.name = 'Arial'
                h2.runs[0].font.size = Pt(12)
                h2.runs[0].font.color.rgb = RGBColor(0,0,0)
                h2.runs[0].bold = False 
            
            # Sub-items (Nivel 3 - Ej: 1.1.1. Formulación...)
            if "subitems" in item:
                for sub in item["subitems"]:
                    h3 = doc.add_heading(sub, level=3)
                    if h3.runs:
                        h3.runs[0].font.name = 'Arial'
                        h3.runs[0].font.size = Pt(12)
                        h3.runs[0].font.color.rgb = RGBColor(0,0,0)
                        h3.runs[0].bold = False

        doc.add_page_break()

def cargar_json(ruta):
    with open(ruta, "r", encoding="utf-8") as f:
        return json.load(f)

def main(json_path, output_path):
    # Cargar datos
    print(f"Cargando datos desde {json_path}...")
    try:
        data = cargar_json(json_path)
    except Exception as e:
        print(f"Error cargando JSON: {e}")
        return
    
    doc = Document()
    configurar_estilos_base(doc)
    
    # 1. Carátula
    configurar_seccion_uni(doc.sections[0])
    crear_caratula_posgrado_oficial(doc, data)
    
    # 2. Preliminares (Dedicatoria, Agradecimientos)
    # Comprobamos si existe la sección preliminares
    prel = data.get("preliminares", {})
    if "dedicatoria" in prel:
        agregar_seccion_preliminar(doc, prel["dedicatoria"].get("titulo"), prel["dedicatoria"].get("texto"))
    if "agradecimientos" in prel:
        agregar_seccion_preliminar(doc, prel["agradecimientos"].get("titulo"), prel["agradecimientos"].get("texto"))
    
    # 3. Índices
    agregar_indices_unificados(doc, data)
    
    # 4. Resumen e Introducción
    if "resumen" in prel:
        agregar_seccion_preliminar(doc, prel["resumen"].get("titulo"), prel["resumen"].get("texto"))
    if "introduccion" in prel:
        agregar_seccion_preliminar(doc, prel["introduccion"].get("titulo"), prel["introduccion"].get("texto"))
    
    # 5. Cuerpo (Capítulos)
    # Configurar numeración arábiga aquí si se desea reiniciar
    sect = doc.add_section(WD_SECTION.NEW_PAGE)
    configurar_seccion_uni(sect)
    pg = OxmlElement("w:pgNumType"); pg.set(qn("w:fmt"), "decimal"); pg.set(qn("w:start"), "1")
    sect._sectPr.append(pg)
    _insertar_n(sect.footer)

    procesar_cuerpo_jerarquico(doc, data.get("cuerpo", []))
    
    # 6. Finales
    fin = data.get("finales", {})
    
    # Referencias
    if "referencias" in fin:
        tit_ref = fin["referencias"]
        h_ref = doc.add_heading(tit_ref, level=1)
        if h_ref.runs:
            h_ref.runs[0].font.name = 'Arial'
            h_ref.runs[0].font.size = Pt(14)
            h_ref.runs[0].font.color.rgb = RGBColor(0,0,0)
        doc.add_paragraph("[Insertar bibliografía aquí]")
        doc.add_page_break()
    
    # Anexos
    if "anexos" in fin:
        tit_anx = fin["anexos"]
        h_anx = doc.add_heading(tit_anx, level=1)
        if h_anx.runs:
            h_anx.runs[0].font.name = 'Arial'
            h_anx.runs[0].font.size = Pt(14)
            h_anx.runs[0].font.color.rgb = RGBColor(0,0,0)
        doc.add_paragraph("[Insertar anexos aquí]")

    # Guardar
    doc.save(output_path)
    print(f"[OK] Documento generado: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        main(sys.argv[1], sys.argv[2])
    else:
        print("Uso: python generador_maestria.py <ruta_json> <ruta_salida>")
