
import json
import os
import sys
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Importamos las herramientas compartidas
# ADAPTACIÓN: Usamos generador_informe_tesis como fuente de utils
from generador_informe_tesis import (
    configurar_seccion_uni,
    configurar_estilos_base,
    agregar_bloque,
    agregar_titulo_formal,
    agregar_nota_estilizada,
    _add_fldSimple,
    encontrar_recurso
)

def crear_caratula_plan(doc, data):
    """Genera la carátula según el formato 'Plan de Trabajo de Tesis' [cite: 79-89]"""
    c = data.get("caratula", {})
    
    # 1. Logo
    ruta_logo = encontrar_recurso("LogoUNI.png")
    if ruta_logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(ruta_logo, width=Inches(2.0))
    else:
        agregar_bloque(doc, "[LOGO UNI]", despues=12)

    # 2. Institución
    agregar_bloque(doc, "UNIVERSIDAD NACIONAL DE INGENIERÍA", negrita=True, tamano=16, despues=6)
    agregar_bloque(doc, c.get("facultad", "").upper(), negrita=True, tamano=14, despues=40)

    # 3. Datos Documento
    agregar_bloque(doc, c.get("tipo_documento", "").upper(), negrita=True, tamano=16, despues=24)
    agregar_bloque(doc, c.get("titulo", ""), negrita=True, tamano=16, despues=24)

    # 4. Grado Objetivo
    agregar_bloque(doc, c.get("frase_grado", ""), tamano=12, despues=6)
    agregar_bloque(doc, c.get("carrera", "").upper(), negrita=True, tamano=14, despues=40)

    # 5. Autor y Asesor
    agregar_bloque(doc, c.get("label_autor", ""), negrita=True, tamano=12, despues=6)
    agregar_bloque(doc, c.get("autor", ""), tamano=12, despues=24)

    agregar_bloque(doc, c.get("label_asesor", ""), negrita=True, tamano=12, despues=6)
    agregar_bloque(doc, c.get("asesor", ""), tamano=12, despues=50)

    # 6. Pie
    agregar_bloque(doc, c.get("lugar_fecha", ""), negrita=True, tamano=12)
    doc.add_page_break()

def generar_indice_simple(doc, secciones):
    """Genera la lista numerada de las 12 secciones [cite: 90-102]"""
    agregar_titulo_formal(doc, "ÍNDICE", espaciado_antes=0)
    
    for sec in secciones:
        num = sec.get("numero", "")
        titulo = sec.get("titulo", "")
        
        p = doc.add_paragraph()
        # Tabulador derecho para el número de página (simulado o automático)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT, 1) # 1 = Puntos
        
        run = p.add_run(f"{num}.\t{titulo}")
        run.bold = False
        run.font.name = 'Arial'; run.font.size = Pt(11)
        
        # Agregamos el tabulador y el campo de página (o un placeholder)
        p.add_run("\t")
        _add_fldSimple(p, "PAGE") # Número de página actual (referencial)

    doc.add_page_break()

def procesar_secciones(doc, secciones):
    """Itera y renderiza los 12 puntos del plan."""
    for sec in secciones:
        num = sec.get("numero", "")
        titulo = sec.get("titulo", "")
        
        # Título Nivel 1 (Ej: 1. TÍTULO)
        texto_tit = f"{num}. {titulo}"
        h = doc.add_heading(texto_tit, level=1)
        if h.runs:
            h.runs[0].font.name = 'Arial'
            h.runs[0].font.size = Pt(12)
            h.runs[0].font.color.rgb = RGBColor(0,0,0)
            h.runs[0].bold = True
        h.paragraph_format.space_after = Pt(12)
        
        # Instrucción (Cuadro guía)
        if "instruccion" in sec:
            agregar_nota_estilizada(doc, sec["instruccion"])
            
        # Listas (para items 8, 9, 10 que tienen subpuntos)
        if "lista" in sec:
            for item in sec["lista"]:
                p = doc.add_paragraph(item)
                p.style = 'List Bullet' # Viñeta estándar de Word
                p.paragraph_format.space_after = Pt(6)
        
        # Espacio para escritura del alumno
        doc.add_paragraph() 
        doc.add_paragraph()

def main(json_path, output_path):
    print(f"Cargando datos desde {json_path}...")
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        print(f"Error cargando JSON: {e}")
        return
    
    doc = Document()
    configurar_estilos_base(doc)
    
    # 1. Carátula
    configurar_seccion_uni(doc.sections[0])
    crear_caratula_plan(doc, data)
    
    # 2. Índice
    generar_indice_simple(doc, data.get("secciones", []))
    
    # 3. Cuerpo del Plan
    procesar_secciones(doc, data.get("secciones", []))
    
    # Guardar
    doc.save(output_path)
    print(f"[OK] Documento generado exitosamente: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        main(sys.argv[1], sys.argv[2])
    else:
        print("Uso: python generador_proyecto_tesis.py <input.json> <output.docx>")
